"""Tool implementations — domain: file_ops."""
from __future__ import annotations

import logging
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from layla.tools.sandbox_core import (
    _SHELL_BLOCKLIST,
    _SHELL_INJECTION_WARN,
    _SHELL_NETWORK_DENYLIST,
    _agent_registry_dir,
    _check_read_freshness,
    _clear_read_freshness,
    _effective_sandbox,
    _get_sandbox,
    _maybe_file_checkpoint,
    _set_read_freshness,
    _shell_executable_base,
    _write_file_limits,
    inside_sandbox,
    shell_command_is_safe_whitelisted,
    shell_command_line,
)

logger = logging.getLogger("layla")

# Injected by layla.tools.registry with the assembled TOOLS dict (same object in every module).
TOOLS: dict = {}
def write_file(path: str, content: str) -> dict:
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    stale = _check_read_freshness(target)
    if stale:
        return {"ok": False, "error": stale, "hint": "use read_file first"}
    try:
        max_bytes, explosion = _write_file_limits()
        raw = (content or "").encode("utf-8", errors="replace")
        if len(raw) > max_bytes:
            return {"ok": False, "error": "content_too_large", "limit_bytes": max_bytes}
        if target.exists():
            try:
                existing_size = target.stat().st_size
                new_size = len(raw)
                if existing_size > 0 and new_size > existing_size * explosion:
                    return {
                        "ok": False,
                        "error": "size_explosion_detected",
                        "existing_bytes": existing_size,
                        "new_bytes": new_size,
                    }
            except Exception:
                pass
    except Exception:
        pass
    _maybe_file_checkpoint(target, "write_file")
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    _clear_read_freshness(target)
    return {"ok": True, "path": str(target)}

def write_files_batch(files: list) -> dict:
    """
    Write multiple files atomically. files: [{path, content}, ...].
    Returns approval_required when gated; otherwise applies all and returns ok.
    """
    if not isinstance(files, list) or not files:
        return {"ok": False, "error": "files must be a non-empty list of {path, content}"}
    written = []
    errors = []
    for i, item in enumerate(files):
        if not isinstance(item, dict):
            errors.append(f"Item {i}: not a dict")
            continue
        path = (item.get("path") or "").strip()
        content = item.get("content", "")
        if not path:
            errors.append(f"Item {i}: missing path")
            continue
        target = Path(path)
        if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
            target = (Path(_effective_sandbox.path) / path).resolve()
        if not inside_sandbox(target):
            errors.append(f"{path}: outside sandbox")
            continue
        try:
            _maybe_file_checkpoint(target, "write_files_batch")
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(content, encoding="utf-8")
            written.append(str(target))
        except Exception as e:
            errors.append(f"{path}: {e}")
    if errors:
        return {"ok": False, "error": "; ".join(errors[:5]), "written": written}
    return {"ok": True, "written": written, "count": len(written)}

def read_file(path: str) -> dict:
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    if not target.is_file():
        return {"ok": False, "error": "Not a file"}
    try:
        content = target.read_text(encoding="utf-8", errors="replace")
        _set_read_freshness(target)
        return {"ok": True, "path": str(target), "content": content[:8000]}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def list_dir(path: str) -> dict:
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}
    try:
        entries = []
        for item in sorted(target.iterdir()):
            entries.append({
                "name": item.name,
                "type": "dir" if item.is_dir() else "file",
            })
        return {"ok": True, "path": str(target), "entries": entries}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def glob_files(pattern: str, root: str) -> dict:
    root_path = Path(root)
    if not root_path.is_absolute() and getattr(_effective_sandbox, "path", None):
        root_path = (Path(_effective_sandbox.path) / root).resolve()
    if not inside_sandbox(root_path):
        return {"ok": False, "error": "Outside sandbox"}
    if not root_path.exists():
        return {"ok": False, "error": "Path not found"}
    try:
        matches = [str(p) for p in root_path.rglob(pattern)][:200]
        return {"ok": True, "matches": matches}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def sync_repo_cognition(
    workspace_roots: str | list | None = None,
    index_semantic: bool = False,
) -> dict:
    """
    Scan canonical docs under each workspace root, build a deterministic cognition pack,
    and persist it for system-head injection (multi-repo supported). Sandbox-restricted.
    workspace_roots: list of paths, JSON array string, or comma-separated paths; default = current sandbox root.
    index_semantic: also run workspace semantic index (slower).
    """
    import json as _json

    roots: list[str] = []
    if workspace_roots is None:
        roots = [str(_get_sandbox())]
    elif isinstance(workspace_roots, list):
        roots = [str(x).strip() for x in workspace_roots if str(x).strip()]
    elif isinstance(workspace_roots, str):
        s = workspace_roots.strip()
        if s.startswith("["):
            try:
                roots = [str(x).strip() for x in _json.loads(s) if str(x).strip()]
            except Exception:
                roots = []
        else:
            roots = [p.strip() for p in s.split(",") if p.strip()]
    if not roots:
        return {"ok": False, "error": "no workspace roots"}
    for r in roots:
        p = Path(r).expanduser().resolve()
        if not inside_sandbox(p):
            return {"ok": False, "error": f"Outside sandbox: {r}"}
    try:
        from services.repo_cognition import sync_repo_cognition as _sync

        return _sync(roots, index_semantic=index_semantic)
    except Exception as e:
        return {"ok": False, "error": str(e)}

def scan_repo(workspace_root: str = "", dry_run: bool = False, max_files: int = 0) -> dict:
    """
    Scan workspace tree, write `.layla/project_memory.json` (file map + structure).
    Merges with existing memory; preserves purpose/complexity/issues when re-scanning.
    Requires allow_write + approval (writes under workspace).
    """
    import runtime_safety

    from services import project_memory as pm

    cfg = runtime_safety.load_config()
    wr = (workspace_root or "").strip()
    root = Path(wr).expanduser().resolve() if wr else _get_sandbox()
    if not inside_sandbox(root):
        return {"ok": False, "error": "Outside sandbox"}
    if not root.is_dir():
        return {"ok": False, "error": "Path not found or not a directory"}
    mf = int(max_files) if int(max_files or 0) > 0 else int(cfg.get("project_memory_max_file_entries", 500) or 500)
    mb = int(cfg.get("project_memory_max_bytes", pm.DEFAULT_MAX_BYTES) or pm.DEFAULT_MAX_BYTES)
    return pm.scan_workspace_into_memory(root, dry_run=bool(dry_run), max_files=max(1, min(5000, mf)), max_bytes=mb)

def update_project_memory(workspace_root: str = "", patch: dict | None = None) -> dict:
    """Merge a JSON patch into `.layla/project_memory.json` (plan, todos, decisions, file notes)."""
    import runtime_safety

    from services import project_memory as pm

    cfg = runtime_safety.load_config()
    wr = (workspace_root or "").strip()
    root = Path(wr).expanduser().resolve() if wr else _get_sandbox()
    if not inside_sandbox(root):
        return {"ok": False, "error": "Outside sandbox"}
    p = patch if isinstance(patch, dict) else {}
    base = pm.load_project_memory(root) or pm.empty_document(str(root))
    mf = int(cfg.get("project_memory_max_file_entries", 500) or 500)
    ml = int(cfg.get("project_memory_max_list_entries", 200) or 200)
    merged = pm.merge_patch(base, p, max_files=max(1, min(5000, mf)), max_list=max(10, min(2000, ml)))
    mb = int(cfg.get("project_memory_max_bytes", pm.DEFAULT_MAX_BYTES) or pm.DEFAULT_MAX_BYTES)
    ok, err = pm.save_project_memory(root, merged, max_bytes=mb)
    if not ok:
        return {"ok": False, "error": err or "save_failed"}
    return {"ok": True, "path": str(pm.memory_file_path(root))}

def parse_gcode(path: str) -> dict:
    """
    Parse G-code / NC file: moves, tools, units, bounds, feed rates.
    Supports .gcode, .nc, .tap, .sbp.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        content = target.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return {"ok": False, "error": str(e)}
    # Parse G-code patterns
    tools = set()
    feed_rates = []
    bounds = {"x": [], "y": [], "z": []}
    units = "mm"  # default
    for line in content.splitlines():
        line = line.strip()
        if not line or line.startswith(";"):
            continue
        # G20 = inches, G21 = mm
        if "G20" in line.upper():
            units = "in"
        elif "G21" in line.upper():
            units = "mm"
        # Tool: T1, T02, M6 T1
        for m in re.finditer(r"\bT(\d+)\b", line, re.I):
            tools.add(int(m.group(1)))
        # Feed: F3000, F100.5
        for m in re.finditer(r"\bF([\d.]+)\b", line, re.I):
            feed_rates.append(float(m.group(1)))
        # Moves: G0/G1 X Y Z
        for m in re.finditer(r"\b[Gg]?[01]\b.*\b([XYZxyz])([-]?[\d.]+)", line):
            axis = m.group(1).lower()
            val = float(m.group(2))
            if axis in bounds:
                bounds[axis].append(val)
    # Summarize bounds
    summary = {}
    for ax, vals in bounds.items():
        if vals:
            summary[ax] = {"min": min(vals), "max": max(vals), "count": len(vals)}
    return {
        "ok": True,
        "path": str(target),
        "units": units,
        "tools": sorted(tools) if tools else [],
        "move_count": sum(len(bounds[ax]) for ax in bounds),
        "feed_rates": list(set(feed_rates))[:20] if feed_rates else [],
        "bounds": summary,
        "lines": len([ln for ln in content.splitlines() if ln.strip() and not ln.strip().startswith(";")]),
    }

def stl_mesh_info(path: str) -> dict:
    """STL mesh stats: vertex count, bounds, volume. Requires trimesh or numpy."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import trimesh
        mesh = trimesh.load(str(target))
        if isinstance(mesh, trimesh.Scene):
            mesh = mesh.dump(concatenate=True)
        bounds = mesh.bounds
        vol = float(mesh.volume) if hasattr(mesh, "volume") else None
        return {
            "ok": True,
            "path": str(target),
            "vertices": int(len(mesh.vertices)),
            "faces": int(len(mesh.faces)) if hasattr(mesh, "faces") else None,
            "bounds": {"min": bounds[0].tolist(), "max": bounds[1].tolist()},
            "volume": vol,
        }
    except ImportError:
        # Fallback: count vertices from ASCII STL
        try:
            content = target.read_text(encoding="utf-8", errors="replace")
            if "vertex" in content.lower():
                verts = re.findall(r"vertex\s+([-\d.e]+)\s+([-\d.e]+)\s+([-\d.e]+)", content, re.I)
                xs, ys, zs = zip(*[[float(a), float(b), float(c)] for a, b, c in verts]) if verts else ([], [], [])
                return {
                    "ok": True,
                    "path": str(target),
                    "vertices": len(verts),
                    "bounds": {"min": [min(xs), min(ys), min(zs)], "max": [max(xs), max(ys), max(zs)]} if verts else {},
                    "fallback": "ascii_parse",
                }
        except Exception as e:
            return {"ok": False, "error": str(e), "hint": "pip install trimesh for full support"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def tail_file(path: str, n: int = 50) -> dict:
    """Return last n lines of a file. Useful for logs."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        with open(target, encoding="utf-8", errors="replace") as f:
            lines = f.readlines()
        tail = lines[-n:] if len(lines) > n else lines
        return {"ok": True, "path": str(target), "lines": "".join(tail)[:8000], "total_lines": len(lines)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def clipboard_read() -> dict:
    """Read text from system clipboard. Requires pyperclip."""
    try:
        import pyperclip
        text = pyperclip.paste()
        return {"ok": True, "text": (text or "")[:10000]}
    except ImportError:
        return {"ok": False, "error": "pyperclip not installed: pip install pyperclip"}

def clipboard_write(text: str) -> dict:
    """Write text to system clipboard. Requires pyperclip."""
    try:
        import pyperclip
        pyperclip.copy(text[:50000])
        return {"ok": True, "length": len(text)}
    except ImportError:
        return {"ok": False, "error": "pyperclip not installed: pip install pyperclip"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def search_replace(root: str, find: str, replace: str, file_glob: str = "*", dry_run: bool = True) -> dict:
    """
    Multi-file find/replace. dry_run=True lists matches without changing. Uses regex if find contains regex chars.
    """
    root_path = Path(root)
    if not inside_sandbox(root_path):
        return {"ok": False, "error": "Outside sandbox"}
    if not root_path.exists():
        return {"ok": False, "error": "Path not found"}
    use_regex = bool(re.search(r"[.*+?^${}()|[\]\\]", find))
    pattern = re.compile(find) if use_regex else None
    matches = []
    for f in root_path.rglob(file_glob):
        if not f.is_file():
            continue
        try:
            content = f.read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue
        if use_regex:
            new_content, n = pattern.subn(replace, content)
        else:
            n = content.count(find)
            new_content = content.replace(find, replace) if n else content
        if n:
            matches.append({"path": str(f), "count": n})
            if not dry_run:
                _maybe_file_checkpoint(f, "search_replace")
                f.write_text(new_content, encoding="utf-8")
    return {"ok": True, "dry_run": dry_run, "matches": matches[:100], "total_files": len(matches)}

def apply_patch(original_path: str, patch_text: str) -> dict:
    """Apply a unified diff patch using unidiff (pure Python, Windows-safe). Creates a backup first."""
    target = Path(original_path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    stale = _check_read_freshness(target)
    if stale:
        return {"ok": False, "error": stale, "hint": "use read_file first"}
    _maybe_file_checkpoint(target, "apply_patch")
    import shutil

    from layla.time_utils import utcnow
    backup = target.with_suffix(
        f".bak_{utcnow().strftime('%Y%m%d_%H%M%S')}{target.suffix}"
    )
    shutil.copy2(str(target), str(backup))
    try:
        import unidiff
        patch_set = unidiff.PatchSet(patch_text.splitlines(keepends=True))
        original_lines = target.read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
        result_lines = list(original_lines)
        for patched_file in patch_set:
            offset = 0
            for hunk in patched_file:
                src_start = hunk.source_start - 1 + offset
                removed = [line.value for line in hunk if line.is_removed]
                added = [line.value for line in hunk if line.is_added]
                # Remove old lines, insert new ones
                result_lines[src_start: src_start + len(removed)] = added
                offset += len(added) - len(removed)
        target.write_text("".join(result_lines), encoding="utf-8")
        _clear_read_freshness(target)
        return {"ok": True, "path": str(target), "backup": str(backup)}
    except Exception as e:
        return {"ok": False, "error": str(e), "backup": str(backup)}

def file_info(path: str) -> dict:
    """Return size, line count (approx), and whether file looks text. Read-only; no approval."""
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}
    if not target.is_file():
        return {"ok": False, "error": "Not a file"}
    try:
        size = target.stat().st_size
        # Sample first 8k to guess text vs binary and count newlines
        raw = target.read_bytes()
        sample = raw[:8192]
        try:
            sample.decode("utf-8", errors="strict")
            is_text = True
        except Exception:
            is_text = False
        line_count = sample.count(b"\n") if is_text else None
        note = "from first 8k only" if is_text and len(raw) > len(sample) else None
        return {"ok": True, "path": str(target), "size_bytes": size, "is_text": is_text, "line_count_sample": line_count, "note": note}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def understand_file_tool(path: str, content: str | None = None) -> dict:
    """Interpret file intent (read-only). path: file path; optional content for in-memory analysis."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.file_understanding import analyze_file
        if content is not None:
            return {"ok": True, **analyze_file(file_path=path, content=content)}
        return {"ok": True, **analyze_file(file_path=path)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def json_query(path: str, query: str = "") -> dict:
    """
    Parse a JSON file and optionally extract a value by dot-notation path.
    query examples: 'key', 'nested.key', 'array.0.field'.
    If no query, returns the full parsed object (truncated).
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import json as _json
        data = _json.loads(target.read_text(encoding="utf-8"))
        if not query:
            return {"ok": True, "data": str(data)[:4000]}
        parts = query.split(".")
        val = data
        for p in parts:
            if isinstance(val, dict):
                val = val[p]
            elif isinstance(val, list):
                val = val[int(p)]
            else:
                return {"ok": False, "error": f"Cannot traverse into {type(val).__name__} at '{p}'"}
        return {"ok": True, "result": val, "result_str": str(val)[:2000]}
    except (KeyError, IndexError) as e:
        return {"ok": False, "error": f"Path not found: {e}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def diff_files(path_a: str, path_b: str) -> dict:
    """Diff two text files. Returns unified diff."""
    for p in (path_a, path_b):
        t = Path(p)
        if not inside_sandbox(t):
            return {"ok": False, "error": f"Outside sandbox: {p}"}
        if not t.exists():
            return {"ok": False, "error": f"File not found: {p}"}
    try:
        import difflib
        a_lines = Path(path_a).read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
        b_lines = Path(path_b).read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
        diff = list(difflib.unified_diff(a_lines, b_lines, fromfile=path_a, tofile=path_b, n=3))
        return {"ok": True, "diff": "".join(diff)[:8000], "changed": len(diff) > 0}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def read_pdf(path: str, max_pages: int = 30) -> dict:
    """
    Extract text from a PDF file using PyMuPDF.
    Returns text per page up to max_pages. Falls back to pypdf if fitz not installed.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    # Try PyMuPDF (fitz) first
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(target))
        try:
            pages = []
            for i, page in enumerate(doc):
                if i >= max_pages:
                    break
                pages.append(f"--- Page {i+1} ---\n{page.get_text()}")
            full = "\n".join(pages)
            return {"ok": True, "path": str(target), "pages": min(len(pages), max_pages), "text": full[:12000]}
        finally:
            doc.close()
    except ImportError:
        pass
    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(target))
        pages = []
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            pages.append(f"--- Page {i+1} ---\n{page.extract_text() or ''}")
        full = "\n".join(pages)
        return {"ok": True, "path": str(target), "pages": min(len(reader.pages), max_pages), "text": full[:12000]}
    except ImportError:
        return {"ok": False, "error": "PDF reading requires PyMuPDF or pypdf: pip install PyMuPDF"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def read_docx(path: str) -> dict:
    """
    Read a Word document (.docx). Returns full text, paragraph list, and table data.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        from docx import Document
        doc = Document(str(target))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables[:10]:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(rows)
        full_text = "\n".join(paragraphs)
        return {
            "ok": True, "path": str(target),
            "paragraphs": len(paragraphs),
            "text": full_text[:10000],
            "tables": tables[:5],
            "table_count": len(doc.tables),
        }
    except ImportError:
        return {"ok": False, "error": "python-docx not installed: pip install python-docx"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def read_excel(path: str, sheet: str = "", max_rows: int = 100) -> dict:
    """
    Read an Excel file (.xlsx/.xls). Returns sheet names, data from target sheet,
    and basic stats. sheet: sheet name or index (default: first sheet).
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import pandas as _pd
        xl = _pd.ExcelFile(str(target))
        sheet_names = xl.sheet_names
        active_sheet = sheet if sheet else sheet_names[0]
        df = xl.parse(active_sheet)
        result: dict = {
            "ok": True, "path": str(target),
            "sheets": sheet_names, "active_sheet": str(active_sheet),
            "rows": len(df), "columns": list(df.columns),
            "sample": df.head(max_rows).to_dict(orient="records"),
            "null_counts": df.isnull().sum().to_dict(),
        }
        try:
            result["stats"] = df.describe().to_dict()
        except Exception:
            pass
        return result
    except ImportError:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(target), read_only=True, data_only=True)
            sheet_names = wb.sheetnames
            ws = wb[sheet] if sheet and sheet in sheet_names else wb.active
            rows = []
            headers: list = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i == 0:
                    headers = [str(c or f"col_{j}") for j, c in enumerate(row)]
                elif i <= max_rows:
                    rows.append(dict(zip(headers, row)))
            wb.close()
            return {"ok": True, "path": str(target), "sheets": sheet_names, "columns": headers, "sample": rows}
        except ImportError:
            return {"ok": False, "error": "Excel reading requires pandas or openpyxl: pip install pandas openpyxl"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def workspace_map(root: str = "", max_files: int = 500, include_content_preview: bool = False) -> dict:
    """
    Build a full intelligence map of a workspace:
    - Directory tree (depth-limited)
    - File count by extension
    - Detected tech stack (languages, frameworks, config files)
    - Entry points (main.py, index.js, Dockerfile, etc.)
    - Key documentation (README, CHANGELOG, AGENTS.md, etc.)
    - Large files + recently modified files
    """
    root_path = Path(root).expanduser().resolve() if root else Path.cwd()
    if not root_path.exists():
        return {"ok": False, "error": "Path not found"}

    IGNORE = {".git", ".venv", "venv", "__pycache__", "node_modules", ".mypy_cache",
              ".pytest_cache", "dist", "build", ".tox", "*.egg-info"}

    all_files: list[Path] = []
    for f in root_path.rglob("*"):
        if f.is_file() and not any(part in IGNORE for part in f.parts):
            all_files.append(f)
            if len(all_files) >= max_files:
                break

    # Extension counts
    ext_counts: dict = {}
    for f in all_files:
        ext = f.suffix.lower() or "(none)"
        ext_counts[ext] = ext_counts.get(ext, 0) + 1

    # Tech stack detection
    STACK_SIGNALS = {
        "Python": {".py", "requirements.txt", "pyproject.toml", "setup.py", "Pipfile"},
        "JavaScript": {".js", ".mjs", "package.json"},
        "TypeScript": {".ts", ".tsx", "tsconfig.json"},
        "Rust": {".rs", "Cargo.toml"},
        "Go": {".go", "go.mod"},
        "Java": {".java", "pom.xml", "build.gradle"},
        "C/C++": {".c", ".cpp", ".h", ".hpp", "CMakeLists.txt"},
        "Docker": {"Dockerfile", "docker-compose.yml", "docker-compose.yaml"},
        "Kubernetes": {".yaml", "k8s", "helm"},
        "FastAPI": {"main.py"},
        "React": {"package.json", ".jsx", ".tsx"},
        "Vue": {".vue"},
    }
    names_and_exts = {f.name for f in all_files} | {f.suffix.lower() for f in all_files}
    detected_stack = [lang for lang, signals in STACK_SIGNALS.items() if signals & names_and_exts]

    # Entry points
    ENTRY_NAMES = {"main.py", "app.py", "server.py", "index.js", "index.ts", "main.rs",
                   "main.go", "Dockerfile", "docker-compose.yml", "manage.py", "wsgi.py", "asgi.py"}
    entry_points = [str(f.relative_to(root_path)) for f in all_files if f.name in ENTRY_NAMES]

    # Key docs
    DOC_NAMES = {"README.md", "AGENTS.md", "ARCHITECTURE.md", "CHANGELOG.md", "CONTRIBUTING.md",
                 "LICENSE", "INSTALL.md", "SECURITY.md", "TODO.md", "NOTES.md"}
    key_docs = {}
    for f in all_files:
        if f.name in DOC_NAMES:
            preview = ""
            if include_content_preview:
                try:
                    preview = f.read_text(encoding="utf-8", errors="replace")[:400]
                except Exception:
                    pass
            key_docs[f.name] = {"path": str(f.relative_to(root_path)), "preview": preview}

    # Largest files
    sorted_by_size = sorted(all_files, key=lambda f: f.stat().st_size, reverse=True)
    largest = [{"path": str(f.relative_to(root_path)), "size_kb": round(f.stat().st_size / 1024, 1)} for f in sorted_by_size[:10]]

    # Recently modified
    sorted_by_mtime = sorted(all_files, key=lambda f: f.stat().st_mtime, reverse=True)
    recent = [{"path": str(f.relative_to(root_path)), "modified": str(__import__("datetime").datetime.fromtimestamp(f.stat().st_mtime))[:16]} for f in sorted_by_mtime[:10]]

    # Directory tree (2-level)
    def tree_level(path: Path, depth: int = 0, max_depth: int = 2) -> list:
        if depth >= max_depth:
            return []
        entries = []
        try:
            for child in sorted(path.iterdir(), key=lambda x: (x.is_file(), x.name)):
                if any(part in IGNORE for part in child.parts):
                    continue
                entry = {"name": child.name, "type": "file" if child.is_file() else "dir"}
                if child.is_dir() and depth < max_depth - 1:
                    entry["children"] = tree_level(child, depth + 1, max_depth)
                entries.append(entry)
        except PermissionError:
            pass
        return entries

    return {
        "ok": True,
        "root": str(root_path),
        "total_files": len(all_files),
        "extensions": dict(sorted(ext_counts.items(), key=lambda x: -x[1])[:20]),
        "tech_stack": detected_stack,
        "entry_points": entry_points,
        "key_docs": key_docs,
        "largest_files": largest,
        "recently_modified": recent,
        "tree": tree_level(root_path),
    }

def yaml_read(path: str) -> dict:
    """Parse a YAML file. Uses PyYAML (pip install pyyaml); simple fallback if not installed."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import yaml as _yaml
        with open(str(target), encoding="utf-8") as f:
            data = _yaml.safe_load(f)
        return {"ok": True, "path": str(target), "data": data, "type": type(data).__name__}
    except ImportError:
        import re as _re
        text = target.read_text(encoding="utf-8", errors="replace")
        data = {}
        for line in text.splitlines():
            m = _re.match(r'^(\w[\w\-\.]*)\s*:\s*(.+)$', line.strip())
            if m:
                data[m.group(1)] = m.group(2).strip()
        return {"ok": True, "path": str(target), "data": data, "note": "Basic parse â€” install pyyaml: pip install pyyaml"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def xml_parse(path_or_text: str) -> dict:
    """Parse XML from a file path or raw string. Returns tree structure (up to 3 levels)."""
    import xml.etree.ElementTree as _ET

    def _to_dict(elem: _ET.Element, depth: int = 0) -> dict:
        node: dict = {"tag": elem.tag, "attribs": dict(elem.attrib)}
        text = (elem.text or "").strip()
        if text:
            node["text"] = text[:500]
        if depth < 3:
            children = [_to_dict(c, depth+1) for c in list(elem)[:20]]
            if children:
                node["children"] = children
        return node

    try:
        p = Path(path_or_text)
        if inside_sandbox(p) and p.exists():
            root = _ET.parse(str(p)).getroot()
            source = "file"
        else:
            root = _ET.fromstring(path_or_text)
            source = "string"
        return {"ok": True, "source": source, "root": _to_dict(root), "element_count": sum(1 for _ in root.iter())}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def hash_file(path: str, algorithm: str = "sha256") -> dict:
    """Compute cryptographic hash of a file. algorithm: md5 | sha1 | sha256 | sha512."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    import hashlib as _hl
    algo = algorithm.lower().replace("-", "")
    if algo not in ("md5", "sha1", "sha256", "sha512"):
        return {"ok": False, "error": "Use md5/sha1/sha256/sha512"}
    try:
        h = _hl.new(algo)
        with open(str(target), "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
        return {"ok": True, "path": str(target), "algorithm": algo, "hash": h.hexdigest(), "size_bytes": target.stat().st_size}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def base64_tool(data: str, mode: str = "encode", encoding: str = "utf-8") -> dict:
    """Base64 encode/decode. mode: encode | decode | encode_url | decode_url."""
    import base64 as _b64
    try:
        if mode == "encode":
            result = _b64.b64encode(data.encode(encoding)).decode("ascii")
        elif mode == "decode":
            result = _b64.b64decode(data).decode(encoding, errors="replace")
        elif mode == "encode_url":
            result = _b64.urlsafe_b64encode(data.encode(encoding)).decode("ascii")
        elif mode == "decode_url":
            result = _b64.urlsafe_b64decode(data).decode(encoding, errors="replace")
        else:
            return {"ok": False, "error": f"Unknown mode: {mode}. Use encode/decode/encode_url/decode_url"}
        return {"ok": True, "mode": mode, "result": result, "input_len": len(data), "output_len": len(result)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def extract_archive(path: str, dest: str = "") -> dict:
    """Extract zip or tar archive. dest: output dir, default same as archive."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Archive not found"}
    out = Path(dest) if dest else target.parent
    if not inside_sandbox(out):
        return {"ok": False, "error": "Destination outside sandbox"}
    try:
        import tarfile
        import zipfile
        out = out.resolve()
        out.mkdir(parents=True, exist_ok=True)
        if target.suffix.lower() in (".zip",):
            with zipfile.ZipFile(target, "r") as z:
                for m in z.namelist():
                    if ".." in m or m.startswith("/") or (m.startswith("\\") if len(m) > 1 else False):
                        continue
                    member_path = (out / m).resolve()
                    try:
                        member_path.relative_to(out)
                    except ValueError:
                        continue
                    try:
                        z.extract(m, out)
                    except Exception:
                        pass
        elif target.suffix.lower() in (".tar", ".gz", ".tgz", ".bz2", ".xz"):
            with tarfile.open(target, "r:*") as t:
                for m in t.getnames():
                    if ".." in m or m.startswith("/") or (m.startswith("\\") if len(m) > 1 else False):
                        continue
                    member_path = (out / m).resolve()
                    try:
                        member_path.relative_to(out)
                    except ValueError:
                        continue
                    try:
                        t.extract(m, out)
                    except Exception:
                        pass
        else:
            return {"ok": False, "error": "Unsupported format. Use .zip, .tar, .tar.gz, .tgz"}
        return {"ok": True, "path": str(target), "dest": str(out)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def create_archive(paths: list, output: str, format: str = "zip") -> dict:
    """Create zip archive from paths. paths: list of files/dirs. output: .zip path."""
    out = Path(output)
    if not inside_sandbox(out):
        return {"ok": False, "error": "Output path outside sandbox"}
    try:
        import zipfile
        out.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
            for p in paths:
                fp = Path(p)
                if not inside_sandbox(fp):
                    continue
                if fp.is_file():
                    z.write(fp, fp.name)
                elif fp.is_dir():
                    for f in fp.rglob("*"):
                        if f.is_file():
                            z.write(f, str(f.relative_to(fp)))
        return {"ok": True, "output": str(out), "files": len(paths)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def write_csv(path: str, rows: list, headers: list | None = None) -> dict:
    """Write CSV file. rows: list of dicts or lists. headers: optional column order."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    try:
        import csv
        target.parent.mkdir(parents=True, exist_ok=True)
        with open(target, "w", newline="", encoding="utf-8") as f:
            if rows and isinstance(rows[0], dict):
                h = headers or list(rows[0].keys())
                w = csv.DictWriter(f, fieldnames=h, extrasaction="ignore")
                w.writeheader()
                w.writerows(rows)
            else:
                w = csv.writer(f)
                if headers:
                    w.writerow(headers)
                w.writerows(rows)
        return {"ok": True, "path": str(target), "rows": len(rows)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def read_toml(path: str) -> dict:
    """Parse TOML file. Returns dict."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import tomllib
        return {"ok": True, "path": str(target), "data": tomllib.loads(target.read_text(encoding="utf-8"))}
    except ImportError:
        try:
            import tomli
            return {"ok": True, "path": str(target), "data": tomli.loads(target.read_text(encoding="utf-8"))}
        except ImportError:
            return {"ok": False, "error": "tomllib (3.11+) or tomli required"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def merge_pdf(paths: list, output: str) -> dict:
    """Merge PDFs into one. paths: list of PDF paths. output: output path."""
    out = Path(output)
    if not inside_sandbox(out):
        return {"ok": False, "error": "Output outside sandbox"}
    for p in paths:
        if not inside_sandbox(Path(p)):
            return {"ok": False, "error": f"Path outside sandbox: {p}"}
    try:
        from pypdf import PdfMerger
        merger = PdfMerger()
        for p in paths:
            if Path(p).exists():
                merger.append(str(p))
        merger.write(str(out))
        merger.close()
        return {"ok": True, "output": str(out), "merged": len(paths)}
    except ImportError:
        return {"ok": False, "error": "pypdf not installed: pip install pypdf"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def list_file_checkpoints(path_filter: str = "", limit: int = 50) -> dict:
    """List recent file checkpoints (pre-write snapshots). Optional path_filter limits to one file."""
    from services.file_checkpoints import list_checkpoints

    pf: str | None = None
    raw = (path_filter or "").strip()
    if raw:
        p = Path(raw)
        if not p.is_absolute():
            p = (_get_sandbox() / raw).resolve()
        else:
            p = p.expanduser().resolve()
        pf = str(p)
    return list_checkpoints(
        workspace_root=_get_sandbox(),
        agent_dir=_agent_registry_dir(),
        path_filter=pf,
        limit=limit,
    )

def restore_file_checkpoint(checkpoint_id: str) -> dict:
    """Restore a file from a checkpoint id (see list_file_checkpoints). Overwrites current file."""
    from services.file_checkpoints import restore_checkpoint

    cid = (checkpoint_id or "").strip()
    if not cid:
        return {"ok": False, "error": "checkpoint_id required"}
    return restore_checkpoint(
        checkpoint_id=cid,
        workspace_root=_get_sandbox(),
        agent_dir=_agent_registry_dir(),
        sandbox_root=_get_sandbox(),
    )

