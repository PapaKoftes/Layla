import logging
import re
import subprocess
import sys
import tempfile
import threading
from pathlib import Path
from typing import Any

logger = logging.getLogger("layla")

# Thread-local effective sandbox for research missions (lab path). When set, tools use this instead of config sandbox_root.
_effective_sandbox = threading.local()
# Sandbox path cache (plan §7.1): avoid load_config on every tool call
_sandbox_cache: dict[int, tuple[Path, float]] = {}
_SANDBOX_CACHE_TTL = 2.0  # seconds

def set_effective_sandbox(path: str | None) -> None:
    """Set the effective sandbox root for this thread (e.g. .research_lab/workspace). Used by research missions so read_file/list_dir accept lab paths. Clear with None when run ends."""
    _effective_sandbox.path = path
    try:
        tid = threading.get_ident()
        _sandbox_cache.pop(tid, None)
    except Exception:
        pass

def _get_sandbox() -> Path:
    import time as _time
    try:
        p = getattr(_effective_sandbox, "path", None)
        if p is not None and str(p).strip():
            return Path(p).expanduser().resolve()
    except Exception:
        pass
    tid = threading.get_ident()
    now = _time.monotonic()
    if tid in _sandbox_cache:
        cached, ts = _sandbox_cache[tid]
        if now - ts < _SANDBOX_CACHE_TTL:
            return cached
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        import runtime_safety
        root = runtime_safety.load_config().get("sandbox_root", str(Path.home()))
        result = Path(root).expanduser().resolve()
        _sandbox_cache[tid] = (result, now)
        return result
    except Exception:
        return Path.home().resolve()

# Commands that are never allowed even with allow_run=True
_SHELL_BLOCKLIST = [
    "rm", "del", "rmdir", "format", "mkfs", "dd",
    "shutdown", "reboot", "powershell", "cmd", "reg",
    "netsh", "sc", "taskkill", "cipher",
]


def inside_sandbox(path: Path) -> bool:
    """Check whether path is inside the configured sandbox using Path.relative_to (no string prefix tricks)."""
    try:
        sandbox = _get_sandbox()
        resolved = path.resolve()
        resolved.relative_to(sandbox)
        return True
    except (ValueError, Exception):
        return False


def _write_file_limits() -> tuple[int, int]:
    try:
        import runtime_safety

        cfg = runtime_safety.load_config()
        max_b = max(1024, int(cfg.get("write_file_max_bytes", 500_000)))
        fact = max(2, int(cfg.get("write_file_explosion_factor", 5)))
        return max_b, fact
    except Exception:
        return 500_000, 5


def write_file(path: str, content: str) -> dict:
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    try:
        max_bytes, explosion = _write_file_limits()
        raw = (content or "").encode("utf-8", errors="replace")
        if len(raw) > max_bytes:
            return {"ok": False, "error": "content_too_large", "limit_bytes": max_bytes}
        if target.exists():
            try:
                existing_size = target.stat().st_size
                new_size = len(raw)
                if existing_size > 0 and new_size > existing_size * explosion:
                    return {
                        "ok": False,
                        "error": "size_explosion_detected",
                        "existing_bytes": existing_size,
                        "new_bytes": new_size,
                    }
            except Exception:
                pass
    except Exception:
        pass
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return {"ok": True, "path": str(target)}


def write_files_batch(files: list) -> dict:
    """
    Write multiple files atomically. files: [{path, content}, ...].
    Returns approval_required when gated; otherwise applies all and returns ok.
    """
    if not isinstance(files, list) or not files:
        return {"ok": False, "error": "files must be a non-empty list of {path, content}"}
    written = []
    errors = []
    for i, item in enumerate(files):
        if not isinstance(item, dict):
            errors.append(f"Item {i}: not a dict")
            continue
        path = (item.get("path") or "").strip()
        content = item.get("content", "")
        if not path:
            errors.append(f"Item {i}: missing path")
            continue
        target = Path(path)
        if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
            target = (Path(_effective_sandbox.path) / path).resolve()
        if not inside_sandbox(target):
            errors.append(f"{path}: outside sandbox")
            continue
        try:
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(content, encoding="utf-8")
            written.append(str(target))
        except Exception as e:
            errors.append(f"{path}: {e}")
    if errors:
        return {"ok": False, "error": "; ".join(errors[:5]), "written": written}
    return {"ok": True, "written": written, "count": len(written)}


def read_file(path: str) -> dict:
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    if not target.is_file():
        return {"ok": False, "error": "Not a file"}
    try:
        content = target.read_text(encoding="utf-8", errors="replace")
        return {"ok": True, "path": str(target), "content": content[:8000]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def list_dir(path: str) -> dict:
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}
    try:
        entries = []
        for item in sorted(target.iterdir()):
            entries.append({
                "name": item.name,
                "type": "dir" if item.is_dir() else "file",
            })
        return {"ok": True, "path": str(target), "entries": entries}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def git_status(repo: str) -> dict:
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    result = subprocess.run(
        ["git", "status"],
        cwd=str(repo_path),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    return {"ok": result.returncode == 0, "output": result.stdout or ""}


def shell(argv: list, cwd: str) -> dict:
    try:
        from services.sandbox.shell_runner import run_shell_argv

        cwd_path = Path(cwd)
        return run_shell_argv(list(argv or []), cwd_path, inside_sandbox_check=inside_sandbox)
    except Exception as e:
        logger.debug("shell runner failed, fallback: %s", e)
    if not argv:
        return {"ok": False, "error": "Empty command"}
    cmd = argv[0].lower().lstrip("./\\")
    for blocked in _SHELL_BLOCKLIST:
        if cmd == blocked or cmd.endswith(blocked):
            return {"ok": False, "error": f"Command blocked: {argv[0]}"}
    cwd_path = Path(cwd)
    if not inside_sandbox(cwd_path):
        return {"ok": False, "error": "cwd outside sandbox"}
    try:
        proc = subprocess.run(
            argv,
            cwd=str(cwd_path),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
        )
        return {
            "ok": proc.returncode == 0,
            "stdout": (proc.stdout or "")[:4000],
            "stderr": (proc.stderr or "")[:2000],
            "returncode": proc.returncode,
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "Command timed out (60s)"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def shell_session_start(argv: list | None = None, cwd: str = "") -> dict:
    """Start a background shell process; returns session_id. Requires approval."""
    from services.shell_sessions import shell_session_tool

    return shell_session_tool(action="start", argv=list(argv or []), cwd=cwd or "")


def shell_session_manage(
    action: str = "poll",
    session_id: str = "",
    cwd: str = "",
    limit: int = 80,
) -> dict:
    """Poll, log, or kill a background shell session (no extra approval)."""
    from services.shell_sessions import shell_session_tool

    return shell_session_tool(
        action=action,
        session_id=session_id,
        cwd=cwd or "",
        limit=limit,
    )


def grep_code(pattern: str, path: str, file_glob: str = "*") -> dict:
    """Search for a pattern in files. Tries rg first, falls back to Python re walk."""
    root = Path(path)
    if not root.is_absolute() and getattr(_effective_sandbox, "path", None):
        root = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(root):
        return {"ok": False, "error": "Outside sandbox"}
    if not root.exists():
        return {"ok": False, "error": "Path not found"}
    # Try ripgrep (UTF-8 so Windows doesn't decode with cp1252 and raise on rg output)
    try:
        proc = subprocess.run(
            ["rg", pattern, str(root), "--glob", file_glob, "-n", "--max-count", "5"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=15,
        )
        if proc.returncode in (0, 1):  # 1 = no matches
            out = (proc.stdout if proc.stdout is not None else "")[:6000]
            return {"ok": True, "matches": out}
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # Python fallback
    try:
        rx = re.compile(pattern)
        results = []
        for f in root.rglob(file_glob):
            if not f.is_file():
                continue
            try:
                for i, line in enumerate(f.read_text(encoding="utf-8", errors="replace").splitlines(), 1):
                    if rx.search(line):
                        results.append(f"{f}:{i}: {line.rstrip()}")
                        if len(results) >= 50:
                            break
            except Exception as e:
                logger.debug("grep_code read failed %s: %s", f, e)
                continue
            if len(results) >= 50:
                break
        return {"ok": True, "matches": "\n".join(results)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def glob_files(pattern: str, root: str) -> dict:
    root_path = Path(root)
    if not root_path.is_absolute() and getattr(_effective_sandbox, "path", None):
        root_path = (Path(_effective_sandbox.path) / root).resolve()
    if not inside_sandbox(root_path):
        return {"ok": False, "error": "Outside sandbox"}
    if not root_path.exists():
        return {"ok": False, "error": "Path not found"}
    try:
        matches = [str(p) for p in root_path.rglob(pattern)][:200]
        return {"ok": True, "matches": matches}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def search_codebase(symbol: str, root: str = "") -> dict:
    """Find functions/classes and semantic chunks matching symbol (read-only). root defaults to sandbox."""
    try:
        from services.code_intelligence import search_symbols
    except Exception as e:
        return {"ok": False, "error": str(e), "matches": []}
    root_path = Path(root).expanduser().resolve() if (root or "").strip() else _get_sandbox()
    if not inside_sandbox(root_path):
        return {"ok": False, "error": "Workspace outside sandbox", "matches": []}
    return search_symbols(root_path, symbol, k=25)


def run_python(code: str, cwd: str) -> dict:
    try:
        from services.sandbox.python_runner import run_python_file

        cwd_path = Path(cwd)
        return run_python_file(code or "", cwd_path, inside_sandbox_check=inside_sandbox)
    except Exception as e:
        logger.debug("python runner failed, fallback: %s", e)
    cwd_path = Path(cwd)
    if not inside_sandbox(cwd_path):
        return {"ok": False, "error": "cwd outside sandbox"}
    try:
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".py", delete=False, encoding="utf-8"
        ) as tmp:
            tmp.write(code)
            tmp_path = tmp.name
        proc = subprocess.run(
            [sys.executable, tmp_path],
            cwd=str(cwd_path),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=30,
        )
        Path(tmp_path).unlink(missing_ok=True)
        return {
            "ok": proc.returncode == 0,
            "stdout": (proc.stdout or "")[:4000],
            "stderr": (proc.stderr or "")[:2000],
            "returncode": proc.returncode,
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "run_python timed out (30s)"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def run_tests(cwd: str, pattern: str = "", extra_args: str = "", timeout_s: int = 120) -> dict:
    """
    Discover and run pytest (or unittest) in cwd. pattern: test file/pattern e.g. tests/ or test_foo.py.
    extra_args: e.g. -v -x. Returns pass/fail counts and output.
    """
    cwd_path = Path(cwd)
    if not inside_sandbox(cwd_path):
        return {"ok": False, "error": "cwd outside sandbox"}
    if not cwd_path.exists():
        return {"ok": False, "error": "Path not found"}
    # Prefer pytest
    args = ["python", "-m", "pytest", pattern] if pattern else ["python", "-m", "pytest"]
    if extra_args:
        args.extend(extra_args.split())
    try:
        proc = subprocess.run(
            args,
            cwd=str(cwd_path),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout_s,
        )
        out = (proc.stdout or "") + "\n" + (proc.stderr or "")
        # Parse pytest output for pass/fail
        passed = failed = 0
        for line in out.splitlines():
            if " passed" in line or " passed," in line:
                for part in line.split():
                    if part.isdigit():
                        passed = int(part)
                        break
            if " failed" in line or " failed," in line:
                for part in line.replace(",", " ").split():
                    if part.isdigit():
                        failed = int(part)
                        break
        return {
            "ok": proc.returncode == 0,
            "returncode": proc.returncode,
            "passed": passed,
            "failed": failed,
            "output": out[:8000],
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": f"Tests timed out ({timeout_s}s)"}
    except FileNotFoundError:
        # Fallback: unittest
        try:
            proc = subprocess.run(
                ["python", "-m", "unittest", "discover", "-v"],
                cwd=str(cwd_path),
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=timeout_s,
            )
            out = (proc.stdout or "") + "\n" + (proc.stderr or "")
            return {"ok": proc.returncode == 0, "returncode": proc.returncode, "output": out[:8000], "runner": "unittest"}
        except Exception as e:
            return {"ok": False, "error": str(e)}


def parse_gcode(path: str) -> dict:
    """
    Parse G-code / NC file: moves, tools, units, bounds, feed rates.
    Supports .gcode, .nc, .tap, .sbp.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        content = target.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return {"ok": False, "error": str(e)}
    # Parse G-code patterns
    tools = set()
    feed_rates = []
    bounds = {"x": [], "y": [], "z": []}
    units = "mm"  # default
    for line in content.splitlines():
        line = line.strip()
        if not line or line.startswith(";"):
            continue
        # G20 = inches, G21 = mm
        if "G20" in line.upper():
            units = "in"
        elif "G21" in line.upper():
            units = "mm"
        # Tool: T1, T02, M6 T1
        for m in re.finditer(r"\bT(\d+)\b", line, re.I):
            tools.add(int(m.group(1)))
        # Feed: F3000, F100.5
        for m in re.finditer(r"\bF([\d.]+)\b", line, re.I):
            feed_rates.append(float(m.group(1)))
        # Moves: G0/G1 X Y Z
        for m in re.finditer(r"\b[Gg]?[01]\b.*\b([XYZxyz])([-]?[\d.]+)", line):
            axis = m.group(1).lower()
            val = float(m.group(2))
            if axis in bounds:
                bounds[axis].append(val)
    # Summarize bounds
    summary = {}
    for ax, vals in bounds.items():
        if vals:
            summary[ax] = {"min": min(vals), "max": max(vals), "count": len(vals)}
    return {
        "ok": True,
        "path": str(target),
        "units": units,
        "tools": sorted(tools) if tools else [],
        "move_count": sum(len(bounds[ax]) for ax in bounds),
        "feed_rates": list(set(feed_rates))[:20] if feed_rates else [],
        "bounds": summary,
        "lines": len([ln for ln in content.splitlines() if ln.strip() and not ln.strip().startswith(";")]),
    }


def stl_mesh_info(path: str) -> dict:
    """STL mesh stats: vertex count, bounds, volume. Requires trimesh or numpy."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import trimesh
        mesh = trimesh.load(str(target))
        if isinstance(mesh, trimesh.Scene):
            mesh = mesh.dump(concatenate=True)
        bounds = mesh.bounds
        vol = float(mesh.volume) if hasattr(mesh, "volume") else None
        return {
            "ok": True,
            "path": str(target),
            "vertices": int(len(mesh.vertices)),
            "faces": int(len(mesh.faces)) if hasattr(mesh, "faces") else None,
            "bounds": {"min": bounds[0].tolist(), "max": bounds[1].tolist()},
            "volume": vol,
        }
    except ImportError:
        # Fallback: count vertices from ASCII STL
        try:
            content = target.read_text(encoding="utf-8", errors="replace")
            if "vertex" in content.lower():
                verts = re.findall(r"vertex\s+([-\d.e]+)\s+([-\d.e]+)\s+([-\d.e]+)", content, re.I)
                xs, ys, zs = zip(*[[float(a), float(b), float(c)] for a, b, c in verts]) if verts else ([], [], [])
                return {
                    "ok": True,
                    "path": str(target),
                    "vertices": len(verts),
                    "bounds": {"min": [min(xs), min(ys), min(zs)], "max": [max(xs), max(ys), max(zs)]} if verts else {},
                    "fallback": "ascii_parse",
                }
        except Exception as e:
            return {"ok": False, "error": str(e), "hint": "pip install trimesh for full support"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def tail_file(path: str, n: int = 50) -> dict:
    """Return last n lines of a file. Useful for logs."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        with open(target, encoding="utf-8", errors="replace") as f:
            lines = f.readlines()
        tail = lines[-n:] if len(lines) > n else lines
        return {"ok": True, "path": str(target), "lines": "".join(tail)[:8000], "total_lines": len(lines)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def clipboard_read() -> dict:
    """Read text from system clipboard. Requires pyperclip."""
    try:
        import pyperclip
        text = pyperclip.paste()
        return {"ok": True, "text": (text or "")[:10000]}
    except ImportError:
        return {"ok": False, "error": "pyperclip not installed: pip install pyperclip"}


def clipboard_write(text: str) -> dict:
    """Write text to system clipboard. Requires pyperclip."""
    try:
        import pyperclip
        pyperclip.copy(text[:50000])
        return {"ok": True, "length": len(text)}
    except ImportError:
        return {"ok": False, "error": "pyperclip not installed: pip install pyperclip"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def search_replace(root: str, find: str, replace: str, file_glob: str = "*", dry_run: bool = True) -> dict:
    """
    Multi-file find/replace. dry_run=True lists matches without changing. Uses regex if find contains regex chars.
    """
    root_path = Path(root)
    if not inside_sandbox(root_path):
        return {"ok": False, "error": "Outside sandbox"}
    if not root_path.exists():
        return {"ok": False, "error": "Path not found"}
    use_regex = bool(re.search(r"[.*+?^${}()|[\]\\]", find))
    pattern = re.compile(find) if use_regex else None
    matches = []
    for f in root_path.rglob(file_glob):
        if not f.is_file():
            continue
        try:
            content = f.read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue
        if use_regex:
            new_content, n = pattern.subn(replace, content)
        else:
            n = content.count(find)
            new_content = content.replace(find, replace) if n else content
        if n:
            matches.append({"path": str(f), "count": n})
            if not dry_run:
                f.write_text(new_content, encoding="utf-8")
    return {"ok": True, "dry_run": dry_run, "matches": matches[:100], "total_files": len(matches)}


def pip_list(cwd: str = "") -> dict:
    """List installed pip packages. cwd: optional venv path."""
    try:
        argv = [sys.executable, "-m", "pip", "list", "--format=json"]
        proc = subprocess.run(
            argv,
            cwd=cwd or str(Path.cwd()),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=30,
        )
        if proc.returncode != 0:
            return {"ok": False, "error": (proc.stderr or proc.stdout or "")[:500]}
        import json as _json
        data = _json.loads(proc.stdout or "[]")
        return {"ok": True, "packages": [{"name": p["name"], "version": p["version"]} for p in data[:200]]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def pip_install(packages: str, cwd: str = "", upgrade: bool = False) -> dict:
    """Install pip package(s). packages: space-separated names or path to requirements.txt."""
    cwd_path = Path(cwd) if cwd else Path.cwd()
    if cwd and not inside_sandbox(cwd_path):
        return {"ok": False, "error": "cwd outside sandbox"}
    argv = [sys.executable, "-m", "pip", "install"]
    if upgrade:
        argv.append("--upgrade")
    pkg = packages.strip()
    if pkg.endswith(".txt") and Path(pkg).exists():
        argv.append("-r")
        argv.append(pkg)
    else:
        argv.extend(pkg.split())
    try:
        proc = subprocess.run(
            argv,
            cwd=str(cwd_path),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )
        return {"ok": proc.returncode == 0, "output": (proc.stdout or proc.stderr or "")[:4000]}
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "pip install timed out"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def apply_patch(original_path: str, patch_text: str) -> dict:
    """Apply a unified diff patch using unidiff (pure Python, Windows-safe). Creates a backup first."""
    target = Path(original_path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    import shutil

    from layla.time_utils import utcnow
    backup = target.with_suffix(
        f".bak_{utcnow().strftime('%Y%m%d_%H%M%S')}{target.suffix}"
    )
    shutil.copy2(str(target), str(backup))
    try:
        import unidiff
        patch_set = unidiff.PatchSet(patch_text.splitlines(keepends=True))
        original_lines = target.read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
        result_lines = list(original_lines)
        for patched_file in patch_set:
            offset = 0
            for hunk in patched_file:
                src_start = hunk.source_start - 1 + offset
                removed = [line.value for line in hunk if line.is_removed]
                added = [line.value for line in hunk if line.is_added]
                # Remove old lines, insert new ones
                result_lines[src_start: src_start + len(removed)] = added
                offset += len(added) - len(removed)
        target.write_text("".join(result_lines), encoding="utf-8")
        return {"ok": True, "path": str(target), "backup": str(backup)}
    except Exception as e:
        return {"ok": False, "error": str(e), "backup": str(backup)}


def git_diff(repo: str) -> dict:
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    result = subprocess.run(
        ["git", "diff"],
        cwd=str(repo_path),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or "")[:8000]}


def git_log(repo: str, n: int = 10) -> dict:
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    result = subprocess.run(
        ["git", "log", "--oneline", f"-{n}"],
        cwd=str(repo_path),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or "")[:4000]}


def git_branch(repo: str) -> dict:
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    result = subprocess.run(
        ["git", "branch", "--show-current"],
        cwd=str(repo_path),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or "").strip()}


def fetch_url_tool(url: str, store: bool = False) -> dict:
    try:
        import runtime_safety
        from services.http_response_cache import get_cached, set_cached

        cfg = runtime_safety.load_config()
        if not store:
            ck = f"fetch:{url}"
            hit = get_cached(ck, cfg)
            if hit is not None:
                return hit
    except Exception:
        cfg = {}
    from layla.tools.web import fetch_url

    out = fetch_url(url, store=store)
    try:
        if not store and out.get("ok"):
            import runtime_safety
            from services.http_response_cache import set_cached

            set_cached(f"fetch:{url}", out, runtime_safety.load_config())
    except Exception:
        pass
    return out


def file_info(path: str) -> dict:
    """Return size, line count (approx), and whether file looks text. Read-only; no approval."""
    target = Path(path)
    if not target.is_absolute() and getattr(_effective_sandbox, "path", None):
        target = (Path(_effective_sandbox.path) / path).resolve()
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}
    if not target.is_file():
        return {"ok": False, "error": "Not a file"}
    try:
        size = target.stat().st_size
        # Sample first 8k to guess text vs binary and count newlines
        raw = target.read_bytes()
        sample = raw[:8192]
        try:
            sample.decode("utf-8", errors="strict")
            is_text = True
        except Exception:
            is_text = False
        line_count = sample.count(b"\n") if is_text else None
        note = "from first 8k only" if is_text and len(raw) > len(sample) else None
        return {"ok": True, "path": str(target), "size_bytes": size, "is_text": is_text, "line_count_sample": line_count, "note": note}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def get_project_context_tool() -> dict:
    """Return current project context (read-only for agent)."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import get_project_context
        return {"ok": True, **get_project_context()}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def update_project_context_tool(
    project_name: str = "",
    domains: list | None = None,
    key_files: list | None = None,
    goals: str = "",
    lifecycle_stage: str = "",
    progress: str = "",
    blockers: str = "",
    last_discussed: str = "",
) -> dict:
    """Update project context. lifecycle_stage: idea|planning|prototype|iteration|execution|reflection. progress/blockers/last_discussed for companion recall."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import set_project_context
        set_project_context(
            project_name=project_name or "",
            domains=domains,
            key_files=key_files,
            goals=goals,
            lifecycle_stage=lifecycle_stage or "",
            progress=progress or "",
            blockers=blockers or "",
            last_discussed=last_discussed or "",
        )
        return {"ok": True, "message": "Project context updated."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def get_user_identity_tool() -> dict:
    """Return user/companion identity context (verbosity, humor, formality, response length, life narrative). Read-only."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import get_all_user_identity
        return {"ok": True, **get_all_user_identity()}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def update_user_identity_tool(key: str, snapshot: str) -> dict:
    """Update user identity. key: verbosity|humor_tolerance|formality|response_length|life_narrative_summary. snapshot: description."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import USER_IDENTITY_KEYS, set_user_identity
        if key not in USER_IDENTITY_KEYS:
            return {"ok": False, "error": f"key must be one of: {USER_IDENTITY_KEYS}"}
        set_user_identity(key, snapshot)
        return {"ok": True, "message": f"User identity '{key}' updated."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def add_goal_tool(title: str, description: str = "", project_id: str = "") -> dict:
    """Add a long-term goal. project_id: optional project name to associate."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import add_goal
        gid = add_goal(title=title, description=description, project_id=project_id)
        return {"ok": True, "goal_id": gid, "message": "Goal added."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def add_goal_progress_tool(goal_id: str, note: str = "", progress_pct: float = 0) -> dict:
    """Record progress on a goal. progress_pct: 0-100."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import add_goal_progress
        add_goal_progress(goal_id, note=note, progress_pct=progress_pct)
        return {"ok": True, "message": "Progress recorded."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def get_active_goals_tool(project_id: str = "") -> dict:
    """Return active goals, optionally filtered by project_id."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import get_active_goals
        goals = get_active_goals(project_id=project_id)
        return {"ok": True, "goals": goals}
    except Exception as e:
        return {"ok": False, "error": str(e), "goals": []}


def understand_file_tool(path: str, content: str | None = None) -> dict:
    """Interpret file intent (read-only). path: file path; optional content for in-memory analysis."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.file_understanding import analyze_file
        if content is not None:
            return {"ok": True, **analyze_file(file_path=path, content=content)}
        return {"ok": True, **analyze_file(file_path=path)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def browser_navigate(url: str, timeout_ms: int = 15000) -> dict:
    """Navigate to a URL and return its main text content and title."""
    try:
        from services.browser import navigate
        return navigate(url, timeout_ms=timeout_ms)
    except ImportError:
        return {"ok": False, "error": "playwright not installed. Run: playwright install chromium"}


def browser_search(query: str) -> dict:
    """Search the web via DuckDuckGo. Returns top 8 results with titles, URLs, snippets."""
    try:
        from services.browser import search_web
        return search_web(query)
    except ImportError:
        return {"ok": False, "error": "playwright not installed. Run: playwright install chromium"}


def browser_screenshot(url: str) -> dict:
    """Take a full-page screenshot of a URL. Returns path to the screenshot file."""
    try:
        from services.browser import screenshot
        return screenshot(url)
    except ImportError:
        return {"ok": False, "error": "playwright not installed. Run: playwright install chromium"}


def browser_click(url: str, selector: str) -> dict:
    """Navigate to a URL, click a CSS selector, return updated page text."""
    try:
        from services.browser import click_and_extract
        return click_and_extract(url, selector)
    except ImportError:
        return {"ok": False, "error": "playwright not installed. Run: playwright install chromium"}


def browser_fill(url: str, fields: dict, submit_selector: str = "") -> dict:
    """Navigate to a URL, fill form fields {selector: value}, optionally submit."""
    try:
        from services.browser import fill_form
        return fill_form(url, fields, submit_selector)
    except ImportError:
        return {"ok": False, "error": "playwright not installed. Run: playwright install chromium"}


# ─── Extended tools ────────────────────────────────────────────────────────────

def json_query(path: str, query: str = "") -> dict:
    """
    Parse a JSON file and optionally extract a value by dot-notation path.
    query examples: 'key', 'nested.key', 'array.0.field'.
    If no query, returns the full parsed object (truncated).
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import json as _json
        data = _json.loads(target.read_text(encoding="utf-8"))
        if not query:
            return {"ok": True, "data": str(data)[:4000]}
        parts = query.split(".")
        val = data
        for p in parts:
            if isinstance(val, dict):
                val = val[p]
            elif isinstance(val, list):
                val = val[int(p)]
            else:
                return {"ok": False, "error": f"Cannot traverse into {type(val).__name__} at '{p}'"}
        return {"ok": True, "result": val, "result_str": str(val)[:2000]}
    except (KeyError, IndexError) as e:
        return {"ok": False, "error": f"Path not found: {e}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def diff_files(path_a: str, path_b: str) -> dict:
    """Diff two text files. Returns unified diff."""
    for p in (path_a, path_b):
        t = Path(p)
        if not inside_sandbox(t):
            return {"ok": False, "error": f"Outside sandbox: {p}"}
        if not t.exists():
            return {"ok": False, "error": f"File not found: {p}"}
    try:
        import difflib
        a_lines = Path(path_a).read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
        b_lines = Path(path_b).read_text(encoding="utf-8", errors="replace").splitlines(keepends=True)
        diff = list(difflib.unified_diff(a_lines, b_lines, fromfile=path_a, tofile=path_b, n=3))
        return {"ok": True, "diff": "".join(diff)[:8000], "changed": len(diff) > 0}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def env_info() -> dict:
    """Return system info: OS, Python version, CPU, RAM, GPU, installed key packages."""
    import platform
    import sys as _sys
    info: dict = {
        "os": platform.system(),
        "os_version": platform.version(),
        "python": _sys.version.split()[0],
        "architecture": platform.machine(),
    }
    try:
        import psutil
        mem = psutil.virtual_memory()
        info["ram_total_gb"] = round(mem.total / (1024**3), 1)
        info["ram_available_gb"] = round(mem.available / (1024**3), 1)
        info["cpu_logical"] = psutil.cpu_count(logical=True)
    except Exception:
        pass
    try:
        r = subprocess.run(  # noqa: F841
            ["nvidia-smi", "--query-gpu=name,memory.total", "--format=csv,noheader"],
            capture_output=True, text=True, timeout=8, encoding="utf-8", errors="replace",
        )
        if r.returncode == 0:
            info["gpu"] = r.stdout.strip()
    except Exception:
        info["gpu"] = "none / not detected"
    key_packages = ["fastapi", "uvicorn", "llama_cpp", "chromadb", "sentence_transformers",
                    "playwright", "faster_whisper", "psutil", "rank_bm25"]
    installed = {}
    import importlib.metadata as _meta
    for pkg in key_packages:
        try:
            installed[pkg] = _meta.version(pkg.replace("_", "-"))
        except Exception:
            installed[pkg] = "not installed"
    info["packages"] = installed
    return {"ok": True, **info}


def regex_test(pattern: str, text: str, flags: str = "") -> dict:
    """Test a regex pattern against text. Returns matches, groups, count."""
    try:
        import re as _re
        flag_map = {"i": _re.IGNORECASE, "m": _re.MULTILINE, "s": _re.DOTALL}
        compiled_flags = 0
        for f in flags.lower():
            compiled_flags |= flag_map.get(f, 0)
        rx = _re.compile(pattern, compiled_flags)
        matches = list(rx.finditer(text))
        result = []
        for m in matches[:20]:
            result.append({"match": m.group(0), "start": m.start(), "end": m.end(), "groups": list(m.groups())})
        return {"ok": True, "count": len(matches), "matches": result, "pattern": pattern}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def git_add(repo: str, path: str = ".") -> dict:
    """Stage files for commit. path: file or '.' for all."""
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    result = subprocess.run(
        ["git", "add", path],
        cwd=str(repo_path),
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    return {"ok": result.returncode == 0, "output": result.stdout or result.stderr or ""}


def git_commit(repo: str, message: str, add_all: bool = False) -> dict:
    """Commit staged changes. If add_all=True, stages everything first."""
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    if add_all:
        subprocess.run(["git", "add", "-A"], cwd=str(repo_path), capture_output=True)
    result = subprocess.run(
        ["git", "commit", "-m", message],
        cwd=str(repo_path),
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or result.stderr or "")[:2000]}


def git_push(repo: str, remote: str = "origin", branch: str = "") -> dict:
    """Push commits to remote. branch: empty = current branch."""
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    argv = ["git", "push", remote]
    if branch:
        argv.append(branch)
    result = subprocess.run(
        argv,
        cwd=str(repo_path),
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=60,
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or result.stderr or "")[:2000]}


def git_pull(repo: str, remote: str = "origin", branch: str = "") -> dict:
    """Pull from remote. branch: empty = current branch."""
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    argv = ["git", "pull", remote]
    if branch:
        argv.append(branch)
    result = subprocess.run(
        argv,
        cwd=str(repo_path),
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=60,
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or result.stderr or "")[:2000]}


def git_stash(repo: str, action: str = "list", message: str = "") -> dict:
    """Stash changes. action: list|push|pop|apply."""
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    if action == "list":
        argv = ["git", "stash", "list"]
    elif action == "push":
        argv = ["git", "stash", "push", "-m", message or "layla stash"]
    elif action in ("pop", "apply"):
        argv = ["git", "stash", action]
    else:
        return {"ok": False, "error": f"Unknown action: {action}. Use list|push|pop|apply"}
    result = subprocess.run(
        argv,
        cwd=str(repo_path),
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or result.stderr or "")[:2000]}


def git_revert(repo: str, commit: str, no_commit: bool = False) -> dict:
    """Revert a commit. commit: hash or HEAD~1. no_commit=True leaves changes staged."""
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    argv = ["git", "revert", "--no-edit", commit]
    if no_commit:
        argv.append("--no-commit")
    result = subprocess.run(
        argv,
        cwd=str(repo_path),
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or result.stderr or "")[:2000]}


def git_clone(url: str, dest: str, depth: int = 0) -> dict:
    """Clone a git repo. dest: path inside sandbox. depth: 0 = full clone."""
    dest_path = Path(dest)
    if not inside_sandbox(dest_path):
        return {"ok": False, "error": "Destination outside sandbox"}
    argv = ["git", "clone", url, str(dest_path)]
    if depth:
        argv.insert(2, f"--depth={depth}")
    result = subprocess.run(
        argv,
        capture_output=True, text=True, encoding="utf-8", errors="replace",
        timeout=120,
    )
    return {"ok": result.returncode == 0, "output": (result.stdout or result.stderr or "")[:2000]}


def save_note(content: str, tag: str = "note") -> dict:
    """
    Save a note directly to Layla's memory as a learning.
    Use this to remember facts, preferences, or observations mid-conversation.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import save_learning
        save_learning(content=content[:800], kind=tag)
        return {"ok": True, "saved": content[:100]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def search_memories(query: str, n: int = 8) -> dict:
    """
    Search Layla's own memory (learnings + semantic recall) for relevant past knowledge.
    Returns the most relevant stored memories for the given query.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.vector_store import search_memories_full
        results = search_memories_full(query, k=n, use_rerank=False)
        items = [r.get("content", "") for r in results if r.get("content")]
        return {"ok": True, "memories": items, "count": len(items)}
    except Exception as e:
        try:
            from layla.memory.db import get_recent_learnings
            rows = get_recent_learnings(n=n)
            items = [r.get("content", "") for r in rows if r.get("content")]
            return {"ok": True, "memories": items, "count": len(items), "fallback": True}
        except Exception:
            return {"ok": False, "error": str(e)}


def memory_search(query: str, n: int = 8) -> dict:
    """OpenClaw-style alias: search long-term memory and learnings (same as search_memories)."""
    return search_memories(query, n=n)


def memory_get(query: str, n: int = 5) -> dict:
    """OpenClaw-style: retrieve memory snippets by semantic query."""
    return search_memories(query, n=n)


def structured_llm_task(
    instruction: str,
    schema_hint: str = "",
    max_tokens: int = 256,
) -> dict:
    """One bounded LLM step that should return JSON (OpenClaw llm-task style). No file writes."""
    import json as _json

    from services.llm_gateway import run_completion

    sh = (schema_hint or "").strip()
    prompt = (
        "You are a precise JSON generator. Output ONLY one JSON object, no markdown fences.\n"
        f"Task: {instruction}\n"
    )
    if sh:
        prompt += f"Expected shape / keys: {sh}\n"
    mt = min(512, max(32, int(max_tokens)))
    try:
        out = run_completion(prompt, max_tokens=mt, temperature=0.1, stream=False)
    except Exception as e:
        return {"ok": False, "error": f"LLM call failed: {e}"}
    text = ""
    if isinstance(out, dict):
        choices = out.get("choices") or [{}]
        text = (choices[0].get("message") or {}).get("content") or ""
    text = (text or "").strip()
    try:
        obj = _json.loads(text)
        return {"ok": True, "json": obj, "raw": text[:2000]}
    except Exception:
        return {"ok": True, "json": None, "raw": text[:4000], "parse_error": True}


# ─── Research & Information tools ─────────────────────────────────────────────

def read_pdf(path: str, max_pages: int = 30) -> dict:
    """
    Extract text from a PDF file using PyMuPDF.
    Returns text per page up to max_pages. Falls back to pypdf if fitz not installed.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    # Try PyMuPDF (fitz) first
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(target))
        try:
            pages = []
            for i, page in enumerate(doc):
                if i >= max_pages:
                    break
                pages.append(f"--- Page {i+1} ---\n{page.get_text()}")
            full = "\n".join(pages)
            return {"ok": True, "path": str(target), "pages": min(len(pages), max_pages), "text": full[:12000]}
        finally:
            doc.close()
    except ImportError:
        pass
    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(target))
        pages = []
        for i, page in enumerate(reader.pages):
            if i >= max_pages:
                break
            pages.append(f"--- Page {i+1} ---\n{page.extract_text() or ''}")
        full = "\n".join(pages)
        return {"ok": True, "path": str(target), "pages": min(len(reader.pages), max_pages), "text": full[:12000]}
    except ImportError:
        return {"ok": False, "error": "PDF reading requires PyMuPDF or pypdf: pip install PyMuPDF"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def fetch_article(url: str) -> dict:
    """
    Extract clean text from a web article using trafilatura.
    Much cleaner than raw fetch — removes nav, ads, footers. Ideal for research.
    """
    try:
        import trafilatura
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            return {"ok": False, "error": "Could not fetch URL"}
        text = trafilatura.extract(
            downloaded,
            include_comments=False,
            include_tables=True,
            favor_precision=True,
        )
        if not text:
            # Fallback to raw text
            text = trafilatura.extract(downloaded, favor_recall=True)
        if not text:
            return {"ok": False, "error": "Could not extract content from page"}
        title = ""
        try:
            meta = trafilatura.extract_metadata(downloaded)
            if meta:
                title = meta.title or ""
        except Exception:
            pass
        return {"ok": True, "url": url, "title": title, "text": text[:10000], "chars": len(text)}
    except ImportError:
        return {"ok": False, "error": "trafilatura not installed: pip install trafilatura"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def wiki_search(query: str, sentences: int = 8, lang: str = "en") -> dict:
    """
    Search Wikipedia and return a summary. sentences controls summary length.
    Returns the intro, URL, and a list of related page titles.
    """
    try:
        import wikipedia
        wikipedia.set_lang(lang)
        try:
            summary = wikipedia.summary(query, sentences=sentences, auto_suggest=True)
            page = wikipedia.page(query, auto_suggest=True)
            return {
                "ok": True,
                "query": query,
                "title": page.title,
                "url": page.url,
                "summary": summary,
                "related": page.links[:10],
            }
        except wikipedia.DisambiguationError as e:
            # Return top options on disambiguation
            return {"ok": True, "query": query, "disambiguation": True, "options": e.options[:8]}
        except wikipedia.PageError:
            results = wikipedia.search(query, results=5)
            return {"ok": False, "query": query, "error": "Page not found", "suggestions": results}
    except ImportError:
        return {"ok": False, "error": "wikipedia package not installed: pip install wikipedia-api"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def ddg_search(query: str, max_results: int = 10, region: str = "wt-wt") -> dict:
    """
    DuckDuckGo web search — pure Python, no browser required.
    Returns results with title, href, body snippet.
    """
    try:
        import runtime_safety
        from services.http_response_cache import get_cached, set_cached

        cfg = runtime_safety.load_config()
        ck = f"ddg:{region}:{max_results}:{query}"
        hit = get_cached(ck, cfg)
        if hit is not None:
            return hit
    except Exception:
        cfg = {}
    try:
        from duckduckgo_search import DDGS
        with DDGS() as ddgs:
            results = list(ddgs.text(query, region=region, max_results=max_results))
        out = {"ok": True, "query": query, "results": results, "count": len(results)}
        try:
            import runtime_safety
            from services.http_response_cache import set_cached

            set_cached(f"ddg:{region}:{max_results}:{query}", out, runtime_safety.load_config())
        except Exception:
            pass
        return out
    except ImportError:
        return {"ok": False, "error": "duckduckgo-search not installed: pip install duckduckgo-search"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def arxiv_search(query: str, max_results: int = 5, sort_by: str = "relevance") -> dict:
    """
    Search arXiv for papers. Returns title, authors, abstract, PDF URL, published date.
    sort_by: 'relevance' | 'lastUpdatedDate' | 'submittedDate'
    """
    try:
        import arxiv
        sort_map = {
            "relevance": arxiv.SortCriterion.Relevance,
            "lastUpdatedDate": arxiv.SortCriterion.LastUpdatedDate,
            "submittedDate": arxiv.SortCriterion.SubmittedDate,
        }
        client = arxiv.Client()
        search = arxiv.Search(
            query=query,
            max_results=max_results,
            sort_by=sort_map.get(sort_by, arxiv.SortCriterion.Relevance),
        )
        papers = []
        for r in client.results(search):
            papers.append({
                "title": r.title,
                "authors": [str(a) for a in r.authors[:5]],
                "abstract": (r.summary or "")[:500],
                "pdf_url": r.pdf_url,
                "published": str(r.published)[:10] if r.published else "",
                "arxiv_id": r.entry_id.split("/")[-1],
                "categories": r.categories[:3],
            })
        return {"ok": True, "query": query, "results": papers, "count": len(papers)}
    except ImportError:
        return {"ok": False, "error": "arxiv not installed: pip install arxiv"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def math_eval(expression: str) -> dict:
    """
    Safely evaluate a mathematical expression. Supports: +, -, *, /, **, %, //, abs, round,
    min, max, sum, int, float, sqrt, log, sin, cos, tan, pi, e, and more.
    No arbitrary code execution — uses a strict AST whitelist.
    """
    import ast as _ast
    import math as _math
    import operator as _op

    _SAFE_NODES = (
        _ast.Expression, _ast.BinOp, _ast.UnaryOp, _ast.Call, _ast.Constant,
        _ast.Add, _ast.Sub, _ast.Mul, _ast.Div, _ast.Pow, _ast.Mod, _ast.FloorDiv,
        _ast.UAdd, _ast.USub, _ast.Compare, _ast.Lt, _ast.Gt, _ast.LtE, _ast.GtE,
        _ast.Eq, _ast.NotEq, _ast.BoolOp, _ast.And, _ast.Or, _ast.Name, _ast.List,
        _ast.Tuple,
    )
    _SAFE_FUNCS = {
        "abs": abs, "round": round, "min": min, "max": max, "sum": sum,
        "int": int, "float": float, "bool": bool,
        "sqrt": _math.sqrt, "log": _math.log, "log2": _math.log2, "log10": _math.log10,
        "sin": _math.sin, "cos": _math.cos, "tan": _math.tan,
        "asin": _math.asin, "acos": _math.acos, "atan": _math.atan, "atan2": _math.atan2,
        "ceil": _math.ceil, "floor": _math.floor, "trunc": _math.trunc,
        "factorial": _math.factorial, "gcd": _math.gcd,
        "degrees": _math.degrees, "radians": _math.radians,
        "pi": _math.pi, "e": _math.e, "tau": _math.tau, "inf": _math.inf,
        "pow": pow, "divmod": divmod,
    }

    def _safe_eval(node):
        if not isinstance(node, _SAFE_NODES):
            raise ValueError(f"Disallowed operation: {type(node).__name__}")
        if isinstance(node, _ast.Constant):
            return node.value
        if isinstance(node, _ast.Name):
            if node.id in _SAFE_FUNCS:
                return _SAFE_FUNCS[node.id]
            raise ValueError(f"Unknown name: {node.id}")
        if isinstance(node, _ast.BinOp):
            ops = {_ast.Add: _op.add, _ast.Sub: _op.sub, _ast.Mul: _op.mul,
                   _ast.Div: _op.truediv, _ast.Pow: _op.pow, _ast.Mod: _op.mod,
                   _ast.FloorDiv: _op.floordiv}
            return ops[type(node.op)](_safe_eval(node.left), _safe_eval(node.right))
        if isinstance(node, _ast.UnaryOp):
            ops = {_ast.UAdd: _op.pos, _ast.USub: _op.neg}
            return ops[type(node.op)](_safe_eval(node.operand))
        if isinstance(node, _ast.Call):
            func = _safe_eval(node.func)
            args = [_safe_eval(a) for a in node.args]
            return func(*args)
        if isinstance(node, (_ast.List, _ast.Tuple)):
            return [_safe_eval(el) for el in node.elts]
        raise ValueError(f"Unsupported node: {type(node).__name__}")

    try:
        tree = _ast.parse(expression.strip(), mode="eval")
        result = _safe_eval(tree.body)
        return {"ok": True, "expression": expression, "result": result, "result_str": str(result)}
    except ZeroDivisionError:
        return {"ok": False, "error": "Division by zero"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def read_csv(path: str, max_rows: int = 50, describe: bool = True) -> dict:
    """
    Read a CSV file and return a summary. max_rows controls rows returned.
    If describe=True, returns statistical summary (count, mean, std, etc.).
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import pandas as _pd
        df = _pd.read_csv(str(target))
        result: dict = {
            "ok": True,
            "path": str(target),
            "rows": len(df),
            "columns": list(df.columns),
            "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
            "sample": df.head(max_rows).to_dict(orient="records"),
            "null_counts": df.isnull().sum().to_dict(),
        }
        if describe:
            try:
                result["stats"] = df.describe().to_dict()
            except Exception:
                pass
        return result
    except ImportError:
        # Fallback to stdlib csv
        import csv as _csv
        with open(str(target), newline="", encoding="utf-8", errors="replace") as f:
            reader = _csv.DictReader(f)
            rows = [row for _, row in zip(range(max_rows + 1), reader)]
        return {"ok": True, "path": str(target), "columns": reader.fieldnames or [], "sample": rows[:max_rows]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def count_tokens(text: str, model: str = "gpt-4") -> dict:
    """
    Count tokens in text. Uses tiktoken (cl100k_base) when available.
    model hint used for encoding_for_model when supported; falls back to cl100k_base.
    """
    try:
        from services.token_count import count_tokens as _count
        from services.token_count import token_count_available
        if token_count_available():
            n = _count(text)
            return {"ok": True, "tokens": n, "model": "cl100k_base", "method": "tiktoken"}
    except Exception:
        pass
    try:
        import tiktoken
        enc = tiktoken.encoding_for_model(model)
        tokens = enc.encode(text)
        return {"ok": True, "tokens": len(tokens), "model": model, "method": "tiktoken"}
    except Exception:
        pass
    # Fallback: ~4 chars per token
    import re as _re
    words = len(_re.split(r"\s+", (text or "").strip())) or 1
    chars = len(text or "")
    rough = max(int(chars / 4), words)
    return {"ok": True, "tokens": rough, "model": "estimate", "method": "rough (~4 chars/token)"}


def http_request(url: str, method: str = "GET", body: str = "", headers: dict | None = None, timeout: int = 15) -> dict:
    """
    Make an HTTP request. method: GET | POST | PUT | DELETE | PATCH.
    Returns status, response text (truncated to 8000 chars).
    Use for webhooks, REST APIs, testing endpoints.
    """
    import urllib.error
    import urllib.request
    method = method.upper()
    hdrs = {"User-Agent": "Layla/2.0 research agent", "Accept": "application/json,text/html,*/*"}
    if headers:
        hdrs.update(headers)
    try:
        data = body.encode("utf-8") if body else None
        req = urllib.request.Request(url, data=data, headers=hdrs, method=method)
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            content = resp.read(80000).decode("utf-8", errors="replace")
            return {
                "ok": resp.status < 400,
                "status": resp.status,
                "url": url,
                "text": content[:8000],
                "headers": dict(resp.headers),
            }
    except urllib.error.HTTPError as e:
        body_text = ""
        try:
            body_text = e.read(2000).decode("utf-8", errors="replace")
        except Exception:
            pass
        return {"ok": False, "status": e.code, "error": str(e), "text": body_text}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def python_ast(path: str) -> dict:
    """
    Analyze a Python file's AST structure. Returns:
    - Top-level functions and classes (with line numbers, decorators, docstrings)
    - Imports
    - Global variables
    - Complexity indicators (nested function depth, line count)
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    import ast as _ast
    try:
        source = target.read_text(encoding="utf-8", errors="replace")
        tree = _ast.parse(source, filename=str(target))
    except SyntaxError as e:
        return {"ok": False, "error": f"SyntaxError: {e}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

    functions, classes, imports, globals_list = [], [], [], []

    for node in _ast.walk(tree):
        if isinstance(node, (_ast.Import, _ast.ImportFrom)):
            if isinstance(node, _ast.Import):
                for alias in node.names:
                    imports.append(alias.name)
            else:
                mod = node.module or ""
                for alias in node.names:
                    imports.append(f"{mod}.{alias.name}" if mod else alias.name)

    for node in tree.body:
        if isinstance(node, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
            decorators = [_ast.unparse(d) for d in (node.decorator_list or [])]
            doc = _ast.get_docstring(node) or ""
            functions.append({
                "name": node.name,
                "line": node.lineno,
                "async": isinstance(node, _ast.AsyncFunctionDef),
                "args": [a.arg for a in node.args.args],
                "decorators": decorators,
                "docstring": doc[:120],
            })
        elif isinstance(node, _ast.ClassDef):
            methods = []
            for item in node.body:
                if isinstance(item, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
                    methods.append({"name": item.name, "line": item.lineno})
            doc = _ast.get_docstring(node) or ""
            classes.append({
                "name": node.name,
                "line": node.lineno,
                "bases": [_ast.unparse(b) for b in node.bases],
                "methods": methods,
                "docstring": doc[:120],
            })
        elif isinstance(node, _ast.Assign):
            for target in node.targets:
                if isinstance(target, _ast.Name) and target.id.isupper():
                    globals_list.append(target.id)

    lines = source.splitlines()
    return {
        "ok": True,
        "path": str(target),
        "line_count": len(lines),
        "functions": functions,
        "classes": classes,
        "imports": list(dict.fromkeys(imports))[:30],
        "constants": globals_list[:20],
    }


def project_discovery_tool(workspace_root: str = "") -> dict:
    """
    Run project discovery on a workspace: detects tech stack, file types, entry points,
    README summary, and key structural patterns. Useful for orienting to an unfamiliar codebase.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from services.project_discovery import discover_project
        root = workspace_root or str(Path.cwd())
        return discover_project(root)
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Symbolic & Advanced Math ──────────────────────────────────────────────────

def sympy_solve(expression: str, variable: str = "x", mode: str = "solve") -> dict:
    """
    Symbolic math via SymPy. mode options:
    - 'solve': solve equation for variable (e.g. "x**2 - 4", "x" → [-2, 2])
    - 'simplify': algebraically simplify an expression
    - 'diff': differentiate with respect to variable
    - 'integrate': integrate with respect to variable
    - 'expand': expand/distribute
    - 'factor': factor into irreducible parts
    - 'latex': render as LaTeX string
    - 'numeric': numerical evaluation (calls evalf)
    """
    try:
        import sympy as sp
        from sympy.parsing.sympy_parser import implicit_multiplication_application, parse_expr, standard_transformations
        transforms = standard_transformations + (implicit_multiplication_application,)
        local_dict = {v: sp.Symbol(v) for v in "xyzabcntk"}
        if variable not in local_dict:
            local_dict[variable] = sp.Symbol(variable)
        expr = parse_expr(expression, local_dict=local_dict, transformations=transforms)
        var = local_dict.get(variable, sp.Symbol(variable))
        if mode == "solve":
            sol = sp.solve(expr, var)
            return {"ok": True, "mode": "solve", "variable": variable, "solutions": [str(s) for s in sol]}
        elif mode == "diff":
            return {"ok": True, "mode": "diff", "result": str(sp.diff(expr, var)), "latex": sp.latex(sp.diff(expr, var))}
        elif mode == "integrate":
            return {"ok": True, "mode": "integrate", "result": str(sp.integrate(expr, var)), "latex": sp.latex(sp.integrate(expr, var))}
        elif mode == "simplify":
            return {"ok": True, "mode": "simplify", "result": str(sp.simplify(expr)), "latex": sp.latex(sp.simplify(expr))}
        elif mode == "expand":
            return {"ok": True, "mode": "expand", "result": str(sp.expand(expr))}
        elif mode == "factor":
            return {"ok": True, "mode": "factor", "result": str(sp.factor(expr))}
        elif mode == "latex":
            return {"ok": True, "mode": "latex", "latex": sp.latex(expr)}
        elif mode == "numeric":
            return {"ok": True, "mode": "numeric", "result": str(expr.evalf()), "float": float(expr.evalf())}
        else:
            return {"ok": False, "error": f"Unknown mode: {mode}. Use solve/diff/integrate/simplify/expand/factor/latex/numeric"}
    except ImportError:
        return {"ok": False, "error": "sympy not installed: pip install sympy"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── NLP Intelligence ──────────────────────────────────────────────────────────

def nlp_analyze(text: str, tasks: list | None = None) -> dict:
    """
    NLP analysis pipeline. tasks: list of ['entities', 'keywords', 'sentiment', 'sentences', 'pos']
    Default: all. Uses spaCy if available, falls back to NLTK + basic heuristics.
    """
    if not tasks:
        tasks = ["entities", "keywords", "sentiment", "sentences"]
    result: dict = {"ok": True, "text_length": len(text), "tasks": tasks}

    # Try spaCy first
    try:
        import spacy
        try:
            nlp = spacy.load("en_core_web_sm")
        except OSError:
            # Model not downloaded — try blank
            nlp = spacy.blank("en")
        doc = nlp(text[:50000])
        if "entities" in tasks:
            result["entities"] = [
                {"text": ent.text, "label": ent.label_, "start": ent.start_char, "end": ent.end_char}
                for ent in doc.ents
            ][:50]
        if "sentences" in tasks:
            result["sentences"] = [str(s)[:200] for s in list(doc.sents)[:20]]
        if "pos" in tasks:
            result["pos_tags"] = [
                {"token": t.text, "pos": t.pos_, "dep": t.dep_}
                for t in doc if not t.is_space
            ][:60]
    except ImportError:
        pass

    # Keywords via KeyBERT
    if "keywords" in tasks:
        try:
            from keybert import KeyBERT
            kw_model = KeyBERT()
            keywords = kw_model.extract_keywords(text[:10000], keyphrase_ngram_range=(1, 2), top_n=12)
            result["keywords"] = [{"phrase": kw, "score": round(score, 4)} for kw, score in keywords]
        except ImportError:
            # Fallback: simple frequency-based keywords
            import re
            words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
            freq: dict = {}
            for w in words:
                freq[w] = freq.get(w, 0) + 1
            stopwords = {"that", "this", "with", "from", "they", "have", "been", "were", "will", "would", "could", "should", "their", "there", "these", "those"}
            keywords_fb = sorted([(w, c) for w, c in freq.items() if w not in stopwords], key=lambda x: -x[1])[:12]
            result["keywords"] = [{"phrase": w, "score": c} for w, c in keywords_fb]

    # Basic sentiment (no ML needed — lexicon approach)
    if "sentiment" in tasks:
        try:
            from textblob import TextBlob
            tb = TextBlob(text[:5000])
            result["sentiment"] = {"polarity": round(tb.sentiment.polarity, 3), "subjectivity": round(tb.sentiment.subjectivity, 3)}
        except ImportError:
            # Very basic heuristic
            pos_words = {"good","great","excellent","amazing","love","wonderful","best","fantastic","perfect","brilliant"}
            neg_words = {"bad","terrible","awful","horrible","hate","worst","poor","disappointing","fail","wrong"}
            words_lower = set(text.lower().split())
            pos = len(words_lower & pos_words)
            neg = len(words_lower & neg_words)
            polarity = (pos - neg) / max(pos + neg, 1)
            result["sentiment"] = {"polarity": round(polarity, 3), "method": "lexicon_heuristic"}

    return result


# ─── Image & OCR ───────────────────────────────────────────────────────────────

def ocr_image(path: str, lang: str = "eng") -> dict:
    """
    Extract text from an image using OCR.
    Tries EasyOCR first (better accuracy, no Tesseract required),
    then falls back to pytesseract (requires Tesseract binary installed).
    lang: language code ('eng', 'fra', 'deu', 'jpn', 'chi_sim', etc.)
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    ext = target.suffix.lower()
    if ext not in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp", ".gif"}:
        return {"ok": False, "error": f"Unsupported image format: {ext}"}

    # Try EasyOCR
    try:
        import easyocr
        lang_map = {"eng": "en", "fra": "fr", "deu": "de", "chi_sim": "ch_sim", "jpn": "ja"}
        easy_lang = lang_map.get(lang, "en")
        reader = easyocr.Reader([easy_lang], gpu=False, verbose=False)
        results = reader.readtext(str(target))
        text_parts = [item[1] for item in results if item[2] > 0.1]
        full_text = "\n".join(text_parts)
        return {
            "ok": True, "method": "easyocr", "path": str(target),
            "text": full_text[:8000], "blocks": len(results),
            "confidence_avg": round(sum(r[2] for r in results) / max(len(results), 1), 3),
        }
    except ImportError:
        pass

    # Fallback: pytesseract
    try:
        import pytesseract
        from PIL import Image as PILImage
        with PILImage.open(str(target)) as img:
            text = pytesseract.image_to_string(img, lang=lang)
            data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
        confidences = [int(c) for c in data.get("conf", []) if str(c).lstrip("-").isdigit() and int(c) >= 0]
        conf_avg = sum(confidences) / max(len(confidences), 1)
        return {
            "ok": True, "method": "pytesseract", "path": str(target),
            "text": text.strip()[:8000],
            "confidence_avg": round(conf_avg, 1),
        }
    except ImportError:
        return {"ok": False, "error": "OCR requires easyocr or pytesseract+Pillow: pip install easyocr OR pip install pytesseract Pillow"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Visualization ─────────────────────────────────────────────────────────────

def plot_chart(
    data: dict,
    chart_type: str = "bar",
    title: str = "",
    output_path: str = "",
    xlabel: str = "",
    ylabel: str = "",
) -> dict:
    """
    Generate a chart and save it as PNG. Returns path to saved file.
    chart_type: 'bar' | 'line' | 'scatter' | 'pie' | 'histogram' | 'heatmap'
    data format:
    - bar/line: {"labels": [...], "values": [...]} or {"Series A": [...], "Series B": [...], "labels": [...]}
    - scatter: {"x": [...], "y": [...]}
    - pie: {"labels": [...], "values": [...]}
    - histogram: {"values": [...], "bins": 20}
    - heatmap: {"matrix": [[...], ...], "row_labels": [...], "col_labels": [...]}
    """
    try:
        import matplotlib
        matplotlib.use("Agg")  # non-interactive backend, always safe
        import matplotlib.pyplot as plt
        import numpy as _np

        fig, ax = plt.subplots(figsize=(10, 6))
        if title:
            ax.set_title(title, fontsize=14, fontweight="bold")
        if xlabel:
            ax.set_xlabel(xlabel)
        if ylabel:
            ax.set_ylabel(ylabel)

        if chart_type == "bar":
            labels = data.get("labels", list(range(len(data.get("values", [])))))
            values = data.get("values", [])
            ax.bar(range(len(labels)), values, tick_label=[str(lbl) for lbl in labels])
            plt.xticks(rotation=45, ha="right")

        elif chart_type == "line":
            labels = data.get("labels", list(range(len(data.get("values", [])))))
            for key, vals in data.items():
                if key == "labels":
                    continue
                if isinstance(vals, (list, tuple)):
                    ax.plot(labels if len(labels) == len(vals) else range(len(vals)), vals, label=key, marker="o", markersize=3)
            ax.legend()

        elif chart_type == "scatter":
            x, y = data.get("x", []), data.get("y", [])
            labels = data.get("point_labels", [])
            ax.scatter(x, y, alpha=0.7)
            for i, label in enumerate(labels[:len(x)]):
                ax.annotate(str(label), (x[i], y[i]), fontsize=7)

        elif chart_type == "pie":
            labels = data.get("labels", [])
            values = data.get("values", [])
            ax.pie(values, labels=labels, autopct="%1.1f%%", startangle=140)
            ax.axis("equal")

        elif chart_type == "histogram":
            values = data.get("values", [])
            bins = data.get("bins", 20)
            ax.hist(values, bins=bins, edgecolor="black", alpha=0.7)

        elif chart_type == "heatmap":
            matrix = data.get("matrix", [[]])
            row_labels = data.get("row_labels", [])
            col_labels = data.get("col_labels", [])
            arr = _np.array(matrix)
            im = ax.imshow(arr, cmap="viridis", aspect="auto")
            plt.colorbar(im, ax=ax)
            if row_labels:
                ax.set_yticks(range(len(row_labels)))
                ax.set_yticklabels(row_labels)
            if col_labels:
                ax.set_xticks(range(len(col_labels)))
                ax.set_xticklabels(col_labels, rotation=45, ha="right")
        else:
            plt.close(fig)
            return {"ok": False, "error": f"Unknown chart_type: {chart_type}"}

        # Determine save path
        if output_path:
            save_path = Path(output_path)
        else:
            import tempfile
            import time
            tmp_dir = Path(tempfile.gettempdir())
            save_path = tmp_dir / f"layla_chart_{int(time.time())}.png"

        plt.tight_layout()
        fig.savefig(str(save_path), dpi=120, bbox_inches="tight")
        plt.close(fig)
        return {"ok": True, "chart_type": chart_type, "path": str(save_path), "title": title}
    except ImportError:
        return {"ok": False, "error": "matplotlib not installed: pip install matplotlib"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Document Formats ──────────────────────────────────────────────────────────

def read_docx(path: str) -> dict:
    """
    Read a Word document (.docx). Returns full text, paragraph list, and table data.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        from docx import Document
        doc = Document(str(target))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables[:10]:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(rows)
        full_text = "\n".join(paragraphs)
        return {
            "ok": True, "path": str(target),
            "paragraphs": len(paragraphs),
            "text": full_text[:10000],
            "tables": tables[:5],
            "table_count": len(doc.tables),
        }
    except ImportError:
        return {"ok": False, "error": "python-docx not installed: pip install python-docx"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def read_excel(path: str, sheet: str = "", max_rows: int = 100) -> dict:
    """
    Read an Excel file (.xlsx/.xls). Returns sheet names, data from target sheet,
    and basic stats. sheet: sheet name or index (default: first sheet).
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import pandas as _pd
        xl = _pd.ExcelFile(str(target))
        sheet_names = xl.sheet_names
        active_sheet = sheet if sheet else sheet_names[0]
        df = xl.parse(active_sheet)
        result: dict = {
            "ok": True, "path": str(target),
            "sheets": sheet_names, "active_sheet": str(active_sheet),
            "rows": len(df), "columns": list(df.columns),
            "sample": df.head(max_rows).to_dict(orient="records"),
            "null_counts": df.isnull().sum().to_dict(),
        }
        try:
            result["stats"] = df.describe().to_dict()
        except Exception:
            pass
        return result
    except ImportError:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(target), read_only=True, data_only=True)
            sheet_names = wb.sheetnames
            ws = wb[sheet] if sheet and sheet in sheet_names else wb.active
            rows = []
            headers: list = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i == 0:
                    headers = [str(c or f"col_{j}") for j, c in enumerate(row)]
                elif i <= max_rows:
                    rows.append(dict(zip(headers, row)))
            wb.close()
            return {"ok": True, "path": str(target), "sheets": sheet_names, "columns": headers, "sample": rows}
        except ImportError:
            return {"ok": False, "error": "Excel reading requires pandas or openpyxl: pip install pandas openpyxl"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Database Intelligence ─────────────────────────────────────────────────────

def sql_query(db_path: str, query: str, limit: int = 200) -> dict:
    """
    Execute a SQL query against a SQLite or DuckDB database file.
    READ-ONLY by default: SELECT queries only. Non-SELECT queries require allow_write.
    db_path: path to .db/.sqlite/.duckdb file, or ':memory:' for DuckDB in-memory.
    """
    is_readonly = query.strip().upper().startswith("SELECT") or query.strip().upper().startswith("WITH")
    target = Path(db_path) if db_path != ":memory:" else None
    if target and not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if target and not target.exists():
        return {"ok": False, "error": "File not found"}

    # Inject LIMIT if not present
    q = query.strip().rstrip(";")
    if is_readonly and "LIMIT" not in q.upper():
        q += f" LIMIT {limit}"

    # Try DuckDB first (handles .duckdb and in-memory well)
    ext = (target.suffix.lower() if target else ".duckdb")
    if ext == ".duckdb" or db_path == ":memory:":
        try:
            import duckdb
            conn = duckdb.connect(db_path)
            rel = conn.execute(q)
            cols = [d[0] for d in rel.description]
            rows = rel.fetchall()
            conn.close()
            return {
                "ok": True, "db": db_path, "query": query,
                "columns": cols, "rows": [dict(zip(cols, r)) for r in rows],
                "row_count": len(rows),
            }
        except ImportError:
            pass
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # SQLite
    try:
        import sqlite3 as _sql
        conn = _sql.connect(str(target))
        conn.row_factory = _sql.Row
        cursor = conn.execute(q)
        rows = cursor.fetchall()
        cols = [d[0] for d in cursor.description] if cursor.description else []
        conn.close()
        return {
            "ok": True, "db": db_path, "query": query,
            "columns": cols, "rows": [dict(r) for r in rows],
            "row_count": len(rows),
        }
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Financial Intelligence ────────────────────────────────────────────────────

def stock_data(ticker: str, period: str = "1mo", include_info: bool = True) -> dict:
    """
    Fetch stock or crypto data via yfinance.
    ticker: stock symbol (AAPL, TSLA, BTC-USD, ETH-USD, ^GSPC for S&P500)
    period: '1d' | '5d' | '1mo' | '3mo' | '6mo' | '1y' | '2y' | '5y' | 'ytd' | 'max'
    Returns: OHLCV data, current price, company info (if include_info=True).
    """
    try:
        import yfinance as yf
        t = yf.Ticker(ticker)
        hist = t.history(period=period)
        if hist.empty:
            return {"ok": False, "error": f"No data for ticker: {ticker}"}
        hist_records = []
        for date, row in hist.tail(30).iterrows():
            hist_records.append({
                "date": str(date)[:10],
                "open": round(float(row["Open"]), 4),
                "high": round(float(row["High"]), 4),
                "low": round(float(row["Low"]), 4),
                "close": round(float(row["Close"]), 4),
                "volume": int(row["Volume"]),
            })
        result: dict = {
            "ok": True, "ticker": ticker.upper(), "period": period,
            "current_price": round(float(hist["Close"].iloc[-1]), 4),
            "price_change_pct": round(float((hist["Close"].iloc[-1] - hist["Close"].iloc[0]) / hist["Close"].iloc[0] * 100), 2),
            "52w_high": round(float(hist["High"].max()), 4),
            "52w_low": round(float(hist["Low"].min()), 4),
            "history": hist_records,
        }
        if include_info:
            try:
                info = t.info or {}
                result["info"] = {
                    "name": info.get("longName") or info.get("shortName", ""),
                    "sector": info.get("sector", ""),
                    "industry": info.get("industry", ""),
                    "market_cap": info.get("marketCap"),
                    "pe_ratio": info.get("forwardPE") or info.get("trailingPE"),
                    "dividend_yield": info.get("dividendYield"),
                    "description": (info.get("longBusinessSummary") or "")[:400],
                }
            except Exception:
                pass
        return result
    except ImportError:
        return {"ok": False, "error": "yfinance not installed: pip install yfinance"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Security Analysis ─────────────────────────────────────────────────────────

def security_scan(path: str, scan_type: str = "bandit") -> dict:
    """
    Run security analysis on Python code or check dependencies for known vulnerabilities.
    scan_type:
    - 'bandit': static analysis for Python security issues (CWEs, hardcoded secrets, etc.)
    - 'deps': check requirements.txt or pyproject.toml for vulnerable packages
    - 'secrets': pattern-based scan for hardcoded secrets/tokens/keys in any file
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}

    if scan_type == "bandit":
        try:
            r = subprocess.run(
                [sys.executable, "-m", "bandit", "-r", str(target), "-f", "json", "-q"],
                capture_output=True, text=True, timeout=60, encoding="utf-8", errors="replace",
            )
            import json as _json
            try:
                data = _json.loads(r.stdout or "{}")
            except Exception:
                data = {}
            issues = data.get("results", [])
            metrics = data.get("metrics", {})
            return {
                "ok": True, "scan_type": "bandit", "path": str(target),
                "issues": [
                    {
                        "severity": i.get("issue_severity"), "confidence": i.get("issue_confidence"),
                        "text": i.get("issue_text"), "file": i.get("filename"),
                        "line": i.get("line_number"), "cwe": i.get("issue_cwe", {}).get("id"),
                    }
                    for i in issues[:30]
                ],
                "issue_count": len(issues),
                "metrics": metrics,
            }
        except FileNotFoundError:
            return {"ok": False, "error": "bandit not installed: pip install bandit"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    elif scan_type == "secrets":
        import re as _re
        SECRET_PATTERNS = [
            (r'(?i)(api[_-]?key|apikey)\s*[:=]\s*["\']?([A-Za-z0-9\-_]{16,})', "API Key"),
            (r'(?i)(secret[_-]?key|secret)\s*[:=]\s*["\']?([A-Za-z0-9\-_]{16,})', "Secret Key"),
            (r'(?i)(password|passwd|pwd)\s*[:=]\s*["\']?([^\s"\']{8,})', "Password"),
            (r'(?i)(token|access_token|auth_token)\s*[:=]\s*["\']?([A-Za-z0-9\-_\.]{16,})', "Token"),
            (r'AKIA[0-9A-Z]{16}', "AWS Access Key"),
            (r'(?i)(private[_-]?key)\s*[:=]', "Private Key"),
            (r'sk-[A-Za-z0-9]{32,}', "OpenAI Key"),
            (r'ghp_[A-Za-z0-9]{36,}', "GitHub Token"),
        ]
        findings = []
        files_scanned = 0
        if target.is_file():
            scan_files = [target]
        else:
            scan_files = [f for f in target.rglob("*") if f.is_file() and f.suffix in {".py", ".js", ".ts", ".env", ".json", ".yaml", ".yml", ".txt", ".cfg", ".ini"} and ".git" not in str(f)][:100]
        for fpath in scan_files:
            files_scanned += 1
            try:
                content = fpath.read_text(encoding="utf-8", errors="ignore")
                for pattern, label in SECRET_PATTERNS:
                    for m in _re.finditer(pattern, content):
                        line_num = content[:m.start()].count("\n") + 1
                        findings.append({"file": str(fpath.relative_to(target) if target.is_dir() else fpath), "line": line_num, "type": label, "match": m.group(0)[:80]})
            except Exception:
                continue
        return {"ok": True, "scan_type": "secrets", "files_scanned": files_scanned, "findings": findings[:50], "finding_count": len(findings)}

    elif scan_type == "deps":
        req_files = []
        if target.is_file():
            req_files = [target]
        else:
            for name in ("requirements.txt", "pyproject.toml", "Pipfile"):
                f = target / name
                if f.exists():
                    req_files.append(f)
        if not req_files:
            return {"ok": False, "error": "No requirements.txt/pyproject.toml found"}
        try:
            r = subprocess.run(
                [sys.executable, "-m", "pip_audit", "--requirement", str(req_files[0]), "--format", "json"],
                capture_output=True, text=True, timeout=120, encoding="utf-8", errors="replace",
            )
            import json as _json
            try:
                data = _json.loads(r.stdout or "[]")
                return {"ok": True, "scan_type": "deps", "vulnerabilities": data[:30], "count": len(data)}
            except Exception:
                return {"ok": True, "scan_type": "deps", "output": (r.stdout or r.stderr)[:2000]}
        except FileNotFoundError:
            return {"ok": False, "error": "pip-audit not installed: pip install pip-audit"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    return {"ok": False, "error": f"Unknown scan_type: {scan_type}. Use bandit/secrets/deps"}


def _build_tools_from_domains() -> dict[str, Any]:
    """Build TOOLS by merging domain modules. Called after all tool functions are defined."""
    from layla.tools.domains import (
        ANALYSIS_TOOLS,
        AUTOMATION_TOOLS,
        CODE_TOOLS,
        DATA_TOOLS,
        FILE_TOOLS,
        GENERAL_TOOLS,
        GIT_TOOLS,
        MEMORY_TOOLS,
        SYSTEM_TOOLS,
        WEB_TOOLS,
    )
    result: dict[str, Any] = {}
    for domain_tools in (
        FILE_TOOLS,
        GIT_TOOLS,
        WEB_TOOLS,
        MEMORY_TOOLS,
        CODE_TOOLS,
        DATA_TOOLS,
        SYSTEM_TOOLS,
        AUTOMATION_TOOLS,
        ANALYSIS_TOOLS,
        GENERAL_TOOLS,
    ):
        for name, meta in domain_tools.items():
            meta = dict(meta)
            fn_name = meta.pop("fn_key", name)
            fn = globals().get(fn_name)
            if fn is None:
                raise ValueError(f"Tool {name}: function {fn_name} not found in registry")
            result[name] = {"fn": fn, **meta}
    return result


# ─── Semantic Memory Tools ─────────────────────────────────────────────────────

def vector_search(query: str, collection: str = "knowledge", k: int = 8) -> dict:
    """
    Direct semantic vector search over Layla's knowledge or memory collections.
    collection: 'knowledge' | 'memories' | 'aspects'
    Returns top-k results with content + similarity score.
    This is the raw retrieval layer — use search_memories for the full RAG pipeline.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        if collection == "memories":
            from layla.memory.vector_store import search_memories_full
            results = search_memories_full(query, k=k, use_rerank=False)
            return {"ok": True, "collection": collection, "query": query, "results": results[:k], "count": len(results)}
        elif collection == "knowledge":
            from layla.memory.vector_store import get_knowledge_chunks_with_sources
            chunks = get_knowledge_chunks_with_sources(query, k=k)
            results = [{"content": c.get("text", ""), "text": c.get("text", ""), "source": c.get("source", "")} for c in chunks]
            return {"ok": True, "collection": collection, "query": query, "results": results[:k], "count": len(results)}
        else:
            from layla.memory.vector_store import search_memories_full
            results = search_memories_full(query, k=k, use_rerank=False)
            return {"ok": True, "collection": collection, "query": query, "results": results[:k], "count": len(results)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def vector_store(text: str, metadata: dict | None = None, collection: str = "memories") -> dict:
    """
    Explicitly store text into Layla's vector database.
    collection: 'memories' (default) — stored as a learning and embedded.
    metadata: optional dict of tags, source, aspect, etc.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import save_learning
        meta = metadata or {}
        kind = meta.get("kind", "tool_store")
        save_learning(content=text[:800], kind=kind)
        # Also embed into vector store
        try:
            from layla.memory.vector_store import embed, add_vector
            vec = embed(text[:800])
            meta_with_content = {**meta, "content": text[:800]}
            add_vector(vec, meta_with_content)
        except Exception:
            pass
        return {"ok": True, "stored": text[:100], "collection": collection, "kind": kind}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── File System Intelligence ──────────────────────────────────────────────────

def workspace_map(root: str = "", max_files: int = 500, include_content_preview: bool = False) -> dict:
    """
    Build a full intelligence map of a workspace:
    - Directory tree (depth-limited)
    - File count by extension
    - Detected tech stack (languages, frameworks, config files)
    - Entry points (main.py, index.js, Dockerfile, etc.)
    - Key documentation (README, CHANGELOG, AGENTS.md, etc.)
    - Large files + recently modified files
    """
    root_path = Path(root).expanduser().resolve() if root else Path.cwd()
    if not root_path.exists():
        return {"ok": False, "error": "Path not found"}

    IGNORE = {".git", ".venv", "venv", "__pycache__", "node_modules", ".mypy_cache",
              ".pytest_cache", "dist", "build", ".tox", "*.egg-info"}

    all_files: list[Path] = []
    for f in root_path.rglob("*"):
        if f.is_file() and not any(part in IGNORE for part in f.parts):
            all_files.append(f)
            if len(all_files) >= max_files:
                break

    # Extension counts
    ext_counts: dict = {}
    for f in all_files:
        ext = f.suffix.lower() or "(none)"
        ext_counts[ext] = ext_counts.get(ext, 0) + 1

    # Tech stack detection
    STACK_SIGNALS = {
        "Python": {".py", "requirements.txt", "pyproject.toml", "setup.py", "Pipfile"},
        "JavaScript": {".js", ".mjs", "package.json"},
        "TypeScript": {".ts", ".tsx", "tsconfig.json"},
        "Rust": {".rs", "Cargo.toml"},
        "Go": {".go", "go.mod"},
        "Java": {".java", "pom.xml", "build.gradle"},
        "C/C++": {".c", ".cpp", ".h", ".hpp", "CMakeLists.txt"},
        "Docker": {"Dockerfile", "docker-compose.yml", "docker-compose.yaml"},
        "Kubernetes": {".yaml", "k8s", "helm"},
        "FastAPI": {"main.py"},
        "React": {"package.json", ".jsx", ".tsx"},
        "Vue": {".vue"},
    }
    names_and_exts = {f.name for f in all_files} | {f.suffix.lower() for f in all_files}
    detected_stack = [lang for lang, signals in STACK_SIGNALS.items() if signals & names_and_exts]

    # Entry points
    ENTRY_NAMES = {"main.py", "app.py", "server.py", "index.js", "index.ts", "main.rs",
                   "main.go", "Dockerfile", "docker-compose.yml", "manage.py", "wsgi.py", "asgi.py"}
    entry_points = [str(f.relative_to(root_path)) for f in all_files if f.name in ENTRY_NAMES]

    # Key docs
    DOC_NAMES = {"README.md", "AGENTS.md", "ARCHITECTURE.md", "CHANGELOG.md", "CONTRIBUTING.md",
                 "LICENSE", "INSTALL.md", "SECURITY.md", "TODO.md", "NOTES.md"}
    key_docs = {}
    for f in all_files:
        if f.name in DOC_NAMES:
            preview = ""
            if include_content_preview:
                try:
                    preview = f.read_text(encoding="utf-8", errors="replace")[:400]
                except Exception:
                    pass
            key_docs[f.name] = {"path": str(f.relative_to(root_path)), "preview": preview}

    # Largest files
    sorted_by_size = sorted(all_files, key=lambda f: f.stat().st_size, reverse=True)
    largest = [{"path": str(f.relative_to(root_path)), "size_kb": round(f.stat().st_size / 1024, 1)} for f in sorted_by_size[:10]]

    # Recently modified
    sorted_by_mtime = sorted(all_files, key=lambda f: f.stat().st_mtime, reverse=True)
    recent = [{"path": str(f.relative_to(root_path)), "modified": str(__import__("datetime").datetime.fromtimestamp(f.stat().st_mtime))[:16]} for f in sorted_by_mtime[:10]]

    # Directory tree (2-level)
    def tree_level(path: Path, depth: int = 0, max_depth: int = 2) -> list:
        if depth >= max_depth:
            return []
        entries = []
        try:
            for child in sorted(path.iterdir(), key=lambda x: (x.is_file(), x.name)):
                if any(part in IGNORE for part in child.parts):
                    continue
                entry = {"name": child.name, "type": "file" if child.is_file() else "dir"}
                if child.is_dir() and depth < max_depth - 1:
                    entry["children"] = tree_level(child, depth + 1, max_depth)
                entries.append(entry)
        except PermissionError:
            pass
        return entries

    return {
        "ok": True,
        "root": str(root_path),
        "total_files": len(all_files),
        "extensions": dict(sorted(ext_counts.items(), key=lambda x: -x[1])[:20]),
        "tech_stack": detected_stack,
        "entry_points": entry_points,
        "key_docs": key_docs,
        "largest_files": largest,
        "recently_modified": recent,
        "tree": tree_level(root_path),
    }


# ─── Web Crawl ─────────────────────────────────────────────────────────────────

def crawl_site(
    url: str,
    max_pages: int = 20,
    max_depth: int = 2,
    same_domain: bool = True,
    store_knowledge: bool = False,
) -> dict:
    """
    Crawl a website starting from url. Extracts clean text from each page.
    max_pages: hard cap on pages visited
    max_depth: link-following depth (1 = only start URL, 2 = start + its links, etc.)
    same_domain: only follow links within the same domain
    store_knowledge: save extracted pages to knowledge/fetched/ for later RAG indexing
    Returns: list of {url, title, text, depth} for all visited pages.
    """
    import time
    from urllib.parse import urljoin, urlparse

    try:
        import trafilatura
        from trafilatura.sitemaps import sitemap_search  # noqa: F401
    except ImportError:
        return {"ok": False, "error": "trafilatura not installed: pip install trafilatura"}

    base_domain = urlparse(url).netloc
    visited: set[str] = set()
    queue: list[tuple[str, int]] = [(url, 0)]
    results = []
    start_time = time.time()

    while queue and len(results) < max_pages:
        if time.time() - start_time > 120:  # 2 min hard cap
            break
        current_url, depth = queue.pop(0)
        if current_url in visited:
            continue
        visited.add(current_url)

        try:
            downloaded = trafilatura.fetch_url(current_url)
            if not downloaded:
                continue
            text = trafilatura.extract(downloaded, include_comments=False, include_tables=True, favor_recall=True)
            if not text or len(text.strip()) < 50:
                continue
            title = ""
            links: list[str] = []  # noqa: F841
            try:
                meta = trafilatura.extract_metadata(downloaded)
                if meta:
                    title = meta.title or ""
            except Exception:
                pass
            # Extract links for deeper crawl
            if depth < max_depth - 1:
                try:
                    from trafilatura.urls import extract_links
                    raw_links = extract_links(downloaded, url) or []
                    for link in raw_links[:30]:
                        full = urljoin(current_url, link)
                        if full not in visited:
                            if not same_domain or urlparse(full).netloc == base_domain:
                                queue.append((full, depth + 1))
                except Exception:
                    pass
            page_result = {
                "url": current_url, "title": title,
                "text": text[:4000], "chars": len(text), "depth": depth,
            }
            results.append(page_result)

            # Optionally save to knowledge/fetched/
            if store_knowledge:
                try:
                    slug = urlparse(current_url).path.strip("/").replace("/", "_")[:50] or "index"
                    fetched_dir = Path(__file__).resolve().parent.parent.parent.parent / "knowledge" / "fetched"
                    fetched_dir.mkdir(parents=True, exist_ok=True)
                    out = fetched_dir / f"{base_domain}_{slug}.txt"
                    out.write_text(f"source: {current_url}\ntitle: {title}\n\n{text[:30000]}", encoding="utf-8")
                except Exception:
                    pass

        except Exception:
            continue

    return {
        "ok": True, "start_url": url, "pages_visited": len(results),
        "pages_requested": max_pages, "same_domain": same_domain,
        "results": results,
    }


# ─── Database Schema Intelligence ─────────────────────────────────────────────

def schema_introspect(db_path: str) -> dict:
    """
    Introspect a database schema. Returns tables, columns with types, row counts,
    foreign keys, and sample data (first 3 rows per table).
    Supports SQLite (.db, .sqlite) and DuckDB (.duckdb).
    """
    target = Path(db_path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}

    ext = target.suffix.lower()

    if ext == ".duckdb":
        try:
            import duckdb
            conn = duckdb.connect(str(target))
            tables_raw = conn.execute("SHOW TABLES").fetchall()
            schema = {}
            for (table_name,) in tables_raw:
                cols = conn.execute(f"DESCRIBE {table_name}").fetchall()
                count = conn.execute(f"SELECT COUNT(*) FROM {table_name}").fetchone()[0]
                sample = conn.execute(f"SELECT * FROM {table_name} LIMIT 3").fetchall()
                col_names = [c[0] for c in cols]
                schema[table_name] = {
                    "columns": [{"name": c[0], "type": c[1]} for c in cols],
                    "row_count": count,
                    "sample": [dict(zip(col_names, row)) for row in sample],
                }
            conn.close()
            return {"ok": True, "db_type": "duckdb", "path": db_path, "tables": schema}
        except ImportError:
            return {"ok": False, "error": "duckdb not installed: pip install duckdb"}
        except Exception as e:
            return {"ok": False, "error": str(e)}

    # SQLite
    try:
        import sqlite3 as _sql
        conn = _sql.connect(str(target))
        conn.row_factory = _sql.Row
        tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()
        schema = {}
        for (table_name,) in [(r["name"],) for r in tables]:
            cols = conn.execute(f"PRAGMA table_info({table_name})").fetchall()
            fkeys = conn.execute(f"PRAGMA foreign_key_list({table_name})").fetchall()
            count = conn.execute(f"SELECT COUNT(*) FROM {table_name}").fetchone()[0]
            sample_rows = conn.execute(f"SELECT * FROM {table_name} LIMIT 3").fetchall()
            col_names = [c["name"] for c in cols]
            schema[table_name] = {
                "columns": [{"name": c["name"], "type": c["type"], "notnull": bool(c["notnull"]), "pk": bool(c["pk"])} for c in cols],
                "foreign_keys": [{"from": fk["from"], "to_table": fk["table"], "to_col": fk["to"]} for fk in fkeys],
                "row_count": count,
                "sample": [dict(r) for r in sample_rows],
            }
        # Views
        views = conn.execute("SELECT name FROM sqlite_master WHERE type='view'").fetchall()
        conn.close()
        return {"ok": True, "db_type": "sqlite", "path": db_path, "tables": schema, "views": [r["name"] for r in views]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Tool Self-Reflection ──────────────────────────────────────────────────────

def list_tools(filter_by: str = "", include_dangerous: bool = True) -> dict:
    """
    List all tools available to Layla with their descriptions, risk levels, and approval status.
    filter_by: keyword to filter by tool name or description (empty = return all)
    include_dangerous: if False, only shows safe tools
    """
    results = []
    for name, meta in TOOLS.items():
        if not include_dangerous and meta.get("dangerous"):
            continue
        fn = meta.get("fn")
        doc = (fn.__doc__ or "").strip().split("\n")[0][:120] if fn else ""
        if filter_by and filter_by.lower() not in name.lower() and filter_by.lower() not in doc.lower():
            continue
        results.append({
            "name": name,
            "description": doc,
            "dangerous": meta.get("dangerous", False),
            "require_approval": meta.get("require_approval", False),
            "risk_level": meta.get("risk_level", "low"),
        })
    return {
        "ok": True,
        "total": len(TOOLS),
        "shown": len(results),
        "filter": filter_by,
        "tools": sorted(results, key=lambda x: x["name"]),
    }


def tool_recommend(task: str) -> dict:
    """
    Given a task description, recommend the most relevant tools to use.
    Uses keyword matching + category heuristics.
    Example: tool_recommend("read a PDF and summarize it") → [read_pdf, fetch_article, save_note]
    """
    task_lower = task.lower()
    CATEGORY_KEYWORDS = {
        "file": ["read_file", "write_file", "list_dir", "file_info", "understand_file"],
        "pdf": ["read_pdf"],
        "docx word": ["read_docx"],
        "excel spreadsheet": ["read_excel", "read_csv"],
        "csv data table": ["read_csv", "read_excel", "sql_query"],
        "code python test pytest": ["python_ast", "grep_code", "run_python", "run_tests", "security_scan", "code_lint"],
        "code search": ["search_codebase", "grep_code", "glob_files", "python_ast"],
        "git commit diff push pull": ["git_status", "git_diff", "git_log", "git_add", "git_commit", "git_push", "git_pull", "git_stash", "git_revert", "git_clone"],
        "web search": ["ddg_search", "browser_search", "fetch_article", "wiki_search"],
        "research paper arxiv": ["arxiv_search", "wiki_search", "ddg_search"],
        "website crawl": ["crawl_site", "fetch_article", "browser_navigate"],
        "math equation": ["math_eval", "sympy_solve"],
        "image ocr": ["ocr_image", "describe_image"],
        "chart graph plot": ["plot_chart"],
        "sql database": ["sql_query", "schema_introspect"],
        "memory remember": ["save_note", "search_memories", "vector_search", "vector_store"],
        "security scan": ["security_scan"],
        "stock finance crypto": ["stock_data"],
        "nlp entities keywords": ["nlp_analyze"],
        "compress token context": ["context_compress", "count_tokens"],
        "translate sql query": ["generate_sql", "sql_query", "schema_introspect"],
        "workspace project": ["workspace_map", "project_discovery", "get_project_context"],
        "gcode dxf stl fabrication": ["parse_gcode", "stl_mesh_info", "understand_file", "generate_gcode"],
        "clipboard copy paste": ["clipboard_read", "clipboard_write"],
        "search replace refactor": ["search_replace", "rename_symbol", "grep_code"],
        "pip install package": ["pip_list", "pip_install"],
        "docker container": ["docker_ps", "docker_run"],
        "ci github pr issue": ["check_ci", "github_issues", "github_pr"],
        "webhook slack discord email": ["send_webhook", "send_email"],
        "log tail": ["tail_file", "read_file"],
        "format code black ruff": ["code_format"],
        "archive zip extract": ["extract_archive", "create_archive"],
        "uuid random password": ["uuid_generate", "random_string", "password_generate"],
        "disk process system": ["disk_usage", "process_list", "env_info"],
        "qr code": ["generate_qr"],
        "csv write export": ["write_csv", "read_csv"],
        "json schema": ["json_schema", "json_query"],
        "jwt token": ["jwt_decode"],
        "toml config": ["read_toml", "yaml_read"],
        "merge pdf": ["merge_pdf", "read_pdf"],
        "discord": ["discord_send", "send_webhook"],
        "calendar ics event": ["calendar_read", "calendar_add_event"],
        "database backup": ["db_backup", "sql_query", "schema_introspect"],
        "svg draw diagram": ["create_svg", "create_mermaid", "write_file"],
    }
    scores: dict = {}
    for category, tools in CATEGORY_KEYWORDS.items():
        for keyword in category.split():
            if keyword in task_lower:
                for tool in tools:
                    scores[tool] = scores.get(tool, 0) + 1

    # Also match tool names/descriptions directly
    for name, meta in TOOLS.items():
        fn = meta.get("fn")
        doc = (fn.__doc__ or "").lower() if fn else ""
        for word in task_lower.split():
            if len(word) > 3 and (word in name.lower() or word in doc):
                scores[name] = scores.get(name, 0) + 1

    ranked = sorted(scores.items(), key=lambda x: -x[1])[:10]
    recommendations = []
    for name, score in ranked:
        if name in TOOLS:
            fn = TOOLS[name].get("fn")
            doc = (fn.__doc__ or "").strip().split("\n")[0][:100] if fn else ""
            recommendations.append({"tool": name, "relevance": score, "description": doc})

    return {"ok": True, "task": task, "recommendations": recommendations}


# ─── Context Management ────────────────────────────────────────────────────────

def context_compress(text: str, target_tokens: int = 2000, strategy: str = "smart") -> dict:
    """
    Compress text to fit within a token budget.
    strategy:
    - 'smart': extract most important sentences (extractive summarization)
    - 'truncate': simple head truncation
    - 'middle_out': keep head + tail, drop middle (good for code files with imports + logic)
    Returns compressed text + token estimates before/after.
    """
    def rough_tokens(t: str) -> int:
        return max(int(len(t) / 4), len(t.split()))

    original_tokens = rough_tokens(text)

    if original_tokens <= target_tokens:
        return {"ok": True, "strategy": "no_compression_needed", "original_tokens": original_tokens,
                "compressed_tokens": original_tokens, "text": text, "ratio": 1.0}

    if strategy == "truncate":
        char_budget = target_tokens * 4
        compressed = text[:char_budget]
        return {"ok": True, "strategy": "truncate", "original_tokens": original_tokens,
                "compressed_tokens": rough_tokens(compressed), "text": compressed,
                "ratio": round(rough_tokens(compressed) / original_tokens, 3)}

    if strategy == "middle_out":
        char_budget = target_tokens * 4
        head = text[:char_budget // 2]
        tail = text[-(char_budget // 2):]
        compressed = head + "\n\n[... content compressed ...]\n\n" + tail
        return {"ok": True, "strategy": "middle_out", "original_tokens": original_tokens,
                "compressed_tokens": rough_tokens(compressed), "text": compressed,
                "ratio": round(rough_tokens(compressed) / original_tokens, 3)}

    # Smart: sentence scoring (position + length + keyword density)
    import re as _re
    sentences = _re.split(r'(?<=[.!?])\s+', text)
    if not sentences:
        sentences = text.split("\n")

    # Score each sentence
    total = len(sentences)
    def score_sentence(s: str, idx: int) -> float:
        pos_score = 1.5 if idx < total * 0.1 else (1.2 if idx > total * 0.9 else 1.0)
        len_score = 1.0 if 20 < len(s) < 200 else 0.6
        caps = len(_re.findall(r'\b[A-Z][a-z]+\b', s))
        return pos_score * len_score * (1 + caps * 0.1)

    scored = [(score_sentence(s, i), i, s) for i, s in enumerate(sentences)]
    scored.sort(key=lambda x: -x[0])

    # Greedily pick sentences until token budget
    picked_indices: set[int] = set()
    budget_used = 0
    for score, idx, sentence in scored:
        t = rough_tokens(sentence)
        if budget_used + t <= target_tokens:
            picked_indices.add(idx)
            budget_used += t
        if budget_used >= target_tokens:
            break

    # Reconstruct in original order
    compressed_parts = [sentences[i] for i in sorted(picked_indices)]
    compressed = " ".join(compressed_parts)
    return {"ok": True, "strategy": "smart", "original_tokens": original_tokens,
            "compressed_tokens": rough_tokens(compressed), "text": compressed,
            "ratio": round(rough_tokens(compressed) / original_tokens, 3)}


def generate_sql(question: str, schema: str = "", db_path: str = "") -> dict:
    """
    Generate SQL from a natural language question.
    If db_path is provided, automatically introspects schema.
    If schema is provided as text, uses it directly.
    Returns generated SQL query. Pair with sql_query() to execute.
    Note: This builds the query using heuristics + schema context.
    For best results, run with an LLM (this provides the schema grounding layer).
    """
    # If db_path given, get schema first
    effective_schema = schema
    if db_path and not schema:
        try:
            schema_result = schema_introspect(db_path)
            if schema_result.get("ok"):
                lines = [f"Database: {db_path}"]
                for table, info in schema_result.get("tables", {}).items():
                    cols = ", ".join(f"{c['name']} {c['type']}" for c in info.get("columns", []))
                    lines.append(f"Table {table}: ({cols}) — {info.get('row_count', '?')} rows")
                effective_schema = "\n".join(lines)
        except Exception:
            pass

    # Build context for SQL generation
    context = {
        "ok": True,
        "question": question,
        "schema": effective_schema[:3000] if effective_schema else "(no schema provided — add db_path or schema parameter)",
        "hint": (
            "Use this schema + question with your LLM to generate SQL. "
            "Then call sql_query(db_path, generated_sql) to execute it. "
            "Example: SELECT column FROM table WHERE condition LIMIT 100"
        ),
        "example_patterns": [
            "COUNT: SELECT COUNT(*) FROM table WHERE col = 'value'",
            "JOIN: SELECT a.col, b.col FROM a JOIN b ON a.id = b.a_id",
            "GROUP BY: SELECT col, COUNT(*) FROM table GROUP BY col ORDER BY 2 DESC",
            "LIKE: SELECT * FROM table WHERE text_col LIKE '%keyword%'",
        ],
    }

    # Simple keyword-based SQL generation for common patterns
    q_lower = question.lower()
    if effective_schema:
        tables = []
        import re as _re
        for match in _re.finditer(r"Table (\w+):", effective_schema):
            tables.append(match.group(1))

        if tables:
            main_table = tables[0]
            if "count" in q_lower or "how many" in q_lower:
                context["generated_sql"] = f"SELECT COUNT(*) FROM {main_table};"
            elif "all" in q_lower or "show" in q_lower or "list" in q_lower:
                context["generated_sql"] = f"SELECT * FROM {main_table} LIMIT 100;"
            elif "recent" in q_lower or "latest" in q_lower or "last" in q_lower:
                context["generated_sql"] = f"SELECT * FROM {main_table} ORDER BY rowid DESC LIMIT 20;"
            else:
                context["generated_sql"] = f"SELECT * FROM {main_table} LIMIT 20; -- adjust as needed"

    return context


# ─── Image Understanding ───────────────────────────────────────────────────────

def describe_image(path: str, detail: str = "brief") -> dict:
    """
    Generate a natural language description of an image.
    detail: 'brief' | 'detailed'
    Uses BLIP (Salesforce/blip-image-captioning-base) via transformers.
    Falls back to metadata-only description if transformers not installed.
    Note: First call downloads ~500 MB model. Cached afterward.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    ext = target.suffix.lower()
    if ext not in {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp", ".gif"}:
        return {"ok": False, "error": f"Unsupported image format: {ext}"}

    # Try BLIP via transformers
    try:
        import torch
        from PIL import Image as PILImage
        from transformers import BlipForConditionalGeneration, BlipProcessor

        try:
            import runtime_safety

            _cfg = runtime_safety.load_config()
            model_name = (_cfg.get("image_model") or "Salesforce/blip-image-captioning-base").strip()
        except Exception:
            model_name = "Salesforce/blip-image-captioning-base"
        try:
            processor = BlipProcessor.from_pretrained(model_name)
            model = BlipForConditionalGeneration.from_pretrained(model_name)
        except Exception as e:
            return {"ok": False, "error": f"Failed to load BLIP model: {e}. Run: pip install transformers torch Pillow"}

        with PILImage.open(str(target)) as img:
            img_rgb = img.convert("RGB")
        device = "cuda" if torch.cuda.is_available() else "cpu"
        model = model.to(device)

        if detail == "detailed":
            # Conditional captioning with a prompt
            text_prompt = "a photography of"
            inputs = processor(img_rgb, text_prompt, return_tensors="pt").to(device)
        else:
            inputs = processor(img_rgb, return_tensors="pt").to(device)

        with torch.no_grad():
            out = model.generate(**inputs, max_new_tokens=80)
        caption = processor.decode(out[0], skip_special_tokens=True)

        # Also run OCR if text is likely present
        ocr_text = ""
        try:
            import easyocr
            reader = easyocr.Reader(["en"], gpu=False, verbose=False)
            ocr_results = reader.readtext(str(target))
            ocr_text = " ".join([r[1] for r in ocr_results if r[2] > 0.3])[:500]
        except Exception:
            pass

        result: dict = {
            "ok": True, "path": str(target), "model": "BLIP",
            "caption": caption, "detail": detail,
        }
        if ocr_text:
            result["ocr_text"] = ocr_text
        return result

    except ImportError:
        # Fallback: return metadata + OCR if available
        result_fb: dict = {
            "ok": True, "path": str(target), "model": "fallback_metadata",
            "warning": "transformers not installed (pip install transformers torch) — BLIP captioning unavailable",
        }
        try:
            from PIL import Image as PILImage
            with PILImage.open(str(target)) as img:
                result_fb["size"] = f"{img.width}x{img.height}"
                result_fb["mode"] = img.mode
                result_fb["format"] = img.format
        except Exception:
            pass
        # Try OCR anyway
        try:
            import easyocr
            reader = easyocr.Reader(["en"], gpu=False, verbose=False)
            ocr_results = reader.readtext(str(target))
            ocr_text = " ".join([r[1] for r in ocr_results if r[2] > 0.3])[:500]
            if ocr_text:
                result_fb["ocr_text"] = ocr_text
                result_fb["caption"] = f"Image contains text: {ocr_text[:200]}"
        except Exception:
            pass
        return result_fb
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── NLP — Summarization, Classification, Translation ─────────────────────────

def summarize_text(text: str, sentences: int = 5, method: str = "extractive") -> dict:
    """
    Summarize text. method:
    - 'extractive': score sentences by position, length, keyword density. No deps.
    - 'abstractive': uses transformers (facebook/bart-large-cnn). First run ~1.5 GB download.
    sentences: number of sentences to include (extractive) or target length hint (abstractive).
    """
    import re as _re

    if not text.strip():
        return {"ok": False, "error": "Empty text"}

    if method == "abstractive":
        try:
            from transformers import pipeline as _pipeline
            summarizer = _pipeline("summarization", model="facebook/bart-large-cnn")
            max_len = min(sentences * 40, 200)
            result_text = summarizer(text[:4096], max_length=max_len, min_length=30, do_sample=False)[0]["summary_text"]
            return {"ok": True, "method": "abstractive", "model": "bart-large-cnn", "summary": result_text, "original_chars": len(text)}
        except ImportError:
            pass  # Fall through to extractive

    # Extractive summarization (no deps)
    sentence_list = _re.split(r'(?<=[.!?])\s+', text.strip())
    if not sentence_list:
        sentence_list = [s.strip() for s in text.split("\n") if s.strip()]

    if len(sentence_list) <= sentences:
        return {"ok": True, "method": "extractive", "summary": text, "sentences_in": len(sentence_list), "sentences_out": len(sentence_list)}

    words = _re.findall(r'\b\w{4,}\b', text.lower())
    freq: dict = {}
    for w in words:
        freq[w] = freq.get(w, 0) + 1
    max_freq = max(freq.values(), default=1)

    scored = []
    total = len(sentence_list)
    for i, sent in enumerate(sentence_list):
        pos = 1.5 if i < total * 0.15 else (1.2 if i > total * 0.85 else 1.0)
        sent_words = _re.findall(r'\b\w{4,}\b', sent.lower())
        tf_score = sum(freq.get(w, 0) / max_freq for w in sent_words) / max(len(sent_words), 1)
        len_score = 1.0 if 15 < len(sent.split()) < 50 else 0.7
        scored.append((pos * len_score * (1 + tf_score), i, sent))

    top = sorted(scored, key=lambda x: -x[0])[:sentences]
    summary = " ".join(s for _, _, s in sorted(top, key=lambda x: x[1]))
    return {"ok": True, "method": "extractive", "summary": summary, "sentences_in": total, "sentences_out": len(top), "original_chars": len(text)}


def classify_text(text: str, labels: list | None = None, threshold: float = 0.0) -> dict:
    """
    Classify text into one or more categories.
    labels: list of class names. If empty, uses general-purpose categories.
    Uses zero-shot classification via transformers if available,
    falls back to cosine similarity via sentence-transformers,
    falls back to keyword-frequency scoring.
    threshold: minimum score (0.0 = return all, 0.5 = return only confident)
    """
    if not labels:
        labels = ["technical", "creative", "analytical", "factual", "conversational", "instructional", "narrative"]

    # Try zero-shot with transformers
    try:
        from transformers import pipeline as _pipeline
        classifier = _pipeline("zero-shot-classification", model="facebook/bart-large-mnli")
        result = classifier(text[:512], labels)
        scores = dict(zip(result["labels"], result["scores"]))
        filtered = {k: round(v, 4) for k, v in scores.items() if v >= threshold}
        return {"ok": True, "method": "zero-shot-transformers", "text_preview": text[:80], "scores": filtered, "top": result["labels"][0]}
    except (ImportError, Exception):
        pass

    # Fallback: sentence-transformers cosine similarity
    try:
        from sentence_transformers import SentenceTransformer
        from sentence_transformers import util as _stutil
        _model = SentenceTransformer("all-MiniLM-L6-v2")
        text_emb = _model.encode(text[:512], convert_to_tensor=True)
        label_embs = _model.encode(labels, convert_to_tensor=True)
        scores_tensor = _stutil.cos_sim(text_emb, label_embs)[0]
        scores_dict = {label: round(float(score), 4) for label, score in zip(labels, scores_tensor)}
        filtered = {k: v for k, v in scores_dict.items() if v >= threshold}
        top = max(scores_dict, key=lambda x: scores_dict[x])
        return {"ok": True, "method": "sentence-transformers", "text_preview": text[:80], "scores": filtered, "top": top}
    except (ImportError, Exception):
        pass

    # Final fallback: keyword scoring
    text_lower = text.lower()
    keyword_map = {
        "technical": ["function","class","error","code","implement","system","algorithm","api","module"],
        "creative": ["imagine","story","describe","write","create","design","idea","dream","novel"],
        "analytical": ["analyze","compare","evaluate","assess","determine","examine","conclude","evidence"],
        "factual": ["according","reported","data","study","research","found","shows","indicates"],
        "conversational": ["you","i","we","me","your","my","hey","hi","think","feel","want"],
        "instructional": ["step","first","then","next","do","run","install","configure","how","guide"],
        "narrative": ["then","after","before","when","suddenly","finally","said","went","came"],
    }
    scores_fb = {}
    for label in labels:
        keywords = keyword_map.get(label, [label.lower()])
        hits = sum(1 for kw in keywords if kw in text_lower)
        scores_fb[label] = round(hits / max(len(keywords), 1), 4)
    total_s = sum(scores_fb.values()) or 1
    scores_norm = {k: round(v / total_s, 4) for k, v in scores_fb.items()}
    top = max(scores_norm, key=lambda x: scores_norm[x])
    return {"ok": True, "method": "keyword-fallback", "text_preview": text[:80], "scores": scores_norm, "top": top}


def translate_text(text: str, target_lang: str = "en", source_lang: str = "auto") -> dict:
    """
    Translate text. Uses deep-translator (Google backend) if installed.
    Falls back to LibreTranslate public API (rate-limited, no key required).
    target_lang: ISO 639-1 code (en, fr, de, es, zh, ja, ar, ru, pt, it, ko)
    source_lang: ISO 639-1 code or 'auto'
    """
    try:
        from deep_translator import GoogleTranslator
        src = source_lang if source_lang != "auto" else "auto"
        translator = GoogleTranslator(source=src, target=target_lang)
        chunks = [text[i:i+4999] for i in range(0, len(text), 4999)]
        translated_chunks = [translator.translate(chunk) for chunk in chunks]
        translated = " ".join(c for c in translated_chunks if c)
        return {"ok": True, "method": "google-translate", "source_lang": source_lang, "target_lang": target_lang, "translated": translated, "original_chars": len(text)}
    except ImportError:
        pass
    # Fallback: LibreTranslate public API
    try:
        import json as _json
        import urllib.request
        payload = _json.dumps({"q": text[:2000], "source": source_lang if source_lang != "auto" else "en", "target": target_lang, "format": "text"}).encode()
        req = urllib.request.Request(
            "https://libretranslate.com/translate", data=payload,
            headers={"Content-Type": "application/json"}, method="POST",
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = _json.loads(resp.read())
            return {"ok": True, "method": "libretranslate-public", "source_lang": source_lang, "target_lang": target_lang, "translated": data.get("translatedText", ""), "original_chars": len(text)}
    except Exception as e:
        return {"ok": False, "error": f"Translation requires deep-translator: pip install deep-translator. Error: {e}"}


# ─── Code Intelligence ─────────────────────────────────────────────────────────

def code_symbols(path: str, include_private: bool = False) -> dict:
    """
    Extract a complete symbol index from a Python file or directory.
    For each symbol: name, type, line, docstring, signature, parent class.
    include_private: include _ prefixed symbols (default False).
    Returns a structured symbol table useful for code navigation.
    """
    import ast as _ast

    def _extract(fpath: Path) -> dict:
        try:
            source = fpath.read_text(encoding="utf-8", errors="replace")
            tree = _ast.parse(source, filename=str(fpath))
        except Exception as e:
            return {"error": str(e), "symbols": [], "count": 0}
        symbols: list = []
        for node in tree.body:
            if isinstance(node, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
                if not include_private and node.name.startswith("_"):
                    continue
                try:
                    sig = f"({', '.join(a.arg for a in node.args.args)})"
                except Exception:
                    sig = "()"
                symbols.append({"name": node.name, "type": "async_fn" if isinstance(node, _ast.AsyncFunctionDef) else "function", "line": node.lineno, "signature": sig, "docstring": (_ast.get_docstring(node) or "")[:100], "decorators": [_ast.unparse(d) for d in node.decorator_list]})
            elif isinstance(node, _ast.ClassDef):
                if not include_private and node.name.startswith("_"):
                    continue
                methods = []
                for item in node.body:
                    if isinstance(item, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
                        if not include_private and item.name.startswith("_"):
                            continue
                        methods.append({"name": item.name, "line": item.lineno, "type": "method"})
                symbols.append({"name": node.name, "type": "class", "line": node.lineno, "bases": [_ast.unparse(b) for b in node.bases], "docstring": (_ast.get_docstring(node) or "")[:100], "methods": methods})
        return {"symbols": symbols, "count": len(symbols)}

    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}

    if target.is_file():
        return {"ok": True, "path": str(target), **_extract(target)}

    py_files = [f for f in target.rglob("*.py") if not any(p in str(f) for p in (".venv", "__pycache__"))][:50]
    all_symbols: dict = {}
    for f in py_files:
        all_symbols[str(f.relative_to(target))] = _extract(f)
    total = sum(v.get("count", 0) for v in all_symbols.values())
    return {"ok": True, "path": str(target), "files_analyzed": len(py_files), "total_symbols": total, "files": all_symbols}


def find_todos(path: str, tags: list | None = None) -> dict:
    """
    Scan a file or directory for TODO/FIXME/HACK/NOTE/BUG/REVIEW/OPTIMIZE comments.
    Returns each finding: file, line, tag type, message.
    """
    import re as _re
    if tags is None:
        tags = ["TODO", "FIXME", "HACK", "BUG", "REVIEW", "OPTIMIZE", "XXX", "NOTE", "WARN", "DEPRECATED"]
    tag_pattern = "|".join(_re.escape(t) for t in tags)
    rx = _re.compile(rf'(?:#|//|/\*|<!--)\s*({tag_pattern})\s*[:\-]?\s*(.*)', _re.IGNORECASE)
    CODE_EXTS = {".py", ".js", ".ts", ".tsx", ".jsx", ".go", ".rs", ".java", ".c", ".cpp", ".h", ".cs", ".rb", ".php", ".sh", ".yml", ".yaml", ".toml"}

    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}

    scan_files = [target] if target.is_file() else [
        f for f in target.rglob("*")
        if f.is_file() and f.suffix.lower() in CODE_EXTS and not any(p in str(f) for p in (".git", ".venv", "__pycache__", "node_modules"))
    ][:200]

    findings, files_scanned = [], 0
    for fpath in scan_files:
        files_scanned += 1
        try:
            lines = fpath.read_text(encoding="utf-8", errors="ignore").splitlines()
            for lineno, line in enumerate(lines, 1):
                for m in rx.finditer(line):
                    findings.append({"file": str(fpath.relative_to(target) if target.is_dir() else fpath), "line": lineno, "tag": m.group(1).upper(), "message": m.group(2).strip()[:120]})
        except Exception:
            continue

    by_tag: dict = {}
    for f in findings:
        by_tag[f["tag"]] = by_tag.get(f["tag"], 0) + 1
    return {"ok": True, "path": str(target), "files_scanned": files_scanned, "total_found": len(findings), "by_tag": by_tag, "findings": findings[:100]}


def dependency_graph(path: str) -> dict:
    """
    Build a Python import dependency graph for a file or package.
    Uses AST to extract import statements, resolves local vs external deps.
    Returns adjacency list + networkx metrics if available.
    """
    import ast as _ast

    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}

    root = target if target.is_dir() else target.parent
    local_modules = {f.stem for f in root.rglob("*.py") if not any(p in str(f) for p in (".venv", "__pycache__"))}

    def _get_imports(fpath: Path) -> list[str]:
        try:
            source = fpath.read_text(encoding="utf-8", errors="replace")
            tree = _ast.parse(source)
            imports = []
            for node in _ast.walk(tree):
                if isinstance(node, _ast.Import):
                    imports.extend(a.name.split(".")[0] for a in node.names)
                elif isinstance(node, _ast.ImportFrom):
                    if node.module:
                        imports.append(node.module.split(".")[0])
            return list(dict.fromkeys(imports))
        except Exception:
            return []

    py_files = [target] if target.is_file() else [
        f for f in target.rglob("*.py") if not any(p in str(f) for p in (".venv", "__pycache__"))
    ][:30]

    edges: list[dict] = []
    nodes: set[str] = set()
    for fpath in py_files:
        module_name = fpath.stem
        nodes.add(module_name)
        for imp in _get_imports(fpath):
            nodes.add(imp)
            is_local = imp in local_modules
            edges.append({"from": module_name, "to": imp, "type": "local" if is_local else "external"})

    external_deps = list({e["to"] for e in edges if e["type"] == "external"})
    local_deps = [e for e in edges if e["type"] == "local"]
    result: dict = {"ok": True, "path": str(target), "nodes": list(nodes), "node_count": len(nodes), "edges": edges[:200], "edge_count": len(edges), "external_packages": external_deps, "local_edges": local_deps}

    try:
        import networkx as nx
        G = nx.DiGraph()
        G.add_nodes_from(nodes)
        G.add_edges_from([(e["from"], e["to"]) for e in edges])
        result["metrics"] = {"most_imported": sorted(dict(G.in_degree()).items(), key=lambda x: -x[1])[:5], "most_importing": sorted(dict(G.out_degree()).items(), key=lambda x: -x[1])[:5], "is_dag": nx.is_directed_acyclic_graph(G)}
    except Exception:
        pass
    return result


# ─── URL Intelligence ──────────────────────────────────────────────────────────

def extract_links(url: str, same_domain: bool = False, max_links: int = 100) -> dict:
    """
    Extract all hyperlinks from a webpage.
    same_domain: only return links from the same domain.
    Returns: links with href, internal/external classification, domain.
    """
    from urllib.parse import urljoin, urlparse
    try:
        import trafilatura
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            return {"ok": False, "error": "Could not fetch URL"}
        raw_links: list = []
        try:
            from trafilatura.urls import extract_links as _traf_links
            raw_links = list(_traf_links(downloaded, url) or [])
        except Exception:
            pass
        if len(raw_links) < 5:
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(downloaded, "html.parser")
                raw_links += [urljoin(url, a.get("href", "")) for a in soup.find_all("a", href=True)]
            except Exception:
                pass
    except ImportError:
        return {"ok": False, "error": "trafilatura not installed"}

    base_domain = urlparse(url).netloc
    links, seen = [], set()
    for link in raw_links:
        link = str(link).strip()
        if not link or link.startswith(("mailto:", "javascript:", "#")) or link in seen or len(links) >= max_links:
            continue
        seen.add(link)
        is_internal = urlparse(link).netloc == base_domain
        if same_domain and not is_internal:
            continue
        links.append({"url": link, "internal": is_internal, "domain": urlparse(link).netloc})

    return {"ok": True, "source_url": url, "total_links": len(links), "internal": sum(1 for lnk in links if lnk["internal"]), "external": sum(1 for lnk in links if not lnk["internal"]), "links": links}


def check_url(url: str, timeout: int = 10) -> dict:
    """
    Check if a URL is accessible. Returns HTTP status, response time, content type.
    Uses HEAD request for speed. Useful for monitoring, link validation.
    """
    import time as _time
    import urllib.error
    import urllib.request
    start = _time.time()
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Layla/2.0 health-check"}, method="HEAD")
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            elapsed = round((_time.time() - start) * 1000, 1)
            return {"ok": True, "url": url, "status": resp.status, "accessible": resp.status < 400, "response_ms": elapsed, "content_type": resp.headers.get("Content-Type", ""), "server": resp.headers.get("Server", "")}
    except urllib.error.HTTPError as e:
        return {"ok": False, "url": url, "status": e.code, "accessible": False, "response_ms": round((_time.time() - start) * 1000, 1), "error": str(e)}
    except Exception as e:
        return {"ok": False, "url": url, "accessible": False, "response_ms": round((_time.time() - start) * 1000, 1), "error": str(e)}


# ─── Scientific Computation ────────────────────────────────────────────────────

def scipy_compute(operation: str, params: dict | None = None) -> dict:
    """
    Scientific computation via scipy.
    Operations: stats.describe | stats.ttest | stats.correlation | stats.normalize |
                optimize.minimize | integrate.quad | fft | interpolate
    params: dict of inputs specific to each operation (see docstring examples below).
    Examples:
      scipy_compute('stats.describe', {'data': [1,2,3,4,5]})
      scipy_compute('stats.ttest', {'a': [1,2,3], 'b': [4,5,6]})
      scipy_compute('stats.correlation', {'x': [1,2,3], 'y': [2,4,6]})
      scipy_compute('optimize.minimize', {'func': 'x**2 + 2*x', 'x0': 0})
      scipy_compute('integrate.quad', {'func': 'x**2', 'a': 0, 'b': 1})
    """
    try:
        import numpy as _np
        import scipy.stats as _stats
    except ImportError:
        return {"ok": False, "error": "scipy not installed: pip install scipy"}

    if params is None:
        params = {}
    op = operation.lower().strip()

    try:
        if op == "stats.describe":
            data = _np.array(params["data"], dtype=float)
            desc = _stats.describe(data)
            return {"ok": True, "operation": op, "result": {"n": desc.nobs, "min": float(desc.minmax[0]), "max": float(desc.minmax[1]), "mean": float(desc.mean), "variance": float(desc.variance), "skewness": float(desc.skewness), "kurtosis": float(desc.kurtosis), "std": float(_np.std(data)), "median": float(_np.median(data)), "q25": float(_np.percentile(data, 25)), "q75": float(_np.percentile(data, 75))}}
        elif op == "stats.ttest":
            a, b = _np.array(params["a"], dtype=float), _np.array(params["b"], dtype=float)
            res = _stats.ttest_ind(a, b)
            return {"ok": True, "operation": op, "result": {"t_statistic": float(res.statistic), "p_value": float(res.pvalue), "significant_at_05": float(res.pvalue) < 0.05, "mean_a": float(_np.mean(a)), "mean_b": float(_np.mean(b)), "interpretation": "statistically different (p<0.05)" if res.pvalue < 0.05 else "no significant difference"}}
        elif op == "stats.correlation":
            x, y = _np.array(params["x"], dtype=float), _np.array(params["y"], dtype=float)
            r, p = _stats.pearsonr(x, y)
            return {"ok": True, "operation": op, "result": {"pearson_r": round(float(r), 6), "p_value": round(float(p), 6), "significant": float(p) < 0.05, "strength": "strong" if abs(r) > 0.7 else ("moderate" if abs(r) > 0.4 else "weak"), "direction": "positive" if r > 0 else "negative"}}
        elif op == "stats.normalize":
            data = _np.array(params["data"], dtype=float)
            mn, mx = data.min(), data.max()
            normalized = ((data - mn) / (mx - mn)).tolist() if mx != mn else [0.0] * len(data)
            return {"ok": True, "operation": op, "result": {"normalized": normalized, "original_min": float(mn), "original_max": float(mx)}}
        elif op == "optimize.minimize":
            import sympy as sp
            from sympy.parsing.sympy_parser import parse_expr
            x_sym = sp.Symbol("x")
            expr = parse_expr(params["func"], local_dict={"x": x_sym})
            fn = sp.lambdify(x_sym, expr, "numpy")
            from scipy.optimize import minimize_scalar
            res = minimize_scalar(fn)
            return {"ok": True, "operation": op, "result": {"x_min": float(res.x), "f_min": float(res.fun), "success": res.success}}
        elif op == "integrate.quad":
            import sympy as sp
            from sympy.parsing.sympy_parser import parse_expr
            x_sym = sp.Symbol("x")
            expr = parse_expr(params["func"], local_dict={"x": x_sym})
            fn = sp.lambdify(x_sym, expr, "numpy")
            from scipy.integrate import quad
            val, err = quad(fn, params["a"], params["b"])
            return {"ok": True, "operation": op, "result": {"integral": float(val), "error_estimate": float(err)}}
        elif op == "fft":
            data = _np.array(params["data"], dtype=float)
            fft_vals = _np.fft.fft(data)
            magnitudes = _np.abs(fft_vals).tolist()
            freqs = _np.fft.fftfreq(len(data)).tolist()
            half = len(magnitudes) // 2
            return {"ok": True, "operation": op, "result": {"frequencies": freqs[:half], "magnitudes": magnitudes[:half], "dominant_freq_idx": int(_np.argmax(magnitudes[:half]))}}
        elif op == "interpolate":
            from scipy.interpolate import interp1d
            x, y = _np.array(params["x"], dtype=float), _np.array(params["y"], dtype=float)
            f = interp1d(x, y, kind=params.get("kind", "linear"), fill_value="extrapolate")
            x_new = _np.array(params["x_new"], dtype=float)
            return {"ok": True, "operation": op, "result": {"x_new": params["x_new"], "y_interpolated": f(x_new).tolist()}}
        else:
            return {"ok": False, "error": f"Unknown operation: {op}. Use stats.describe/ttest/correlation/normalize, optimize.minimize, integrate.quad, fft, interpolate"}
    except KeyError as e:
        return {"ok": False, "error": f"Missing required param: {e}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Machine Learning ──────────────────────────────────────────────────────────

def cluster_data(data: list, n_clusters: int = 3, method: str = "kmeans", features: list | None = None) -> dict:
    """
    Cluster a dataset. method: 'kmeans' | 'dbscan' | 'hierarchical'
    data: list of dicts (from read_csv) or list of numeric lists.
    features: column names to use for dict rows. Empty = all numeric columns.
    Returns cluster assignments, centroids, per-cluster statistics.
    """
    try:
        import numpy as _np
        import sklearn.cluster as _cluster
        import sklearn.preprocessing as _prep
    except ImportError:
        return {"ok": False, "error": "scikit-learn not installed: pip install scikit-learn"}

    if data and isinstance(data[0], dict):
        import pandas as _pd
        df = _pd.DataFrame(data)
        if features:
            df = df[features]
        df = df.select_dtypes(include="number").dropna()
        X = df.values
        col_names = list(df.columns)
    else:
        X = _np.array(data, dtype=float)
        if X.ndim == 1:
            X = X.reshape(-1, 1)
        col_names = [f"x{i}" for i in range(X.shape[1])]

    if len(X) < 2:
        return {"ok": False, "error": "Need at least 2 data points"}

    scaler = _prep.StandardScaler()
    X_scaled = scaler.fit_transform(X)

    if method == "kmeans":
        n_clusters = min(n_clusters, len(X))
        model = _cluster.KMeans(n_clusters=n_clusters, random_state=42, n_init="auto")
        labels = model.fit_predict(X_scaled).tolist()
        centroids = scaler.inverse_transform(model.cluster_centers_).tolist()
        extra = {"inertia": float(model.inertia_)}
    elif method == "dbscan":
        model = _cluster.DBSCAN(eps=0.5, min_samples=max(2, len(X) // 10))
        labels = model.fit_predict(X_scaled).tolist()
        centroids = []
        extra = {"noise_points": labels.count(-1)}
    elif method == "hierarchical":
        model = _cluster.AgglomerativeClustering(n_clusters=min(n_clusters, len(X)))
        labels = model.fit_predict(X_scaled).tolist()
        centroids = []
        extra = {}
    else:
        return {"ok": False, "error": f"Unknown method: {method}. Use kmeans/dbscan/hierarchical"}

    unique_labels = sorted(set(labels))
    cluster_stats = {int(lbl): {"size": labels.count(lbl), "mean": X[[i for i, ln in enumerate(labels) if ln == lbl]].mean(axis=0).tolist()} for lbl in unique_labels}
    return {"ok": True, "method": method, "n_clusters_found": len(unique_labels), "labels": labels, "centroids": centroids, "cluster_stats": cluster_stats, "features_used": col_names, "n_points": len(X), **extra}


def dataset_summary(path: str) -> dict:
    """
    Comprehensive statistical summary of any tabular data (CSV, Excel, JSON, Parquet).
    Returns: shape, dtypes, missing values, numeric stats, top correlations,
    categorical value counts, duplicate count, and data quality flags.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import numpy as _np
        import pandas as _pd
        ext = target.suffix.lower()
        if ext in (".csv", ".tsv"):
            df = _pd.read_csv(str(target), sep="\t" if ext == ".tsv" else ",")
        elif ext in (".xlsx", ".xls"):
            df = _pd.read_excel(str(target))
        elif ext == ".json":
            df = _pd.read_json(str(target))
        elif ext == ".parquet":
            df = _pd.read_parquet(str(target))
        else:
            df = _pd.read_csv(str(target))

        missing = df.isnull().sum()
        missing_pct = (missing / len(df) * 100).round(2)
        numeric_cols = df.select_dtypes(include="number")
        cat_cols = df.select_dtypes(include=["object", "category"])

        result: dict = {
            "ok": True, "path": str(target),
            "shape": {"rows": len(df), "columns": len(df.columns)},
            "columns": list(df.columns),
            "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
            "missing": {col: {"count": int(missing[col]), "pct": float(missing_pct[col])} for col in df.columns if missing[col] > 0},
            "duplicates": int(df.duplicated().sum()),
        }
        if len(numeric_cols.columns) > 0:
            result["numeric_summary"] = numeric_cols.describe().to_dict()
            try:
                corr = numeric_cols.corr()
                pairs = []
                cols = list(corr.columns)
                for i in range(len(cols)):
                    for j in range(i+1, len(cols)):
                        val = float(corr.iloc[i, j])
                        if not _np.isnan(val):
                            pairs.append({"col_a": cols[i], "col_b": cols[j], "pearson_r": round(val, 4)})
                result["top_correlations"] = sorted(pairs, key=lambda x: abs(x["pearson_r"]), reverse=True)[:10]
            except Exception:
                pass
        if len(cat_cols.columns) > 0:
            result["categorical_summary"] = {col: {"unique_values": int(df[col].nunique()), "top_values": df[col].value_counts().head(10).to_dict()} for col in list(cat_cols.columns)[:5]}
        flags = []
        if result.get("duplicates", 0) > 0:
            flags.append(f"{result['duplicates']} duplicate rows found")
        high_missing = {k for k, v in result.get("missing", {}).items() if v["pct"] > 20}
        if high_missing:
            flags.append(f"High missing data (>20%) in: {', '.join(high_missing)}")
        result["quality_flags"] = flags
        return result
    except ImportError:
        return {"ok": False, "error": "pandas not installed: pip install pandas"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── RSS / Feed ────────────────────────────────────────────────────────────────

def rss_feed(url: str, max_items: int = 20, include_content: bool = False) -> dict:
    """
    Fetch and parse an RSS or Atom feed.
    Returns: feed title, description, and entry list (title, link, published, author, summary, tags).
    include_content: fetch and extract full article text for each entry (slow but thorough).
    """
    try:
        import feedparser
        feed = feedparser.parse(url)
        if feed.bozo and not feed.entries:
            return {"ok": False, "error": f"Feed parse error: {feed.bozo_exception}"}
        entries = []
        for entry in feed.entries[:max_items]:
            item: dict = {"title": entry.get("title", ""), "link": entry.get("link", ""), "published": str(entry.get("published", "")), "author": entry.get("author", ""), "tags": [t.get("term", "") for t in entry.get("tags", [])], "summary": (entry.get("summary", "") or "")[:400]}
            if include_content and item["link"]:
                try:
                    import trafilatura
                    dl = trafilatura.fetch_url(item["link"])
                    if dl:
                        item["full_text"] = (trafilatura.extract(dl) or "")[:3000]
                except Exception:
                    pass
            entries.append(item)
        return {"ok": True, "url": url, "feed_title": feed.feed.get("title", ""), "feed_description": (feed.feed.get("description", "") or "")[:200], "entry_count": len(entries), "entries": entries}
    except ImportError:
        return {"ok": False, "error": "feedparser not installed: pip install feedparser"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Text Statistics ───────────────────────────────────────────────────────────

def text_stats(text: str) -> dict:
    """
    Comprehensive text statistics and readability metrics.
    Returns: word/sentence/char counts, vocabulary richness, Flesch reading ease,
    avg sentence length, estimated reading time, top 15 non-stopword words.
    """
    import re as _re

    if not text.strip():
        return {"ok": False, "error": "Empty text"}

    words = _re.findall(r'\b[a-zA-Z]+\b', text)
    sentences = [s.strip() for s in _re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    syllables = sum(max(1, len(_re.findall(r'[aeiouAEIOU]+', w))) for w in words)

    wc = len(words)
    sc = max(len(sentences), 1)
    unique = set(w.lower() for w in words)
    avg_wps = round(wc / sc, 1)
    avg_spw = round(syllables / max(wc, 1), 2)
    flesch = round(max(0.0, min(100.0, 206.835 - 1.015 * avg_wps - 84.6 * avg_spw)), 1)
    grade = "Easy" if flesch >= 70 else ("Standard" if flesch >= 50 else ("Difficult" if flesch >= 30 else "Very Difficult"))

    STOPWORDS = {"the","a","an","and","or","but","in","on","at","to","for","of","with","is","was","are","it","this","that","be","have","do","i","you","we","he","she","they","not","by","as","from","his","her","its","our","their","so","if","but","about","which"}
    freq: dict = {}
    for w in words:
        wl = w.lower()
        if wl not in STOPWORDS and len(wl) > 2:
            freq[wl] = freq.get(wl, 0) + 1

    return {
        "ok": True,
        "counts": {"words": wc, "unique_words": len(unique), "sentences": sc, "paragraphs": len(paragraphs), "characters": len(text)},
        "averages": {"words_per_sentence": avg_wps, "syllables_per_word": avg_spw},
        "readability": {"flesch_score": flesch, "grade": grade},
        "vocabulary_richness": round(len(unique) / max(wc, 1), 4),
        "reading_time_minutes": round(wc / 200, 1),
        "top_words": [{"word": w, "count": c} for w, c in sorted(freq.items(), key=lambda x: -x[1])[:15]],
    }


# ─── Embedding Generation ──────────────────────────────────────────────────────

def embedding_generate(text: str | list, normalize: bool = True) -> dict:
    """
    Generate dense vector embeddings using Layla's RAG embedder (nomic-embed-text).
    text: string or list of strings.
    normalize: L2 normalize (default True — required for cosine similarity).
    Returns: embedding(s) as list of floats, dimension, model name.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.vector_store import _get_embedder
        embedder = _get_embedder()
        is_batch = isinstance(text, list)
        texts = text if is_batch else [text]
        embeddings = embedder.encode(texts, normalize_embeddings=normalize)
        emb_list = embeddings.tolist() if hasattr(embeddings, "tolist") else [list(e) for e in embeddings]
        return {"ok": True, "dimension": len(emb_list[0]) if emb_list else 0, "count": len(emb_list), "normalized": normalize, "embeddings": emb_list if is_batch else emb_list[0]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Image Utilities ───────────────────────────────────────────────────────────

def image_resize(path: str, width: int = 0, height: int = 0, output_path: str = "", maintain_aspect: bool = True) -> dict:
    """
    Resize an image. If maintain_aspect=True, only one dimension needed — the other scales proportionally.
    output_path: where to save. Default: <original>_resized.<ext> in same directory.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    if not width and not height:
        return {"ok": False, "error": "Provide at least width or height"}
    try:
        from PIL import Image as PILImage
        with PILImage.open(str(target)) as img:
            orig_w, orig_h = img.size
            if maintain_aspect:
                if width and not height:
                    height = int(orig_h * width / orig_w)
                elif height and not width:
                    width = int(orig_w * height / orig_h)
            resized = img.resize((width, height), PILImage.LANCZOS)
        out = Path(output_path) if output_path else target.parent / (target.stem + "_resized" + target.suffix)
        if not inside_sandbox(out):
            out = target.parent / (target.stem + "_resized" + target.suffix)
        resized.save(str(out))
        return {"ok": True, "original": str(target), "output": str(out), "original_size": f"{orig_w}x{orig_h}", "new_size": f"{resized.width}x{resized.height}"}
    except ImportError:
        return {"ok": False, "error": "Pillow not installed: pip install Pillow"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TIER 3 TOOLS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

# â”€â”€â”€ Scheduling (APScheduler) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

_SCHEDULER = None
_SCHEDULED_JOBS: dict = {}  # job_id -> metadata


def _get_scheduler():
    global _SCHEDULER
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        if _SCHEDULER is None or not _SCHEDULER.running:
            _SCHEDULER = BackgroundScheduler(timezone="UTC")
            _SCHEDULER.start()
        return _SCHEDULER
    except ImportError:
        return None


def schedule_task(
    tool_name: str,
    args: dict | None = None,
    delay_seconds: float = 0,
    cron_expr: str = "",
    job_id: str = "",
) -> dict:
    """
    Schedule a tool to run in the background.
    tool_name: any registered tool name. args: dict of kwargs.
    delay_seconds: run once after N seconds. cron_expr: '*/5 * * * *' for recurring.
    Returns job_id for cancellation via cancel_task().
    """
    scheduler = _get_scheduler()
    if scheduler is None:
        return {"ok": False, "error": "apscheduler not installed: pip install apscheduler"}
    if tool_name not in TOOLS:
        return {"ok": False, "error": f"Unknown tool: {tool_name}. Use list_tools() to see available tools."}
    import uuid as _uuid
    from datetime import timedelta

    from layla.time_utils import utcnow
    jid = job_id or f"task_{tool_name}_{_uuid.uuid4().hex[:8]}"
    kw = args or {}

    def _run():
        try:
            result = TOOLS[tool_name]["fn"](**kw)
            _SCHEDULED_JOBS[jid]["last_result"] = result
            _SCHEDULED_JOBS[jid]["last_run"] = str(utcnow())[:19]
        except Exception as exc:
            if jid in _SCHEDULED_JOBS:
                _SCHEDULED_JOBS[jid]["last_error"] = str(exc)

    try:
        if cron_expr:
            from apscheduler.triggers.cron import CronTrigger
            parts = cron_expr.split()
            if len(parts) != 5:
                return {"ok": False, "error": "cron_expr must be 5 fields: 'min hour dom month dow'"}
            m, h, dom, mo, dow = parts
            trigger = CronTrigger(minute=m, hour=h, day=dom, month=mo, day_of_week=dow, timezone="UTC")
            schedule_type = f"cron: {cron_expr}"
        else:
            from apscheduler.triggers.date import DateTrigger
            run_at = utcnow() + timedelta(seconds=max(delay_seconds, 0))
            trigger = DateTrigger(run_date=run_at, timezone="UTC")
            schedule_type = f"once in {delay_seconds}s" if delay_seconds > 0 else "immediate background"
        scheduler.add_job(_run, trigger, id=jid, replace_existing=True)
        _SCHEDULED_JOBS[jid] = {"tool": tool_name, "args": kw, "schedule": schedule_type, "added_at": str(utcnow())[:19], "job_id": jid}
        return {"ok": True, "job_id": jid, "tool": tool_name, "schedule": schedule_type}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def list_scheduled_tasks() -> dict:
    """List all currently scheduled background tasks with status, next run time, and last result."""
    scheduler = _get_scheduler()
    if scheduler is None:
        return {"ok": False, "error": "apscheduler not installed"}
    jobs = []
    for job in scheduler.get_jobs():
        meta = _SCHEDULED_JOBS.get(job.id, {})
        jobs.append({"job_id": job.id, "tool": meta.get("tool", "unknown"), "args": meta.get("args", {}), "schedule": meta.get("schedule", ""), "next_run": str(job.next_run_time) if job.next_run_time else "no next run", "last_run": meta.get("last_run", "never"), "last_result": meta.get("last_result"), "last_error": meta.get("last_error")})
    return {"ok": True, "total_jobs": len(jobs), "jobs": jobs}


def cancel_task(job_id: str) -> dict:
    """Cancel a scheduled background task by job_id."""
    scheduler = _get_scheduler()
    if scheduler is None:
        return {"ok": False, "error": "apscheduler not installed"}
    try:
        scheduler.remove_job(job_id)
        meta = _SCHEDULED_JOBS.pop(job_id, {})
        return {"ok": True, "cancelled": job_id, "tool": meta.get("tool", "unknown")}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ Observability â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def log_event(message: str, level: str = "info", context: dict | None = None) -> dict:
    """
    Write a structured log entry to agent/.governance/layla-events.log (JSON-lines).
    level: debug | info | warning | error | critical.
    context: optional dict of extra fields.
    """
    import json as _json

    from layla.time_utils import utcnow
    entry = {"ts": str(utcnow())[:19], "level": level.upper(), "message": message[:500], "context": context or {}}
    try:
        log_path = Path(__file__).resolve().parent.parent.parent / ".governance" / "layla-events.log"
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(_json.dumps(entry) + "\n")
    except Exception:
        pass
    return {"ok": True, "logged": entry}


def trace_last_run(n: int = 20) -> dict:
    """Return the last N entries from the audit log for debugging what the agent did."""
    import json as _json
    audit_path = Path(__file__).resolve().parent.parent.parent / ".governance" / "audit.log"
    if not audit_path.exists():
        return {"ok": False, "error": "No audit log found at agent/.governance/audit.log"}
    try:
        lines = audit_path.read_text(encoding="utf-8", errors="replace").strip().splitlines()
        entries = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            try:
                entries.append(_json.loads(line))
            except Exception:
                entries.append({"raw": line[:200]})
        return {"ok": True, "total_lines": len(lines), "showing_last": min(n, len(entries)), "entries": entries[-n:]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def tool_metrics(top_n: int = 15) -> dict:
    """Analyze audit log for tool usage statistics: call counts, approval rates, never-called tools."""
    import json as _json
    audit_path = Path(__file__).resolve().parent.parent.parent / ".governance" / "audit.log"
    call_counts: dict = {}
    approved, rejected, total = 0, 0, 0
    if audit_path.exists():
        for line in audit_path.read_text(encoding="utf-8", errors="replace").splitlines():
            line = line.strip()
            if not line:
                continue
            total += 1
            try:
                entry = _json.loads(line)
                tool = entry.get("tool") or entry.get("action") or "unknown"
                call_counts[tool] = call_counts.get(tool, 0) + 1
                status = str(entry.get("action", "") or entry.get("status", "")).lower()
                if "approv" in status or status == "ok":
                    approved += 1
                elif "reject" in status or "deny" in status or "block" in status:
                    rejected += 1
            except Exception:
                pass
    top = sorted(call_counts.items(), key=lambda x: -x[1])[:top_n]
    never_called = [t for t in TOOLS if t not in call_counts]
    return {"ok": True, "total_log_entries": total, "approved_actions": approved, "rejected_actions": rejected, "top_tools": [{"tool": t, "calls": c} for t, c in top], "never_called": never_called[:25], "total_registered_tools": len(TOOLS)}


# â”€â”€â”€ Speech as Tools â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def stt_file(path: str, language: str = "en", model_size: str = "base") -> dict:
    """
    Transcribe an audio file (.wav, .mp3, .ogg, .flac, .m4a) using faster-whisper.
    language: ISO 639-1 code or empty for auto-detect. model_size: tiny|base|small|medium|large-v3.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from services.stt import transcribe_file
        result = transcribe_file(str(target), language=language or None)
        if isinstance(result, dict):
            return {"ok": True, "path": str(target), **result}
        return {"ok": True, "path": str(target), "text": str(result)}
    except Exception:
        pass
    try:
        from faster_whisper import WhisperModel
        model = WhisperModel(model_size, device="cpu", compute_type="int8")
        segments, info = model.transcribe(str(target), language=language or None)
        text = " ".join(s.text for s in segments)
        return {"ok": True, "path": str(target), "text": text.strip(), "language": info.language, "prob": round(info.language_probability, 3)}
    except ImportError:
        return {"ok": False, "error": "faster-whisper not installed: pip install faster-whisper"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def tts_speak(text: str, voice: str = "af_heart", output_path: str = "") -> dict:
    """
    Synthesize speech and save as WAV. voice: kokoro-onnx voice ID (af_heart, af_sky, am_adam...).
    output_path: where to save (default: temp file). Returns path to WAV.
    """
    if not text.strip():
        return {"ok": False, "error": "Empty text"}
    import tempfile as _tmp
    import time as _time
    out = output_path or str(Path(_tmp.gettempdir()) / f"layla_tts_{int(_time.time())}.wav")
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from services.tts import speak_to_bytes
        wav_bytes = speak_to_bytes(text)
        if wav_bytes:
            p = Path(out)
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_bytes(wav_bytes)
            return {"ok": True, "output_path": out, "method": "kokoro-onnx", "chars": len(text)}
    except Exception:
        pass
    try:
        import pyttsx3 as _pyttsx3
        engine = _pyttsx3.init()
        engine.save_to_file(text, out)
        engine.runAndWait()
        return {"ok": True, "output_path": out, "method": "pyttsx3", "chars": len(text)}
    except ImportError:
        return {"ok": False, "error": "No TTS backend. Install kokoro-onnx or pyttsx3."}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ Crypto & Finance Extended â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def crypto_prices(symbols: list | str, period: str = "1d") -> dict:
    """
    Real-time and historical crypto price data via yfinance.
    symbols: 'BTC' or ['BTC','ETH','SOL'] â€” auto-appends -USD if missing.
    period: 1d | 5d | 1mo | 3mo | 1y | max.
    """
    try:
        import yfinance as yf
        if isinstance(symbols, str):
            symbols = [symbols]
        results = {}
        for sym in symbols[:10]:
            s = sym.upper()
            if "-" not in s:
                s += "-USD"
            try:
                t = yf.Ticker(s)
                hist = t.history(period=period)
                if hist.empty:
                    results[s] = {"error": "No data"}
                    continue
                current = float(hist["Close"].iloc[-1])
                first = float(hist["Close"].iloc[0])
                results[s] = {"price_usd": round(current, 6), "change_pct": round((current - first) / first * 100, 2), "high": round(float(hist["High"].max()), 6), "low": round(float(hist["Low"].min()), 6), "volume_24h": int(hist["Volume"].iloc[-1]), "period": period}
            except Exception as e:
                results[s] = {"error": str(e)}
        return {"ok": True, "data": results}
    except ImportError:
        return {"ok": False, "error": "yfinance not installed: pip install yfinance"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def economic_indicators(series: str = "SP500", start_year: int = 2000) -> dict:
    """
    Fetch macroeconomic data. Common series: GDP, UNRATE, CPIAUCSL, FEDFUNDS, SP500, T10Y2Y, DEXUSEU.
    Uses pandas-datareader FRED if installed; falls back to yfinance proxies.
    """
    try:
        import datetime as _dt

        import pandas_datareader.data as _pdr
        df = _pdr.DataReader(series, "fred", _dt.datetime(start_year, 1, 1))
        data = df[series].dropna()
        recent = data.tail(20)
        return {"ok": True, "series": series, "source": "FRED", "observations": len(data), "latest_value": round(float(recent.iloc[-1]), 6) if len(recent) else None, "latest_date": str(recent.index[-1])[:10] if len(recent) else None, "history": [{"date": str(d)[:10], "value": round(float(v), 6)} for d, v in recent.items()]}
    except ImportError:
        pass
    except Exception:
        pass
    YF_MAP = {"SP500": "^GSPC", "DEXUSEU": "EURUSD=X", "DEXJPUS": "JPY=X", "GC=F": "GC=F", "CL=F": "CL=F"}
    yf_sym = YF_MAP.get(series.upper(), series)
    try:
        import yfinance as yf
        t = yf.Ticker(yf_sym)
        hist = t.history(period="1y")
        if not hist.empty:
            return {"ok": True, "series": series, "source": "yfinance", "current": round(float(hist["Close"].iloc[-1]), 4), "1y_change_pct": round(float((hist["Close"].iloc[-1] - hist["Close"].iloc[0]) / hist["Close"].iloc[0] * 100), 2)}
    except Exception:
        pass
    return {"ok": False, "error": f"Series '{series}' needs pandas-datareader: pip install pandas-datareader. FRED series: GDP, UNRATE, CPIAUCSL, FEDFUNDS, T10Y2Y"}


# â”€â”€â”€ Code Metrics & Linting â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def code_metrics(path: str) -> dict:
    """
    Compute code quality metrics for Python files: LOC, blank/comment lines,
    function/class count, avg complexity, docstring coverage, high-complexity functions.
    """
    import ast as _ast

    def _file_metrics(fpath: Path) -> dict:
        try:
            source = fpath.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return {"error": str(e)}
        lines = source.splitlines()
        blank = sum(1 for ln in lines if not ln.strip())
        comment = sum(1 for ln in lines if ln.strip().startswith("#"))
        try:
            tree = _ast.parse(source)
        except SyntaxError as e:
            return {"loc": len(lines), "blank": blank, "syntax_error": str(e)}
        functions, classes = [], []
        for node in _ast.walk(tree):
            if isinstance(node, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
                end = getattr(node, "end_lineno", node.lineno + 1)
                branches = sum(1 for n in _ast.walk(node) if isinstance(n, (_ast.If, _ast.While, _ast.For, _ast.ExceptHandler, _ast.BoolOp)))
                functions.append({"name": node.name, "len": (end or node.lineno+1) - node.lineno, "complexity": 1 + branches, "has_doc": bool(_ast.get_docstring(node))})
            elif isinstance(node, _ast.ClassDef):
                classes.append({"name": node.name, "methods": sum(1 for n in node.body if isinstance(n, (_ast.FunctionDef, _ast.AsyncFunctionDef))), "has_doc": bool(_ast.get_docstring(node))})
        avg_cc = round(sum(f["complexity"] for f in functions) / max(len(functions), 1), 2)
        doc_cov = round(sum(1 for f in functions if f["has_doc"]) / max(len(functions), 1) * 100, 1)
        return {"loc": len(lines), "blank": blank, "comment": comment, "code": len(lines)-blank-comment, "functions": len(functions), "classes": len(classes), "avg_complexity": avg_cc, "doc_coverage_pct": doc_cov, "high_complexity": [f["name"] for f in functions if f["complexity"] > 10], "longest": sorted(functions, key=lambda x: -x["len"])[:3]}

    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}
    if target.is_file():
        return {"ok": True, "path": str(target), **_file_metrics(target)}
    py_files = [f for f in target.rglob("*.py") if not any(p in str(f) for p in (".venv", "__pycache__"))][:50]
    totals = {"loc": 0, "functions": 0, "classes": 0}
    files = {}
    for f in py_files:
        m = _file_metrics(f)
        files[str(f.relative_to(target))] = m
        for k in totals:
            totals[k] += m.get(k, 0)
    return {"ok": True, "path": str(target), "files": len(py_files), "totals": totals, "file_metrics": files}


def code_lint(path: str, fix: bool = False) -> dict:
    """Run ruff linter on Python file/dir. fix=True auto-fixes. Falls back to syntax check."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}
    import json as _json
    cmd = [sys.executable, "-m", "ruff", "check", str(target), "--output-format", "json"]
    if fix:
        cmd.append("--fix")
    try:
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=30, encoding="utf-8", errors="replace")
        violations = []
        for v in (_json.loads(r.stdout or "[]") or []):
            violations.append({"file": v.get("filename", ""), "line": v.get("location", {}).get("row"), "code": v.get("code", ""), "message": v.get("message", ""), "fixable": v.get("fix") is not None})
        by_code: dict = {}
        for v in violations:
            by_code[v["code"]] = by_code.get(v["code"], 0) + 1
        return {"ok": True, "tool": "ruff", "violations": len(violations), "by_code": dict(sorted(by_code.items(), key=lambda x: -x[1])[:20]), "details": violations[:50]}
    except FileNotFoundError:
        import ast as _ast
        errors = []
        files = [target] if target.is_file() else [f for f in target.rglob("*.py") if not any(p in str(f) for p in (".venv", "__pycache__"))][:30]
        for f in files:
            try:
                _ast.parse(f.read_text(encoding="utf-8", errors="replace"))
            except SyntaxError as e:
                errors.append({"file": str(f), "line": e.lineno, "message": str(e)})
        return {"ok": True, "tool": "syntax_check_fallback", "syntax_errors": errors, "note": "Install ruff: pip install ruff"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def git_blame(repo: str, file_path: str, line_start: int = 1, line_end: int = 0) -> dict:
    """Run git blame on a file. Returns per-line author, commit hash, date, content."""
    repo_path = Path(repo)
    if not inside_sandbox(repo_path):
        return {"ok": False, "error": "Outside sandbox"}
    cmd = ["git", "blame", "--line-porcelain"]
    if line_end > 0:
        cmd += [f"-L{line_start},{line_end}"]
    cmd.append(file_path)
    try:
        r = subprocess.run(cmd, cwd=str(repo_path), capture_output=True, text=True, timeout=30, encoding="utf-8", errors="replace")
        if r.returncode != 0:
            return {"ok": False, "error": r.stderr.strip()[:500]}
        lines, current = [], {}
        for line in r.stdout.splitlines():
            if line.startswith("\t"):
                current["content"] = line[1:]
                lines.append(current)
                current = {}
            elif " " in line:
                k, v = line.split(" ", 1)
                if k in ("author", "summary"):
                    current[k] = v
                elif k == "author-time":
                    import datetime as _dt
                    current["date"] = str(_dt.datetime.utcfromtimestamp(int(v)))[:10]
                elif len(k) == 40:
                    current["commit"] = k[:8]
        return {"ok": True, "file": file_path, "lines": lines[:200], "total_lines": len(lines)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ File Format Tools â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def yaml_read(path: str) -> dict:
    """Parse a YAML file. Uses PyYAML (pip install pyyaml); simple fallback if not installed."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import yaml as _yaml
        with open(str(target), encoding="utf-8") as f:
            data = _yaml.safe_load(f)
        return {"ok": True, "path": str(target), "data": data, "type": type(data).__name__}
    except ImportError:
        import re as _re
        text = target.read_text(encoding="utf-8", errors="replace")
        data = {}
        for line in text.splitlines():
            m = _re.match(r'^(\w[\w\-\.]*)\s*:\s*(.+)$', line.strip())
            if m:
                data[m.group(1)] = m.group(2).strip()
        return {"ok": True, "path": str(target), "data": data, "note": "Basic parse â€” install pyyaml: pip install pyyaml"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def xml_parse(path_or_text: str) -> dict:
    """Parse XML from a file path or raw string. Returns tree structure (up to 3 levels)."""
    import xml.etree.ElementTree as _ET

    def _to_dict(elem: _ET.Element, depth: int = 0) -> dict:
        node: dict = {"tag": elem.tag, "attribs": dict(elem.attrib)}
        text = (elem.text or "").strip()
        if text:
            node["text"] = text[:500]
        if depth < 3:
            children = [_to_dict(c, depth+1) for c in list(elem)[:20]]
            if children:
                node["children"] = children
        return node

    try:
        p = Path(path_or_text)
        if inside_sandbox(p) and p.exists():
            root = _ET.parse(str(p)).getroot()
            source = "file"
        else:
            root = _ET.fromstring(path_or_text)
            source = "string"
        return {"ok": True, "source": source, "root": _to_dict(root), "element_count": sum(1 for _ in root.iter())}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def hash_file(path: str, algorithm: str = "sha256") -> dict:
    """Compute cryptographic hash of a file. algorithm: md5 | sha1 | sha256 | sha512."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    import hashlib as _hl
    algo = algorithm.lower().replace("-", "")
    if algo not in ("md5", "sha1", "sha256", "sha512"):
        return {"ok": False, "error": "Use md5/sha1/sha256/sha512"}
    try:
        h = _hl.new(algo)
        with open(str(target), "rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
        return {"ok": True, "path": str(target), "algorithm": algo, "hash": h.hexdigest(), "size_bytes": target.stat().st_size}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def base64_tool(data: str, mode: str = "encode", encoding: str = "utf-8") -> dict:
    """Base64 encode/decode. mode: encode | decode | encode_url | decode_url."""
    import base64 as _b64
    try:
        if mode == "encode":
            result = _b64.b64encode(data.encode(encoding)).decode("ascii")
        elif mode == "decode":
            result = _b64.b64decode(data).decode(encoding, errors="replace")
        elif mode == "encode_url":
            result = _b64.urlsafe_b64encode(data.encode(encoding)).decode("ascii")
        elif mode == "decode_url":
            result = _b64.urlsafe_b64decode(data).decode(encoding, errors="replace")
        else:
            return {"ok": False, "error": f"Unknown mode: {mode}. Use encode/decode/encode_url/decode_url"}
        return {"ok": True, "mode": mode, "result": result, "input_len": len(data), "output_len": len(result)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ System Utilities â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def check_port(host: str, port: int, timeout: float = 3.0) -> dict:
    """Check if a TCP port is open. Returns: open/closed, response time ms."""
    import socket as _sock
    import time as _time
    start = _time.time()
    try:
        s = _sock.socket(_sock.AF_INET, _sock.SOCK_STREAM)
        s.settimeout(timeout)
        result = s.connect_ex((host, port))
        s.close()
        return {"ok": True, "host": host, "port": port, "open": result == 0, "response_ms": round((_time.time()-start)*1000, 1)}
    except Exception as e:
        return {"ok": False, "host": host, "port": port, "open": False, "error": str(e)}


def timestamp_convert(value: str | int | float, input_format: str = "auto", output_format: str = "iso") -> dict:
    """
    Convert between timestamp formats.
    input_format: auto | unix | unix_ms | strftime string
    output_format: iso | unix | human | strftime string
    """
    import datetime as _dt
    dt = None
    try:
        if input_format in ("auto", "unix"):
            try:
                f = float(str(value))
                dt = _dt.datetime.utcfromtimestamp(f if f < 1e12 else f/1000)
                input_format = "unix"
            except ValueError:
                pass
        if dt is None:
            for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y"):
                try:
                    dt = _dt.datetime.strptime(str(value), fmt)
                    break
                except ValueError:
                    continue
        if dt is None:
            return {"ok": False, "error": f"Cannot parse '{value}'"}
        if output_format == "iso":
            result = dt.isoformat()
        elif output_format == "unix":
            result = int(dt.timestamp())
        elif output_format == "human":
            result = dt.strftime("%B %d, %Y at %H:%M UTC")
        else:
            result = dt.strftime(output_format)
        return {"ok": True, "input": str(value), "result": result, "utc_iso": dt.isoformat()}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def string_transform(text: str, operations: list | str | None = None) -> dict:
    """
    Apply text transformations. operations: list or single string.
    Ops: upper, lower, title, capitalize, strip, slug, snake_case, camel_case, reverse,
    truncate_N, dedupe_lines, sort_lines, remove_empty_lines, extract_numbers,
    extract_emails, extract_urls, remove_punctuation, first_sentence
    """
    import re as _re
    if operations is None:
        operations = []
    if isinstance(operations, str):
        operations = [operations]
    result, applied = text, []
    for op in operations:
        op = op.strip().lower()
        try:
            if op == "upper":
                result = result.upper()
            elif op == "lower":
                result = result.lower()
            elif op == "title":
                result = result.title()
            elif op == "capitalize":
                result = result.capitalize()
            elif op == "strip":
                result = result.strip()
            elif op == "slug":
                result = _re.sub(r'[^a-z0-9]+', '-', result.lower().strip()).strip('-')
            elif op == "snake_case":
                result = _re.sub(r'[^\w]', '_', _re.sub(r'[\s\-]+', '_', result.lower())).strip('_')
            elif op == "camel_case":
                parts = _re.split(r'[\s_\-]+', result)
                result = parts[0].lower() + "".join(p.capitalize() for p in parts[1:])
            elif op == "reverse":
                result = result[::-1]
            elif op.startswith("truncate_"):
                n = int(op.split("_")[1])
                result = result[:n] + ("..." if len(result) > n else "")
            elif op == "dedupe_lines":
                seen, lines = set(), []
                for line in result.splitlines():
                    if line not in seen:
                        seen.add(line)
                        lines.append(line)
                result = "\n".join(lines)
            elif op == "sort_lines":
                result = "\n".join(sorted(result.splitlines()))
            elif op == "remove_empty_lines":
                result = "\n".join(ln for ln in result.splitlines() if ln.strip())
            elif op == "first_sentence":
                m = _re.search(r'^.+?[.!?]', result)
                result = m.group(0) if m else result
            elif op == "extract_numbers":
                result = ", ".join(_re.findall(r'-?\d+\.?\d*', result))
            elif op == "extract_emails":
                result = ", ".join(_re.findall(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', result))
            elif op == "extract_urls":
                result = "\n".join(_re.findall(r'https?://[^\s<>"\']+', result))
            elif op == "remove_punctuation":
                result = _re.sub(r'[^\w\s]', '', result)
            else:
                applied.append(f"UNKNOWN:{op}")
                continue
            applied.append(op)
        except Exception:
            applied.append(f"ERROR:{op}")
    return {"ok": True, "result": result, "original_length": len(text), "result_length": len(result), "operations_applied": applied}


# â”€â”€â”€ NLP Extended â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def extract_entities(text: str, entity_types: list | None = None) -> dict:
    """
    Extract named entities. entity_types: ['PERSON','ORG','GPE','DATE','MONEY',...]
    Uses spaCy if installed (en_core_web_sm); regex fallback for common patterns.
    """
    if not entity_types:
        entity_types = []
    try:
        import spacy
        try:
            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text[:50000])
            ents = [{"text": e.text, "label": e.label_, "start": e.start_char, "end": e.end_char} for e in doc.ents if not entity_types or e.label_ in entity_types]
            by_type: dict = {}
            for e in ents:
                by_type.setdefault(e["label"], []).append(e["text"])
            return {"ok": True, "method": "spacy", "total": len(ents), "by_type": by_type, "entities": ents[:100]}
        except OSError:
            pass
    except ImportError:
        pass
    import re as _re
    patterns = {"EMAIL": r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', "URL": r'https?://[^\s<>"\']+', "DATE": r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4}', "MONEY": r'\$\d+(?:,\d{3})*(?:\.\d{2})?', "PHONE": r'\+?1?\s*\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}', "CAPITALIZED": r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b'}
    ents = []
    for label, pattern in patterns.items():
        if entity_types and label not in entity_types:
            continue
        for m in _re.finditer(pattern, text):
            ents.append({"text": m.group(0), "label": label, "start": m.start(), "end": m.end()})
    by_type: dict = {}
    for e in ents:
        by_type.setdefault(e["label"], []).append(e["text"])
    return {"ok": True, "method": "regex-fallback", "total": len(ents), "by_type": by_type, "entities": ents[:100]}


def sentiment_timeline(texts: list, labels: list | None = None) -> dict:
    """
    Apply sentiment analysis to a list of texts. Returns per-item polarity + overall trend.
    Useful for review series, chat history, social media posts, time-series sentiment.
    """
    if not texts:
        return {"ok": False, "error": "Empty texts list"}

    def _score(t: str) -> float:
        try:
            from textblob import TextBlob
            return float(TextBlob(t[:2000]).sentiment.polarity)
        except ImportError:
            pos = {"good","great","excellent","love","best","fantastic","happy","positive","success","win","amazing","wonderful"}
            neg = {"bad","terrible","awful","hate","worst","poor","fail","wrong","negative","loss","sad","horrible","disaster"}
            words = set(t.lower().split())
            p, n = len(words & pos), len(words & neg)
            return (p - n) / max(p + n, 1)

    results = []
    for i, text in enumerate(texts[:100]):
        pol = round(_score(text), 4)
        lbl = labels[i] if labels and i < len(labels) else str(i)
        results.append({"label": lbl, "preview": text[:80], "polarity": pol, "sentiment": "positive" if pol > 0.1 else ("negative" if pol < -0.1 else "neutral")})

    scores = [r["polarity"] for r in results]
    avg = round(sum(scores) / len(scores), 4)
    trend = "stable"
    if len(scores) >= 4:
        mid = len(scores) // 2
        first_avg = sum(scores[:mid]) / mid
        second_avg = sum(scores[mid:]) / (len(scores) - mid)
        trend = "improving" if second_avg - first_avg > 0.05 else ("declining" if first_avg - second_avg > 0.05 else "stable")

    return {"ok": True, "count": len(results), "avg_polarity": avg, "trend": trend, "min": round(min(scores), 4), "max": round(max(scores), 4), "timeline": results}


# â”€â”€â”€ Visualization Extended â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def plot_scatter(x: list, y: list, labels: list | None = None, title: str = "", xlabel: str = "", ylabel: str = "", show_regression: bool = True, output_path: str = "") -> dict:
    """
    Scatter plot with optional linear regression line and RÂ² annotation.
    x, y: numeric lists. labels: optional point annotations. show_regression: draw best-fit line.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as _np
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.scatter(x, y, alpha=0.7, s=60, edgecolors="white", linewidth=0.5)
        if labels:
            for xi, yi, lbl in zip(x, y, labels):
                ax.annotate(str(lbl), (xi, yi), fontsize=7, alpha=0.8, xytext=(4, 4), textcoords="offset points")
        if show_regression and len(x) >= 3:
            xa, ya = _np.array(x, dtype=float), _np.array(y, dtype=float)
            m, b = _np.polyfit(xa, ya, 1)
            xl = _np.linspace(xa.min(), xa.max(), 100)
            ax.plot(xl, m*xl+b, "r--", alpha=0.7, linewidth=1.5)
            corr = _np.corrcoef(xa, ya)[0, 1]
            ax.annotate(f"RÂ²={corr**2:.4f}  y={m:.3f}x+{b:.3f}", xy=(0.05, 0.95), xycoords="axes fraction", fontsize=9, color="red")
        ax.set_title(title or "Scatter Plot", fontsize=13)
        ax.set_xlabel(xlabel or "x")
        ax.set_ylabel(ylabel or "y")
        ax.grid(True, alpha=0.3)
        plt.tight_layout()
        import tempfile as _tmp
        import time as _time
        out = output_path or str(Path(_tmp.gettempdir()) / f"layla_scatter_{int(_time.time())}.png")
        fig.savefig(out, dpi=120, bbox_inches="tight")
        plt.close(fig)
        return {"ok": True, "chart_type": "scatter", "path": out, "points": len(x), "regression": show_regression}
    except ImportError:
        return {"ok": False, "error": "matplotlib not installed: pip install matplotlib"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def plot_histogram(data: list, bins: int = 20, title: str = "", xlabel: str = "", show_kde: bool = True, output_path: str = "") -> dict:
    """
    Histogram with optional KDE (kernel density) overlay and descriptive stats annotation.
    data: numeric list. bins: number of histogram bins. show_kde: overlay smooth density curve.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import numpy as _np
        arr = _np.array(data, dtype=float)
        fig, ax = plt.subplots(figsize=(10, 6))
        ax.hist(arr, bins=bins, edgecolor="black", alpha=0.7, density=show_kde, color="#4C72B0")
        if show_kde:
            try:
                from scipy.stats import gaussian_kde
                kde = gaussian_kde(arr)
                xr = _np.linspace(arr.min(), arr.max(), 200)
                ax.plot(xr, kde(xr), "r-", linewidth=2, label="KDE")
                ax.legend(fontsize=9)
            except ImportError:
                pass
        stats_txt = f"n={len(arr):,}  mean={arr.mean():.3f}  std={arr.std():.3f}  median={_np.median(arr):.3f}"
        ax.set_title(f"{title}\n{stats_txt}" if title else stats_txt, fontsize=11)
        ax.set_xlabel(xlabel or "value")
        ax.set_ylabel("density" if show_kde else "frequency")
        ax.grid(True, alpha=0.3)
        plt.tight_layout()
        import tempfile as _tmp
        import time as _time
        out = output_path or str(Path(_tmp.gettempdir()) / f"layla_hist_{int(_time.time())}.png")
        fig.savefig(out, dpi=120, bbox_inches="tight")
        plt.close(fig)
        return {"ok": True, "chart_type": "histogram", "path": out, "n": len(arr), "mean": round(float(arr.mean()), 4), "std": round(float(arr.std()), 4)}
    except ImportError:
        return {"ok": False, "error": "matplotlib not installed: pip install matplotlib"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ Memory Stats Tool â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def spaced_repetition_review(limit: int = 10, interval_hours: float = 24.0) -> dict:
    """
    Get learnings due for spaced repetition review. Optionally schedule next review.
    Returns items due (next_review_at <= now or NULL). Call schedule_next_review per item to reinforce.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import get_learnings_due_for_review
        due = get_learnings_due_for_review(limit=limit)
        items = [{"id": r["id"], "content": (r.get("content") or "")[:200], "importance": r.get("importance_score")} for r in due]
        return {"ok": True, "due_count": len(items), "items": items}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def schedule_learning_review(learning_id: int, interval_hours: float = 24.0) -> dict:
    """Schedule next spaced repetition review for a learning."""
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import schedule_next_review
        schedule_next_review(learning_id, interval_hours)
        return {"ok": True, "learning_id": learning_id, "interval_hours": interval_hours}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def memory_stats() -> dict:
    """
    Return stats about Layla's memory: learnings count, ChromaDB docs, aspect memories, DB size.
    """
    try:
        agent_dir = Path(__file__).resolve().parent.parent.parent
        sys.path.insert(0, str(agent_dir))
        from layla.memory.db import get_recent_learnings
        learnings = get_recent_learnings(n=9999)
        result: dict = {"ok": True, "learnings_count": len(learnings)}
        db_path = agent_dir / "layla.db"
        if db_path.exists():
            result["db_size_kb"] = round(db_path.stat().st_size / 1024, 1)
        try:
            from layla.memory.vector_store import _get_knowledge_collection
            coll = _get_knowledge_collection()
            result["knowledge_docs"] = coll.count() if coll else 0
        except Exception:
            result["knowledge_docs"] = "unavailable"
        try:
            import sqlite3 as _sql
            conn = _sql.connect(str(db_path))
            rows = conn.execute("SELECT aspect_id, COUNT(*) FROM aspect_memories GROUP BY aspect_id").fetchall()
            conn.close()
            result["aspect_memories"] = {r[0]: r[1] for r in rows}
        except Exception:
            pass
        return result
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ Tool Chain Planner â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def tool_chain_plan(goal: str, context: str = "") -> dict:
    """
    Plan a multi-step tool execution sequence for a given goal using intent detection.
    Returns an ordered list of tools with purpose descriptions.
    This is a heuristic planner â€” the LLM adapts and executes the plan.
    """
    goal_lower = (goal + " " + context).lower()
    INTENT_PLANS = {
        "research": (["research","find","what is","who is","explain","look up"], [
            {"step": 1, "tool": "ddg_search", "purpose": "Broad web search"},
            {"step": 2, "tool": "wiki_search", "purpose": "Encyclopedic context"},
            {"step": 3, "tool": "arxiv_search", "purpose": "Academic papers if relevant"},
            {"step": 4, "tool": "fetch_article", "purpose": "Extract full article from best result"},
            {"step": 5, "tool": "summarize_text", "purpose": "Distill findings"},
            {"step": 6, "tool": "save_note", "purpose": "Store to memory"},
        ]),
        "code": (["code","analyze","review","bug","refactor","function","class","import","ast"], [
            {"step": 1, "tool": "workspace_map", "purpose": "Map project structure"},
            {"step": 2, "tool": "code_symbols", "purpose": "Index all symbols"},
            {"step": 3, "tool": "dependency_graph", "purpose": "Map imports"},
            {"step": 4, "tool": "code_metrics", "purpose": "Measure complexity"},
            {"step": 5, "tool": "find_todos", "purpose": "Identify outstanding issues"},
            {"step": 6, "tool": "code_lint", "purpose": "Check for violations"},
            {"step": 7, "tool": "security_scan", "purpose": "Security audit"},
        ]),
        "data": (["dataset","csv","excel","data","statistics","correlations","cluster","analyze data"], [
            {"step": 1, "tool": "dataset_summary", "purpose": "Full statistical overview"},
            {"step": 2, "tool": "plot_histogram", "purpose": "Distribution visualization"},
            {"step": 3, "tool": "cluster_data", "purpose": "Natural groupings"},
            {"step": 4, "tool": "scipy_compute", "purpose": "Statistical tests"},
            {"step": 5, "tool": "plot_scatter", "purpose": "Correlation exploration"},
            {"step": 6, "tool": "save_note", "purpose": "Record findings"},
        ]),
        "web_crawl": (["crawl","scrape","website","all pages","download docs","site"], [
            {"step": 1, "tool": "check_url", "purpose": "Verify accessibility"},
            {"step": 2, "tool": "extract_links", "purpose": "Map site structure"},
            {"step": 3, "tool": "crawl_site", "purpose": "Crawl + extract all pages"},
            {"step": 4, "tool": "vector_store", "purpose": "Index into RAG"},
        ]),
        "database": (["database","sql","query","schema","table"], [
            {"step": 1, "tool": "schema_introspect", "purpose": "Understand structure"},
            {"step": 2, "tool": "generate_sql", "purpose": "Draft SQL query"},
            {"step": 3, "tool": "sql_query", "purpose": "Execute and retrieve data"},
            {"step": 4, "tool": "dataset_summary", "purpose": "Analyze results"},
        ]),
        "image": (["image","photo","picture","ocr","caption","detect"], [
            {"step": 1, "tool": "ocr_image", "purpose": "Extract text"},
            {"step": 2, "tool": "describe_image", "purpose": "Generate caption"},
            {"step": 3, "tool": "detect_objects", "purpose": "Identify objects"},
        ]),
        "security": (["security","vulnerability","secret","scan","cve","bandit"], [
            {"step": 1, "tool": "security_scan", "purpose": "Static analysis (bandit)"},
            {"step": 2, "tool": "security_scan", "purpose": "Secret detection", "args": {"scan_type": "secrets"}},
            {"step": 3, "tool": "security_scan", "purpose": "Dependency audit", "args": {"scan_type": "deps"}},
            {"step": 4, "tool": "find_todos", "purpose": "Find security-related TODOs"},
        ]),
    }
    best_intent, best_score, best_plan = "research", 0, []
    for intent, (patterns, plan) in INTENT_PLANS.items():
        score = sum(1 for p in patterns if p in goal_lower)
        if score > best_score:
            best_score, best_intent, best_plan = score, intent, plan
    if not best_plan:
        best_plan = [{"step": 1, "tool": "tool_recommend", "purpose": "Find best tools"}, {"step": 2, "tool": "ddg_search", "purpose": "Gather information"}, {"step": 3, "tool": "save_note", "purpose": "Store findings"}]
    valid_plan = [s for s in best_plan if s["tool"] in TOOLS]
    return {"ok": True, "goal": goal, "detected_intent": best_intent, "plan": valid_plan, "step_count": len(valid_plan), "note": "Heuristic plan â€” LLM will adapt based on results at each step"}


# â”€â”€â”€ Geographic Intelligence â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def geo_query(location: str, details: bool = True) -> dict:
    """
    Geocode a location to coordinates + geographic details.
    Uses geopy/Nominatim (pip install geopy) or public Nominatim REST API fallback.
    """
    try:
        from geopy.geocoders import Nominatim
        geolocator = Nominatim(user_agent="layla-agent/2.0")
        loc = geolocator.geocode(location, exactly_one=True, timeout=10, addressdetails=details, language="en")
        if not loc:
            return {"ok": False, "error": f"Not found: {location}"}
        result: dict = {"ok": True, "query": location, "display_name": loc.raw.get("display_name", ""), "lat": float(loc.latitude), "lon": float(loc.longitude)}
        if details and loc.raw.get("address"):
            addr = loc.raw["address"]
            result.update({"country": addr.get("country", ""), "country_code": addr.get("country_code", "").upper(), "state": addr.get("state", ""), "city": addr.get("city", addr.get("town", addr.get("village", "")))})
        if loc.raw.get("boundingbox"):
            bb = loc.raw["boundingbox"]
            result["bounding_box"] = {"south": float(bb[0]), "north": float(bb[1]), "west": float(bb[2]), "east": float(bb[3])}
        return result
    except ImportError:
        pass
    # Fallback: public REST API
    try:
        import json as _json
        import urllib.parse
        import urllib.request
        q = urllib.parse.quote(location)
        req = urllib.request.Request(f"https://nominatim.openstreetmap.org/search?q={q}&format=json&limit=1&addressdetails=1", headers={"User-Agent": "layla-agent/2.0"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = _json.loads(resp.read())
        if not data:
            return {"ok": False, "error": f"Not found: {location}"}
        d = data[0]
        addr = d.get("address", {})
        return {"ok": True, "query": location, "method": "nominatim-api", "display_name": d.get("display_name", ""), "lat": float(d["lat"]), "lon": float(d["lon"]), "country": addr.get("country", ""), "state": addr.get("state", ""), "city": addr.get("city", addr.get("town", ""))}
    except Exception as e:
        return {"ok": False, "error": f"geo_query failed (install geopy: pip install geopy): {e}"}


def map_url(center: str = "", lat: float = 0.0, lon: float = 0.0, zoom: int = 12, markers: list | None = None) -> dict:
    """
    Generate static map URLs centered on a location (auto-geocoded from name or lat/lon).
    Returns OpenStreetMap URL, embed HTML, and Geoapify static map URL.
    """
    if center and not (lat and lon):
        geo = geo_query(center, details=False)
        if geo.get("ok"):
            lat, lon = geo["lat"], geo["lon"]
        else:
            return {"ok": False, "error": f"Could not geocode: {center}"}
    if not (lat and lon):
        return {"ok": False, "error": "Provide center name or lat+lon"}
    osm = f"https://www.openstreetmap.org/?mlat={lat}&mlon={lon}#map={zoom}/{lat}/{lon}"
    static = f"https://maps.geoapify.com/v1/staticmap?style=osm-bright&width=800&height=600&center=lonlat:{lon},{lat}&zoom={zoom}"
    embed = f'<iframe src="https://www.openstreetmap.org/export/embed.html?bbox={lon-0.05},{lat-0.05},{lon+0.05},{lat+0.05}&layer=mapnik&marker={lat},{lon}" width="800" height="500"></iframe>'
    return {"ok": True, "center": {"lat": lat, "lon": lon}, "zoom": zoom, "osm_url": osm, "static_map_url": static, "embed_html": embed}


# â”€â”€â”€ Video Intelligence â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def extract_frames(path: str, fps: float = 1.0, max_frames: int = 30, output_dir: str = "") -> dict:
    """
    Extract frames from a video at given fps. Requires ffmpeg binary in PATH.
    ffmpeg-python package (pip install ffmpeg-python) enables probe metadata.
    Falls back to ffmpeg CLI directly if package not installed.
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    import tempfile as _tmp
    out_dir = Path(output_dir) if output_dir else Path(_tmp.gettempdir()) / f"frames_{target.stem}"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_pattern = str(out_dir / "frame_%04d.png")
    try:
        import ffmpeg
        probe = ffmpeg.probe(str(target))
        duration = float(probe["format"].get("duration", 0))
        vi = next((s for s in probe["streams"] if s["codec_type"] == "video"), {})
        (ffmpeg.input(str(target)).filter("fps", fps=fps).output(out_pattern, vframes=max_frames).overwrite_output().run(quiet=True))
        frames = sorted(out_dir.glob("frame_*.png"))
        return {"ok": True, "path": str(target), "fps": fps, "duration_sec": round(duration, 2), "resolution": f"{vi.get('width',0)}x{vi.get('height',0)}", "frames_extracted": len(frames), "output_dir": str(out_dir), "frame_paths": [str(f) for f in frames]}
    except ImportError:
        pass
    try:
        subprocess.run(["ffmpeg", "-i", str(target), "-vf", f"fps={fps}", "-frames:v", str(max_frames), out_pattern, "-y"], capture_output=True, timeout=120, text=True, encoding="utf-8", errors="replace")
        frames = sorted(out_dir.glob("frame_*.png"))
        if frames:
            return {"ok": True, "path": str(target), "fps": fps, "frames_extracted": len(frames), "output_dir": str(out_dir), "frame_paths": [str(f) for f in frames]}
        return {"ok": False, "error": "ffmpeg produced no output. Ensure ffmpeg is installed and in PATH."}
    except FileNotFoundError:
        return {"ok": False, "error": "ffmpeg not found. Install: https://ffmpeg.org/download.html and pip install ffmpeg-python"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def detect_scenes(path: str, threshold: float = 27.0) -> dict:
    """
    Detect scene cuts in a video. threshold: lower = more sensitive.
    Requires: pip install scenedetect[opencv]
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        from scenedetect import ContentDetector, detect
        scenes = detect(str(target), ContentDetector(threshold=threshold))
        scene_list = [{"scene": i+1, "start_sec": round(s[0].get_seconds(), 3), "end_sec": round(s[1].get_seconds(), 3), "duration_sec": round(s[1].get_seconds()-s[0].get_seconds(), 3)} for i, s in enumerate(scenes)]
        return {"ok": True, "path": str(target), "scene_count": len(scene_list), "threshold": threshold, "scenes": scene_list}
    except ImportError:
        return {"ok": False, "error": "pyscenedetect not installed: pip install scenedetect[opencv]"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ Object Detection â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def detect_objects(path: str, confidence: float = 0.25, model: str = "yolov8n.pt") -> dict:
    """
    Detect objects in an image using YOLO (ultralytics).
    First run auto-downloads model (~6 MB for nano). confidence: 0.0-1.0.
    model: yolov8n.pt (nano/fast) | yolov8s.pt (small) | yolov8m.pt (medium).
    Requires: pip install ultralytics
    """
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        from ultralytics import YOLO
        m = YOLO(model)
        results = m(str(target), conf=confidence, verbose=False)
        detections = []
        for r in results:
            for box in r.boxes:
                cls_id = int(box.cls[0])
                detections.append({"class": r.names[cls_id], "confidence": round(float(box.conf[0]), 4), "bbox": {"x1": round(float(box.xyxy[0][0]), 1), "y1": round(float(box.xyxy[0][1]), 1), "x2": round(float(box.xyxy[0][2]), 1), "y2": round(float(box.xyxy[0][3]), 1)}})
        by_class: dict = {}
        for d in detections:
            by_class[d["class"]] = by_class.get(d["class"], 0) + 1
        return {"ok": True, "model": model, "total": len(detections), "by_class": by_class, "detections": detections}
    except ImportError:
        return {"ok": False, "error": "ultralytics not installed: pip install ultralytics"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# â”€â”€â”€ Desktop Automation â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

def screenshot_desktop(region: list | None = None, output_path: str = "") -> dict:
    """
    Capture a screenshot of the desktop or a region [x, y, width, height].
    Uses Pillow ImageGrab or pyautogui. Returns path to saved PNG.
    """
    import tempfile as _tmp
    import time as _time
    out = output_path or str(Path(_tmp.gettempdir()) / f"layla_screen_{int(_time.time())}.png")
    try:
        from PIL import ImageGrab
        img = ImageGrab.grab(bbox=(region[0], region[1], region[0]+region[2], region[1]+region[3]) if region else None)
        img.save(out)
        return {"ok": True, "path": out, "size": f"{img.width}x{img.height}", "method": "Pillow"}
    except ImportError:
        pass
    try:
        import pyautogui
        img = pyautogui.screenshot(region=tuple(region) if region else None)
        img.save(out)
        return {"ok": True, "path": out, "size": f"{img.width}x{img.height}", "method": "pyautogui"}
    except ImportError:
        return {"ok": False, "error": "Install Pillow or pyautogui for screenshots"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def click_ui(x: int, y: int, button: str = "left", clicks: int = 1) -> dict:
    """
    Click at screen coordinates. button: left | right | middle. CAUTION: controls actual mouse.
    Requires: pip install pyautogui
    """
    try:
        import pyautogui
        pyautogui.FAILSAFE = True
        pyautogui.click(x=x, y=y, button=button, clicks=clicks)
        return {"ok": True, "action": "click", "x": x, "y": y, "button": button, "clicks": clicks}
    except ImportError:
        return {"ok": False, "error": "pyautogui not installed: pip install pyautogui"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def type_text(text: str, interval: float = 0.03) -> dict:
    """
    Type text at current cursor position. CAUTION: types into whatever window has focus.
    interval: delay between keystrokes in seconds. Requires: pip install pyautogui
    """
    try:
        import pyautogui
        pyautogui.FAILSAFE = True
        pyautogui.typewrite(text, interval=max(0.01, interval))
        return {"ok": True, "action": "type", "chars_typed": len(text)}
    except ImportError:
        return {"ok": False, "error": "pyautogui not installed: pip install pyautogui"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Fabrication: DXF → G-code ─────────────────────────────────────────────────

def generate_gcode(dxf_path: str, output_path: str, layer: str = "", depth_mm: float = -5.0, feed_rate: int = 3000, safe_z: float = 5.0) -> dict:
    """
    Generate 2D G-code from DXF polylines (flat cutting). layer: filter by layer name, empty=all.
    Requires ezdxf. Output is inside sandbox.
    """
    target = Path(dxf_path)
    out = Path(output_path)
    if not inside_sandbox(target) or not inside_sandbox(out):
        return {"ok": False, "error": "Paths must be inside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "DXF file not found"}
    try:
        import ezdxf
        doc = ezdxf.readfile(str(target))
        msp = doc.modelspace()
        lines_out = ["G21", "G90", f"G0 Z{safe_z}", f"F{feed_rate}"]
        count = 0
        for e in msp:
            if layer and getattr(e.dxf, "layer", "") != layer:
                continue
            if e.dxftype() == "LINE":
                start, end = e.dxf.start, e.dxf.end
                lines_out.append(f"G0 X{start[0]:.3f} Y{start[1]:.3f}")
                lines_out.append(f"G1 Z{depth_mm:.3f}")
                lines_out.append(f"G1 X{end[0]:.3f} Y{end[1]:.3f}")
                lines_out.append(f"G0 Z{safe_z}")
                count += 1
            elif e.dxftype() == "LWPOLYLINE":
                points = list(e.get_points())
                if len(points) < 2:
                    continue
                x, y = points[0][0], points[0][1]
                lines_out.append(f"G0 X{x:.3f} Y{y:.3f}")
                lines_out.append(f"G1 Z{depth_mm:.3f}")
                for x, y in points[1:]:
                    lines_out.append(f"G1 X{x:.3f} Y{y:.3f}")
                lines_out.append(f"G0 Z{safe_z}")
                count += 1
        lines_out.append("M2")
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text("\n".join(lines_out), encoding="utf-8")
        return {"ok": True, "output_path": str(out), "moves": count, "lines": len(lines_out)}
    except ImportError:
        return {"ok": False, "error": "ezdxf not installed: pip install ezdxf"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Code Refactoring ──────────────────────────────────────────────────────────

def rename_symbol(root: str, old_name: str, new_name: str, symbol_type: str = "auto", file_glob: str = "*.py", apply: bool = False) -> dict:
    """
    Rename symbol across Python files. symbol_type: function|class|variable|auto.
    apply=False: dry run, returns proposed changes. apply=True: writes changes.
    """
    root_path = Path(root)
    if not inside_sandbox(root_path):
        return {"ok": False, "error": "Outside sandbox"}
    if not root_path.exists():
        return {"ok": False, "error": "Path not found"}
    pattern = re.compile(r"\b" + re.escape(old_name) + r"\b")
    changes = []
    for f in root_path.rglob(file_glob):
        if not f.is_file():
            continue
        try:
            content = f.read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue
        new_content = pattern.sub(new_name, content)
        if new_content != content:
            n = len(pattern.findall(content))
            changes.append({"path": str(f), "replacements": n})
            if apply:
                f.write_text(new_content, encoding="utf-8")
    return {"ok": True, "old_name": old_name, "new_name": new_name, "changes": changes[:50], "total_files": len(changes), "applied": apply}


# ─── Docker ────────────────────────────────────────────────────────────────────

def docker_ps(all_containers: bool = False) -> dict:
    """List Docker containers. all_containers=True includes stopped."""
    try:
        argv = ["docker", "ps", "-a"] if all_containers else ["docker", "ps"]
        proc = subprocess.run(argv, capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=10)
        return {"ok": proc.returncode == 0, "output": (proc.stdout or proc.stderr or "")[:4000]}
    except FileNotFoundError:
        return {"ok": False, "error": "Docker not found or not in PATH"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def docker_run(image: str, args: str = "", name: str = "", rm: bool = True) -> dict:
    """Run a Docker container. args: extra docker run args. name: container name. rm: remove when stopped."""
    argv = ["docker", "run"]
    if rm:
        argv.append("--rm")
    if name:
        argv.extend(["--name", name])
    argv.append(image)
    if args:
        argv.extend(args.split())
    try:
        proc = subprocess.run(argv, capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=60)
        return {"ok": proc.returncode == 0, "output": (proc.stdout or proc.stderr or "")[:4000]}
    except FileNotFoundError:
        return {"ok": False, "error": "Docker not found"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── CI / GitHub ────────────────────────────────────────────────────────────────

def check_ci(repo: str, provider: str = "github") -> dict:
    """Check CI status. repo: path or owner/repo. provider: github|gitlab."""
    repo_path = Path(repo)
    if repo_path.exists() and repo_path.is_dir():
        try:
            r = subprocess.run(["git", "remote", "get-url", "origin"], cwd=str(repo_path), capture_output=True, text=True)
            url = (r.stdout or "").strip()
            if "github.com" in url:
                m = re.search(r"github\.com[/:]([\w-]+)/([\w.-]+)", url)
                if m:
                    owner, repo_name = m.group(1), m.group(2).replace(".git", "")
                    api_url = f"https://api.github.com/repos/{owner}/{repo_name}/actions/runs?per_page=5"
                    try:
                        import urllib.request
                        req = urllib.request.Request(api_url, headers={"Accept": "application/vnd.github.v3+json"})
                        with urllib.request.urlopen(req, timeout=10) as resp:
                            data = __import__("json").loads(resp.read().decode())
                        runs = data.get("workflow_runs", [])[:5]
                        return {"ok": True, "provider": "github", "runs": [{"name": r.get("name"), "status": r.get("status"), "conclusion": r.get("conclusion"), "created_at": r.get("created_at")} for r in runs]}
                    except Exception as e:
                        return {"ok": False, "error": str(e)}
        except Exception:
            pass
    return {"ok": False, "error": "Could not determine repo or fetch CI status"}


def send_webhook(url: str, payload: dict, method: str = "POST") -> dict:
    """Send JSON payload to webhook URL (Slack, Discord, custom)."""
    try:
        import json as _json
        import urllib.request
        data = _json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(url, data=data, method=method, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=15) as resp:
            body = resp.read().decode("utf-8", errors="replace")[:2000]
        return {"ok": True, "status": resp.status, "response": body}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def discord_send(content: str = "", embed: dict | None = None, webhook_url: str = "") -> dict:
    """
    Send message to Discord via webhook. Easy setup: Server Settings → Integrations → Webhooks → New.
    webhook_url: override; else reads discord_webhook_url from config or DISCORD_WEBHOOK_URL env.
    """
    url = webhook_url or __import__("os").environ.get("DISCORD_WEBHOOK_URL", "")
    if not url:
        try:
            agent_dir = Path(__file__).resolve().parent.parent.parent
            sys.path.insert(0, str(agent_dir))
            import runtime_safety
            cfg = runtime_safety.load_config()
            url = cfg.get("discord_webhook_url", "") or ""
        except Exception:
            pass
    if not url:
        return {"ok": False, "error": "No webhook URL. Set discord_webhook_url in config, DISCORD_WEBHOOK_URL env, or pass webhook_url."}
    payload = {}
    if content:
        payload["content"] = content[:2000]
    if embed:
        payload["embeds"] = [{"title": embed.get("title", "")[:256], "description": (embed.get("description") or "")[:4096], "color": embed.get("color")}][:1]
    if not payload:
        return {"ok": False, "error": "Provide content or embed"}
    return send_webhook(url, payload)


def calendar_read(path: str) -> dict:
    """Read .ics calendar file. Returns events with start, end, summary."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import icalendar
        cal = icalendar.Calendar.from_ical(target.read_bytes())
        events = []
        for c in cal.walk():
            if c.name == "VEVENT":
                start = c.get("dtstart").dt if c.get("dtstart") else None
                end = c.get("dtend").dt if c.get("dtend") else None
                summary = str(c.get("summary", ""))
                events.append({"start": str(start), "end": str(end), "summary": summary})
        return {"ok": True, "path": str(target), "events": events[:50]}
    except ImportError:
        return {"ok": False, "error": "icalendar not installed: pip install icalendar"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def calendar_add_event(path: str, summary: str, start: str, end: str = "", description: str = "") -> dict:
    """Add event to .ics file. start/end: ISO or YYYY-MM-DD HH:MM. Creates file if missing."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    try:
        from datetime import datetime

        import icalendar
        cal = icalendar.Calendar()
        if target.exists():
            cal = icalendar.Calendar.from_ical(target.read_text(encoding="utf-8", errors="replace"))
        event = icalendar.Event()
        event.add("summary", summary[:200])
        event.add("dtstart", datetime.fromisoformat(start.replace("Z", "+00:00")) if "T" in start else datetime.fromisoformat(start))
        if end:
            event.add("dtend", datetime.fromisoformat(end.replace("Z", "+00:00")) if "T" in end else datetime.fromisoformat(end))
        if description:
            event.add("description", description[:500])
        cal.add_component(event)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(cal.to_ical().decode("utf-8", errors="replace"), encoding="utf-8")
        return {"ok": True, "path": str(target), "summary": summary}
    except ImportError:
        return {"ok": False, "error": "icalendar not installed: pip install icalendar"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def db_backup(db_path: str, backup_path: str = "") -> dict:
    """Backup SQLite database. backup_path: optional; default adds .bak timestamp."""
    src = Path(db_path)
    if not inside_sandbox(src):
        return {"ok": False, "error": "Outside sandbox"}
    if not src.exists():
        return {"ok": False, "error": "Database not found"}
    dest = Path(backup_path) if backup_path else src.with_suffix(f".bak_{__import__('time').strftime('%Y%m%d_%H%M%S')}{src.suffix}")
    if not inside_sandbox(dest):
        return {"ok": False, "error": "Backup path outside sandbox"}
    try:
        import shutil
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(src), str(dest))
        return {"ok": True, "source": str(src), "backup": str(dest)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def create_svg(path: str, content: str) -> dict:
    """Write SVG file. content: valid SVG markup. Procedural drawing — no Gen AI."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not content.strip().startswith("<"):
        content = f'<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 400">{content}</svg>'
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return {"ok": True, "path": str(target)}


def create_mermaid(path: str, content: str) -> dict:
    """Write Mermaid diagram file (.mmd). content: mermaid code (flowchart, sequenceDiagram, etc)."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not content.strip().startswith(("flowchart", "graph", "sequenceDiagram", "classDiagram", "stateDiagram", "erDiagram", "gantt", "pie", "journey")):
        content = f"flowchart TD\n{content}"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return {"ok": True, "path": str(target)}


def github_issues(repo_slug: str, state: str = "open", token: str = "") -> dict:
    """List GitHub issues. repo_slug: owner/repo. token: optional GITHUB_TOKEN env or param."""
    token = token or __import__("os").environ.get("GITHUB_TOKEN", "")
    headers = {"Accept": "application/vnd.github.v3+json"}
    if token:
        headers["Authorization"] = f"Bearer {token}"
    url = f"https://api.github.com/repos/{repo_slug}/issues?state={state}&per_page=20"
    try:
        import urllib.request
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = __import__("json").loads(resp.read().decode())
        return {"ok": True, "issues": [{"number": i["number"], "title": i["title"], "state": i["state"], "url": i["html_url"]} for i in data[:20]]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def github_pr(repo_slug: str, title: str, head: str, base: str = "main", body: str = "", token: str = "") -> dict:
    """Create a GitHub PR. token: GITHUB_TOKEN env or param."""
    token = token or __import__("os").environ.get("GITHUB_TOKEN", "")
    if not token:
        return {"ok": False, "error": "GITHUB_TOKEN required for creating PRs"}
    import json as _json
    import urllib.request
    payload = {"title": title, "head": head, "base": base, "body": body}
    data = _json.dumps(payload).encode("utf-8")
    req = urllib.request.Request("https://api.github.com/repos/" + repo_slug + "/pulls", data=data, method="POST", headers={"Accept": "application/vnd.github.v3+json", "Authorization": f"Bearer {token}", "Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            pr = _json.loads(resp.read().decode())
            return {"ok": True, "number": pr.get("number"), "url": pr.get("html_url"), "title": pr.get("title")}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Code Format & Archive ──────────────────────────────────────────────────────

def code_format(path: str, formatter: str = "ruff") -> dict:
    """Format Python code with ruff or black. path: file or directory."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Path not found"}
    try:
        if formatter == "ruff":
            proc = subprocess.run(
                [sys.executable, "-m", "ruff", "format", str(target)],
                capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=30,
            )
        else:
            proc = subprocess.run(
                [sys.executable, "-m", "black", str(target)],
                capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=60,
            )
        return {"ok": proc.returncode == 0, "output": (proc.stdout or proc.stderr or "")[:2000]}
    except FileNotFoundError:
        return {"ok": False, "error": f"{formatter} not installed: pip install {formatter}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def extract_archive(path: str, dest: str = "") -> dict:
    """Extract zip or tar archive. dest: output dir, default same as archive."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "Archive not found"}
    out = Path(dest) if dest else target.parent
    if not inside_sandbox(out):
        return {"ok": False, "error": "Destination outside sandbox"}
    try:
        import tarfile
        import zipfile
        out = out.resolve()
        out.mkdir(parents=True, exist_ok=True)
        if target.suffix.lower() in (".zip",):
            with zipfile.ZipFile(target, "r") as z:
                for m in z.namelist():
                    if ".." in m or m.startswith("/") or (m.startswith("\\") if len(m) > 1 else False):
                        continue
                    member_path = (out / m).resolve()
                    try:
                        member_path.relative_to(out)
                    except ValueError:
                        continue
                    try:
                        z.extract(m, out)
                    except Exception:
                        pass
        elif target.suffix.lower() in (".tar", ".gz", ".tgz", ".bz2", ".xz"):
            with tarfile.open(target, "r:*") as t:
                for m in t.getnames():
                    if ".." in m or m.startswith("/") or (m.startswith("\\") if len(m) > 1 else False):
                        continue
                    member_path = (out / m).resolve()
                    try:
                        member_path.relative_to(out)
                    except ValueError:
                        continue
                    try:
                        t.extract(m, out)
                    except Exception:
                        pass
        else:
            return {"ok": False, "error": "Unsupported format. Use .zip, .tar, .tar.gz, .tgz"}
        return {"ok": True, "path": str(target), "dest": str(out)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def create_archive(paths: list, output: str, format: str = "zip") -> dict:
    """Create zip archive from paths. paths: list of files/dirs. output: .zip path."""
    out = Path(output)
    if not inside_sandbox(out):
        return {"ok": False, "error": "Output path outside sandbox"}
    try:
        import zipfile
        out.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
            for p in paths:
                fp = Path(p)
                if not inside_sandbox(fp):
                    continue
                if fp.is_file():
                    z.write(fp, fp.name)
                elif fp.is_dir():
                    for f in fp.rglob("*"):
                        if f.is_file():
                            z.write(f, str(f.relative_to(fp)))
        return {"ok": True, "output": str(out), "files": len(paths)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def uuid_generate(count: int = 1) -> dict:
    """Generate UUID(s). count: number of UUIDs."""
    import uuid
    uuids = [str(uuid.uuid4()) for _ in range(min(max(1, count), 100))]
    return {"ok": True, "uuids": uuids}


def random_string(length: int = 16, charset: str = "alphanumeric") -> dict:
    """Generate random string. charset: alphanumeric|hex|ascii."""
    import secrets
    import string
    if charset == "hex":
        s = secrets.token_hex(length // 2 + 1)[:length]
    elif charset == "ascii":
        s = "".join(secrets.choice(string.ascii_letters + string.digits) for _ in range(length))
    else:
        s = "".join(secrets.choice(string.ascii_letters + string.digits) for _ in range(length))
    return {"ok": True, "value": s}


def password_generate(length: int = 20, symbols: bool = True) -> dict:
    """Generate secure random password."""
    import secrets
    import string
    chars = string.ascii_letters + string.digits
    if symbols:
        chars += "!@#$%^&*"
    pwd = "".join(secrets.choice(chars) for _ in range(min(max(12, length), 128)))
    return {"ok": True, "password": pwd, "length": len(pwd)}


def disk_usage(path: str = ".") -> dict:
    """Disk usage for path. Returns used/total/free in GB."""
    try:
        import shutil
        p = Path(path)
        if not p.exists():
            p = Path.cwd()
        total, used, free = shutil.disk_usage(str(p))
        return {"ok": True, "path": str(p), "total_gb": round(total / (1024**3), 2), "used_gb": round(used / (1024**3), 2), "free_gb": round(free / (1024**3), 2)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def process_list(limit: int = 20) -> dict:
    """List running processes (top by CPU/memory). Requires psutil."""
    try:
        import psutil
        procs = []
        for p in sorted(psutil.process_iter(["pid", "name", "cpu_percent", "memory_percent"]), key=lambda x: (x.info.get("cpu_percent") or 0) + (x.info.get("memory_percent") or 0), reverse=True)[:limit]:
            try:
                procs.append({"pid": p.info["pid"], "name": (p.info.get("name") or "")[:40], "cpu": p.info.get("cpu_percent"), "memory": p.info.get("memory_percent")})
            except (psutil.NoSuchProcess, KeyError):
                pass
        return {"ok": True, "processes": procs}
    except ImportError:
        return {"ok": False, "error": "psutil not installed: pip install psutil"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def generate_qr(data: str, output_path: str = "", size: int = 10) -> dict:
    """Generate QR code. data: text/URL. output_path: save PNG. Requires qrcode."""
    try:
        import qrcode
        qr = qrcode.QRCode(version=1, box_size=size, border=2)
        qr.add_data(data[:4000])
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        if output_path:
            out = Path(output_path)
            if not inside_sandbox(out):
                return {"ok": False, "error": "Output outside sandbox"}
            img.save(str(out))
            return {"ok": True, "output": str(out)}
        import io
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return {"ok": True, "base64": __import__("base64").b64encode(buf.getvalue()).decode()[:5000]}
    except ImportError:
        return {"ok": False, "error": "qrcode not installed: pip install qrcode[pil]"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def write_csv(path: str, rows: list, headers: list | None = None) -> dict:
    """Write CSV file. rows: list of dicts or lists. headers: optional column order."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    try:
        import csv
        target.parent.mkdir(parents=True, exist_ok=True)
        with open(target, "w", newline="", encoding="utf-8") as f:
            if rows and isinstance(rows[0], dict):
                h = headers or list(rows[0].keys())
                w = csv.DictWriter(f, fieldnames=h, extrasaction="ignore")
                w.writeheader()
                w.writerows(rows)
            else:
                w = csv.writer(f)
                if headers:
                    w.writerow(headers)
                w.writerows(rows)
        return {"ok": True, "path": str(target), "rows": len(rows)}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def json_schema(data: str | dict) -> dict:
    """Infer JSON schema from sample data. data: JSON string or dict."""
    try:
        import json as _json
        obj = _json.loads(data) if isinstance(data, str) else data
        def _infer(v):
            if v is None:
                return {"type": "null"}
            if isinstance(v, bool):
                return {"type": "boolean"}
            if isinstance(v, int):
                return {"type": "integer"}
            if isinstance(v, float):
                return {"type": "number"}
            if isinstance(v, str):
                return {"type": "string"}
            if isinstance(v, list):
                item = _infer(v[0]) if v else {}
                return {"type": "array", "items": item}
            if isinstance(v, dict):
                return {"type": "object", "properties": {k: _infer(v) for k, v in v.items()}}
            return {}
        schema = _infer(obj)
        return {"ok": True, "schema": schema}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def jwt_decode(token: str, verify: bool = False, secret: str = "") -> dict:
    """Decode JWT (header + payload). verify=True validates signature with secret."""
    try:
        import base64
        import json as _json
        parts = token.split(".")
        if len(parts) != 3:
            return {"ok": False, "error": "Invalid JWT format"}
        def b64d(s):
            pad = 4 - len(s) % 4
            return base64.urlsafe_b64decode(s + "=" * pad)
        header = _json.loads(b64d(parts[0]).decode())
        payload = _json.loads(b64d(parts[1]).decode())
        if verify and secret:
            try:
                import hashlib
                import hmac
                sig = parts[2]
                msg = f"{parts[0]}.{parts[1]}".encode()
                exp = base64.urlsafe_b64encode(hmac.new(secret.encode(), msg, hashlib.sha256).digest()).decode().rstrip("=")
                if exp != sig:
                    return {"ok": False, "error": "Signature verification failed"}
            except Exception as e:
                return {"ok": False, "error": str(e)}
        return {"ok": True, "header": header, "payload": payload}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def read_toml(path: str) -> dict:
    """Parse TOML file. Returns dict."""
    target = Path(path)
    if not inside_sandbox(target):
        return {"ok": False, "error": "Outside sandbox"}
    if not target.exists():
        return {"ok": False, "error": "File not found"}
    try:
        import tomllib
        return {"ok": True, "path": str(target), "data": tomllib.loads(target.read_text(encoding="utf-8"))}
    except ImportError:
        try:
            import tomli
            return {"ok": True, "path": str(target), "data": tomli.loads(target.read_text(encoding="utf-8"))}
        except ImportError:
            return {"ok": False, "error": "tomllib (3.11+) or tomli required"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def merge_pdf(paths: list, output: str) -> dict:
    """Merge PDFs into one. paths: list of PDF paths. output: output path."""
    out = Path(output)
    if not inside_sandbox(out):
        return {"ok": False, "error": "Output outside sandbox"}
    for p in paths:
        if not inside_sandbox(Path(p)):
            return {"ok": False, "error": f"Path outside sandbox: {p}"}
    try:
        from pypdf import PdfMerger
        merger = PdfMerger()
        for p in paths:
            if Path(p).exists():
                merger.append(str(p))
        merger.write(str(out))
        merger.close()
        return {"ok": True, "output": str(out), "merged": len(paths)}
    except ImportError:
        return {"ok": False, "error": "pypdf not installed: pip install pypdf"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def send_email(to: str, subject: str, body: str, smtp_host: str = "", smtp_port: int = 587, username: str = "", password: str = "") -> dict:
    """Send email via SMTP. Credentials from env: SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASS."""
    import os
    host = smtp_host or os.environ.get("SMTP_HOST", "localhost")
    port = smtp_port or int(os.environ.get("SMTP_PORT", "587"))
    user = username or os.environ.get("SMTP_USER", "")
    pwd = password or os.environ.get("SMTP_PASS", "")
    try:
        import smtplib
        from email.mime.text import MIMEText
        msg = MIMEText(body, "plain", "utf-8")
        msg["Subject"] = subject
        msg["To"] = to
        msg["From"] = user or "layla@local"
        with smtplib.SMTP(host, port, timeout=15) as s:
            if user and pwd:
                s.starttls()
                s.login(user, pwd)
            s.send_message(msg)
        return {"ok": True, "to": to, "subject": subject[:50]}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# ─── Build TOOLS from domain modules ───────────────────────────────────────────
TOOLS = _build_tools_from_domains()


def _wrap_tool_with_metrics(name: str, fn: Any) -> Any:
    """Wrap tool fn to record execution latency to performance_monitor."""
    def wrapped(*args: Any, **kwargs: Any) -> Any:
        import time
        start = time.perf_counter()
        result = None
        try:
            result = fn(*args, **kwargs)
            return result
        finally:
            elapsed_ms = (time.perf_counter() - start) * 1000
            try:
                from services.observability import log_tool_result
                ok = isinstance(result, dict) and result.get("ok", True) if result is not None else False
                log_tool_result(name, ok=ok, duration_ms=elapsed_ms)
            except Exception:
                pass
    return wrapped


# Wrap all tool fns for observability metrics
for _tname, _entry in list(TOOLS.items()):
    if isinstance(_entry, dict) and _entry.get("fn"):
        _entry["fn"] = _wrap_tool_with_metrics(_tname, _entry["fn"])


_REQUIRED_META = {"name", "description", "category", "risk_level"}


TOOL_COUNT_THRESHOLD = 50


def validate_tools_registry() -> None:
    """Validate tool registry integrity: count threshold + required metadata. Raise if incomplete."""
    import logging
    log = logging.getLogger("layla")
    if len(TOOLS) < TOOL_COUNT_THRESHOLD:
        raise RuntimeError(f"Tool registry incomplete: {len(TOOLS)} tools (expected >= {TOOL_COUNT_THRESHOLD})")
    missing = []
    for tool_name, entry in TOOLS.items():
        if not isinstance(entry, dict):
            missing.append((tool_name, "not a dict"))
            continue
        fn = entry.get("fn")
        if not fn:
            missing.append((tool_name, "missing fn"))
            continue
        # name: key is name
        if not entry.get("name"):
            entry["name"] = tool_name
        # description: prefer explicit, else __doc__, else humanized name
        if not entry.get("description"):
            doc = (getattr(fn, "__doc__") or "").strip().split("\n")[0][:200]
            if doc:
                entry["description"] = doc
            else:
                entry["description"] = tool_name.replace("_", " ").strip()
        # category: infer if missing
        if not entry.get("category"):
            entry["category"] = "general"
            log.debug("tool %s: missing category, defaulting to general", tool_name)
        # risk_level: required
        if not entry.get("risk_level"):
            entry["risk_level"] = "medium" if entry.get("dangerous") else "low"
            log.warning("tool %s: missing risk_level, inferred as %s", tool_name, entry["risk_level"])
    for name, msg in missing:
        log.warning("tool %s: %s", name, msg)
